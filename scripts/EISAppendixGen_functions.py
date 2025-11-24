import pandas as pd
import numpy as np
import docx
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
import matplotlib.pyplot as plt
import calendar
import os
from time import strptime
from storage_to_elevation import storage_to_elevation
from ec_to_cl import ec_to_cl
from math import floor, ceil
from datetime import datetime, timedelta
from docx_caption_formatter import add_caption_byfield, add_caption_water_supply
import shutil
import copy
import subprocess
from pydsstools.heclib.dss import HecDss

pd.options.mode.chained_assignment = None

def get_locations(location_crosswalk_path, fields):
    """
    Gets location names from field codes passed

    Parameters
    ----------
    location_crosswalk_path: string
        Path and file name for xlsx file containing location names and field codes
    fields: list of strings
        Names of the fields to be processed

    Returns
    ----------
    locations : list of str
        List of the location titles from the crosswalk file given in location_crosswalk_path

    """
    #Read in crosswalk as a df
    crosswalk = pd.read_excel(location_crosswalk_path)

    #Look up each field code's corresponding location title and add to a list
    locations = []
    for i, field in enumerate(fields):
        if type(field) ==str:
            locations.append(crosswalk.loc[crosswalk["DSSPartB"] == field, "Location (Title)"].values[0])
        elif type(field)==tuple:
            if i>0 and field == fields[i-1]:
                use_index = 1
            else:
                use_index = 0
            #If filter_search is provided, then filter the parameter column of the dataframe. This is used to get the elevation figure/table titles, since they're under the same dsspartb as the storage titles.
            locations.append(crosswalk.loc[(crosswalk['DSSPartB'] == field[0]) &(crosswalk.Parameter == field[1]), "Location (Title)"].values[use_index])

    return locations

def get_locations_params(location_crosswalk_path, fields):
    """
    Gets location names from field codes passed

    Parameters
    ----------
    location_crosswalk_path: string
        Path and file name for xlsx file containing location names and field codes
    fields: list of strings
        Names of the fields to be processed

    Returns
    ----------
    locations: list of str
        list of the parameter names corresponding to the locations, based on what's in the crosswalk file.

    """
    #Read in crosswalk as a df
    crosswalk = pd.read_excel(location_crosswalk_path)

    #Look up each field code's corresponding location title and add to a list
    locations = []
    for i, field in enumerate(fields):
        if type(field) ==str:
            locations.append(crosswalk.loc[crosswalk["DSSPartB"] == field, "Parameter"].values[0])
        elif type(field)==tuple:
            if i>0 and field == fields[i-1]:
                use_index = 1
            else:
                use_index = 0
            #If filter_search is provided, then filter the parameter column of the dataframe. This is used to get the elevation figure/table titles, since they're under the same dsspartb as the storage titles.
            locations.append(crosswalk.loc[(crosswalk['DSSPartB'] == field[0]) &(crosswalk.Parameter == field[1]), "Parameter"].values[use_index])

    return locations

def get_location_wytypes(location_crosswalk_path, fields):
    """
    Gets location names from field codes passed

    Parameters
    ----------
    location_crosswalk_path: string
        Path and file name for xlsx file containing location names and field codes
    fields: list of strings
        Names of the fields to be processed

    Returns
    ---------
    wytype_list:  list of str
        list of water year types corresponding to each location, based on what's in the crosswalk file.

    """
    #Read in crosswalk as a df
    crosswalk = pd.read_excel(location_crosswalk_path)

    #Look up each field code's corresponding location title and add to a list
    wytype_list = []
    for i, field in enumerate(fields):
        if type(field) ==str:
            wytype_list.append(crosswalk.loc[crosswalk["DSSPartB"] == field, "Water Year Type Index"].values[0])
        elif type(field)==tuple:
            if i>0 and field == fields[i-1]:
                use_index = 1
            else:
                use_index = 0
            #If filter_search is provided, then filter the parameter column of the dataframe. This is used to get the elevation figure/table titles, since they're under the same dsspartb as the storage titles.
            wytype_list.append(crosswalk.loc[(crosswalk['DSSPartB'] == field[0]) &(crosswalk.Parameter == field[1]), "Water Year Type Index"].values[use_index])

    return wytype_list

def calculate_supply_fields(s_inputs, s_formulas, s_wy_flags_path):
    """
    Reads in data and calculated the specific categories for the water supply appendix
    Parameters
    ----------
    s_inputs: str
        Path for DSS inputs
    s_formulas: str
        Path for the excel sheet with all the formulas
    s_wy_flags_path: path
        Path to the WYTs

    Returns
    -------
    df_final: dataframe
        Data for the tables
    df_exceedances: dataframe
        Dataframe with the exceedances
    """

    # DSS data
    df_inputs = pd.read_excel(s_inputs)

    # replace 'Baseline' with 'NAA'
    df_inputs.replace({'Baseline': 'NAA'}, inplace=True)

    # Formulas for calculated fields
    df_formulas = pd.read_excel(s_formulas, sheet_name='annual')

    # Will hold the outputs
    df_output = pd.DataFrame(index=pd.MultiIndex.from_product([df_inputs['Scenario'].unique(), df_inputs['Year'].unique()]))

    # Go through each sub field and calculate it
    for row_index, row in df_formulas.iterrows():
        s_formula = row['Formula']
        sl_add_fields = [field.strip() for field in s_formula.split(' + ')]
        s_stat = 'sum' if row['Statistic'] == 'Sum' else 'mean'
        if row['Resolution'] == 'Annual':
            if row['Annual_ Month_ Range'] == 'MarFeb':
                df_output.loc[:, row['Title']] = df_inputs.groupby(['Scenario', 'DY'])[sl_add_fields].agg(s_stat).agg(s_stat, axis=1)
            elif row['Annual_ Month_ Range'] == 'JanDec':
                df_output.loc[:, row['Title']] = df_inputs.groupby(['Scenario', 'Year'])[sl_add_fields].agg(s_stat).agg(s_stat, axis=1)
            elif row['Annual_ Month_ Range'] == 'OctSep':
                df_output.loc[:, row['Title']] = df_inputs.groupby(['Scenario', 'WY'])[sl_add_fields].agg(s_stat).agg(s_stat, axis=1)
        elif row['Resolution'] == 'SingleMonth':
            i_month = list(calendar.month_abbr).index(row['Annual_ Month_ Range'])
            df_output.loc[:, row['Title']] = df_inputs[df_inputs['Month'] == i_month].groupby(['Scenario', 'Year'])[sl_add_fields].agg(s_stat).agg(s_stat, axis=1)

    # drop 1921 partial year
    df_output.drop(index=[1921], level=1, inplace=True)

    # formulas for the final categories
    df_final_formulas = pd.read_excel(s_formulas, sheet_name='final', index_col=[0, 1])

    # data to hold the table data, skip 2021 since none have that full year
    df_final = pd.DataFrame(index=df_output.drop(index=[2021], level=1).index, columns=df_final_formulas.index)

    # go through each final formula and calculate it
    for row_index, row in df_final_formulas.iterrows():

        # Pull out description and units
        s_description = row['Description']
        s_units = row['Units']

        # remove description and units
        row = row.iloc[2:]

        # Fields to add up
        ls_fields = row[~row.isna()].values
        if len(ls_fields) == 0:
            # add in description and units
            df_final.loc['Description', row_index] = s_description
            df_final.loc['Units', row_index] = s_units
            continue

        # add them up and insert into final data frame
        df_final[row_index] = df_output[ls_fields].sum(axis=1)

        # add in description and units
        df_final.loc['Description', row_index] = s_description
        df_final.loc['Units', row_index] = s_units

    # save the descriptions and fields
    df_temp = df_final.loc[['Description', 'Units'], :]

    # calculate the fields that are not in the formulas
    df_final[('Total For All Regions', 'Total Supplies')] = df_final[['Sacramento River Hydrologic Region', 'San Joaquin River Hydrologic Region (not including Friant-Kern and Madera Canal water users)',
                                                                      'San Francisco Bay Hydrologic Region', 'Central Coast Hydrologic Region', 'Tulare Lake Hydrologic Region (not including Friant-Kern Canal water users)',
                                                                      'South Lahontan Hydrologic Region', 'South Coast Hydrologic Region']].sum(axis=1)

    df_final[('North of Delta', 'SWP Ag')] = df_final[('North of Delta', 'SWP Ag')].fillna(0)
    df_final[('Total CVP North of Delta', 'Total CVP Ag and M&I')] = df_final[[('North of Delta', 'CVP Ag'), ('North of Delta', 'CVP M&I')]].sum(axis=1)
    df_final[('Total SWP North of Delta', 'Total SWP Ag and M&I')] = df_final[[('North of Delta', 'SWP Ag'), ('North of Delta', 'SWP M&I')]].sum(axis=1)
    df_final[('Total North of Delta', 'Total Ag and M&I Deliveries')] = df_final[[('Total CVP North of Delta', 'Total CVP Ag and M&I'), ('Total SWP North of Delta', 'Total SWP Ag and M&I')]].sum(axis=1)
    df_final[('Total CVP South of Delta', 'Total CVP Ag and M&I')] = df_final[[('South of Delta', 'CVP Ag'), ('South of Delta', 'CVP M&I')]].sum(axis=1)
    df_final[('Total SWP South of Delta', 'Total SWP Ag and M&I')] = df_final[[('South of Delta', 'SWP Ag'), ('South of Delta', 'SWP M&I')]].sum(axis=1)
    df_final[('Total South of Delta', 'Total Ag and M&I Deliveries')] = df_final[[('Total CVP South of Delta', 'Total CVP Ag and M&I'), ('Total SWP South of Delta', 'Total SWP Ag and M&I')]].sum(axis=1)

    # replace the descriptions and units
    df_final.loc[['Description', 'Units'], :] = df_temp

    # the WYTs for each year
    wy_flags_all = pd.read_excel(s_wy_flags_path, index_col=0)

    # add in long term average and dry and critical average
    for scenario in df_final.index.get_level_values(0).unique():
        if scenario in ['Description', 'Units']:
            continue

        # Long term average
        df_final.loc[(scenario, 'Long Term'), :] = df_final.loc[scenario, :].mean()
        li_dry_crit_years = wy_flags_all[wy_flags_all['40-30-30'].isin([4, 5])].index
        if 2021 in li_dry_crit_years:
            li_dry_crit_years = li_dry_crit_years.drop(2021)

            # dry and crit years
        df_final.loc[(scenario, 'Dry and Critical'), :] = df_final.loc[scenario, :].loc[li_dry_crit_years, :].mean()

    # read in exceedance plot formulas
    df_exceedance_formulas = pd.read_excel(s_formulas, sheet_name='exceedance', index_col=0)

    df_exceedances = pd.DataFrame(index=pd.MultiIndex.from_product([df_inputs['Scenario'].unique(), df_inputs['Year'].unique()]), columns=df_exceedance_formulas.index)

    # calculate each exceedance field
    for row_index, row in df_exceedance_formulas.iterrows():

        # fields to add up
        ls_fields = row[~row.isna()].values
        ls_add_fields = [field for field in ls_fields if field[0] != '-']
        ls_subtract_fields = [field[1:] for field in ls_fields if field[0] == '-']

        # If they are all water year, inclue 2021
        if np.all(df_formulas[df_formulas['Title'].isin(ls_fields)]['Annual_ Month_ Range'] == 'OctSep'):
            # add them up and insert into final data frame
            df_exceedances[row_index] = df_output[ls_add_fields].sum(axis=1) - df_output[ls_subtract_fields].sum(axis=1)
        else:
            # add them up and insert into final data frame
            df_exceedances[row_index] = df_output.drop(index=[2021], level=1)[ls_add_fields].sum(axis=1) - df_output.drop(index=[2021], level=1)[ls_subtract_fields].sum(axis=1)

    return df_final, df_exceedances


def parse_dssReader_output(dss_path, runs, field, report_type, convert_to_elevation= False, convert_to_cl=False,  orig_unit = 'TAF', storage_elevation_fn = ''):
    """
    Reads DSS output from reader for desired runs and field

    Parameters
    ----------
    dss_path: string
        Path and file name for xlsx file containing DSSReader Output
    runs: list of strings
        Names of the runs to be processed
    report_type: string
        Type of report being generated. Used to check whether or not it's a temperature report
    field: string
        Current field being processed
    convert_to_elevation: bool
        True if you are converting storage to elevation. Need to also set the orig_unit field to the original storage
        unit
    orig_unit: str
        Original storage unit (Currently only have "TAF" implemented). Used for storage to elevation conversion.

    storage_elevation_fn: str
        Optional. Default is "". Filename of Excel file containing storage-elev relationships for CalSim. (Storage-elev
        tables taken from lookup/res_info.table in CalSim wresl code.

    :returns
    t_dfs: list of pandas dataframes
        List of dataframe containing data for this location. Each dataframe corresponds to a run. Dataframe's has columns
        for WY, and each month.

        For temperature, daily values will be averaged to monthly.

    """
    #Read DSS Output from specified path for specified field
    dss_output = pd.read_excel(dss_path)

    dss_output = dss_output[["Month", "Scenario", "WY", field]]

    if report_type in ["temperature"]:
        #If temperature or DSM2 data is being read, convert daily data to monthly by averaging
        #scenario = dss_output.loc[0, "Scenario"]
        #dss_output.drop(columns = ["Scenario"], inplace = True)

        monthly_data = dss_output.groupby(["Scenario", "WY", "Month"]).mean()
        monthly_data.reset_index(inplace=True)
        dss_output = monthly_data

        #Drop rows with flag value for missing data
        rows_to_drop = (dss_output[dss_output.columns[3:]] < -100).any(axis=1)
        dss_output = dss_output[~rows_to_drop]

        #dss_output["Scenario"] = scenario
    #If we want elevation, need to convert from storage
    if convert_to_elevation:
        #Convert storage to elevation
        df_elevations = storage_to_elevation(dss_output, field, storage_elevation_fn, orig_unit = orig_unit)
        #Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_elevations
    if convert_to_cl:
        #Convert EC (microsiemens/cm) to mg/L Cl using the regression relationship defined as eqn 2 in
        #https://www.waterboards.ca.gov/waterrights/water_issues/programs/bay_delta/california_waterfix/exhibits/docs/ccc_cccwa/CCC-SC_25.pdf
        df_cl = ec_to_cl(dss_output, field, orig_unit = orig_unit)
        #Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_cl

    # Create df for each alternative/run and reformat
    run_dfs = []
    for run in runs:
        if run == "NAA":
            run_df = dss_output.loc[dss_output["Scenario"] == "Baseline"]
        else:
            run_df = dss_output.loc[dss_output["Scenario"] == run]

        run_df["month_name"] = " "

        #Add abbrievated month name to df for tables and plotting later
        for index, row in run_df.iterrows():
            run_df.loc[index, "month_name"] = calendar.month_abbr[int(row["Month"])]
        #Drop unneeded columns
        run_df.drop(columns=["Month", "Scenario"], inplace=True)
        run_dfs.append(run_df)

    #Transpose dfs to be in correct format for tables
    t_dfs = []
    for run_df in run_dfs:
        transposed_df = pd.DataFrame(
            columns=["WY", "Oct", "Nov", "Dec", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep"])
        #One row for each WY consisting of a column for each monthly EC value
        for year in np.unique(run_df["WY"]):
            year_t = run_df.loc[run_df["WY"] == year]
            year_t.set_index("month_name", inplace=True)
            year_t = year_t.transpose()
            year_t.insert(0, "WY", year)
            year_t.reset_index(drop=True, inplace=True)
            #Add each year as new row to df
            transposed_df = pd.concat([transposed_df, year_t.iloc[1:2]], axis=0, ignore_index=True)

        t_dfs.append(transposed_df)

    return t_dfs

def parse_dssReader_annualavg(dss_path, runs, field, report_type, convert_to_elevation= False, convert_to_cl=False,  orig_unit = 'TAF', storage_elevation_fn = ''):
    """
        Reads DSS output from reader for desired runs and field

        Parameters
        ----------
        dss_path: string
            Path and file name for xlsx file containing DSSReader Output
        runs: list of strings
            Names of the runs to be processed
        report_type: string
            Type of report being generated. Used to check whether or not it's a temperature report
        yr_start_month: int
            Integer calendar month (1= Jan) for start of each year used for the annual avg. 1 indicates a January start, so the
            average is computed using Jan1 through Dec 31 values.
        field: string
            Current field being processed
        convert_to_elevation: bool
            True if you are converting storage to elevation. Need to also set the orig_unit field to the original storage
            unit
        orig_unit: str
            Original storage unit (Currently only have "TAF" implemented). Used for storage to elevation conversion.

        """
    # Read DSS Output from specified path for specified field
    dss_output = pd.read_excel(dss_path)

    dss_output = dss_output[['Date', "Scenario", "WY", field]]
    dss_output['Date'] = pd.to_datetime(dss_output.Date, format = '%Y-%m-%d')

    if report_type in ["temperature"]:
        # Drop rows with flag value for missing data
        rows_to_drop = (dss_output[dss_output.columns[3:]] < -100).any(axis=1)
        dss_output = dss_output[~rows_to_drop]

        #Yearly averages (for each water year). This assumes daily data.
        i_first_complete_wy =  dss_output.Date.min().year +2 if dss_output.Date.min()> datetime(dss_output.Date.min().year, 10, 1,0) else dss_output.Date.min().year + 1
        i_last_complete_wy = dss_output.Date.max().year if dss_output.Date.max()>=datetime(dss_output.Date.max().year, 10,31,0) else dss_output.Date.max().year -1
        yrly_data = dss_output.loc[dss_output.WY.isin(range(i_first_complete_wy, i_last_complete_wy+1))].groupby(["Scenario", "WY"])[[field]].mean()
        yrly_data.reset_index(inplace=True)
        dss_output = yrly_data



        # dss_output["Scenario"] = scenario
    # If we want elevation, need to convert from storage
    if convert_to_elevation:
        # Convert storage to elevation
        df_elevations = storage_to_elevation(dss_output, field, storage_elevation_fn, orig_unit=orig_unit)
        # Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_elevations
        raise ValueError("Annual Averaging for elevations has not yet been implemented.")
    if convert_to_cl:
        # Convert EC (microsiemens/cm) to mg/L Cl using the regression relationship defined as eqn 2 in
        # https://www.waterboards.ca.gov/waterrights/water_issues/programs/bay_delta/california_waterfix/exhibits/docs/ccc_cccwa/CCC-SC_25.pdf
        df_cl = ec_to_cl(dss_output, field, orig_unit=orig_unit)
        # Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_cl
        raise ValueError("Annual Averaging for cloride has not yet been implemented.")

    # Create df for each alternative/run and reformat
    df_all_runs = pd.DataFrame(index = range(i_first_complete_wy, i_last_complete_wy + 1))
    for run in runs:
        if run == "NAA":
            run_df = dss_output.loc[dss_output["Scenario"] == "Baseline"]
        else:
            run_df = dss_output.loc[dss_output["Scenario"] == run]
        # Add this run's data to the dataframe of all run data.
        df_all_runs [run_df.Scenario.unique()[0]] = run_df.copy(deep = True).set_index("WY")[field]

    return df_all_runs

def parse_dssReader_calendaryr(dss_path, runs, field, report_type,  convert_to_elevation= False, convert_to_cl=False,  orig_unit = 'TAF', storage_elevation_fn = '', shastabin_data = '', use_calendar_yr=False):
    """
    Reads DSS output from reader for desired runs and field. Returns a dataframe with calendar year as the index and months (Jan - Dec) + shastabin flag, as columns.

    Parameters
    ----------
    dss_path: string
        Path and file name for xlsx file containing DSSReader Output
    runs: list of strings
        Names of the runs to be processed
    field: string
        Current field being processed
    report_type: string
        Type of report being generated. Used to check whether or not it's a temperature report
    convert_to_elevation: bool
        True if you are converting storage to elevation. Need to also set the orig_unit field to the original storage
        unit
    convert_to_cl: bool
        If the values need to be converted to chloride
    orig_unit: str
        Original storage unit (Currently only have "TAF" implemented). Used for storage to elevation conversion.
    storage_elevation_fn: str
        file name of storage-elevation table data. See note in function description.
    shastabin_data: str
        Path to shastabin data
    use_calendar_yr: bool
        Group by the calendar year (True) or water year (False)

    """
    #Read DSS Output from specified path for specified field
    dss_output = pd.read_excel(dss_path)
    dss_output = dss_output[['Date',"Month", "Scenario", field]]

    #Create a column for the Calendar Year (will be used to find the corresponding Shasta Bin type)
    if use_calendar_yr:
        dss_output['Year'] = dss_output.Date.dt.year
    else:
        dss_output['Year'] = np.where(dss_output.Date.dt.month < 10, dss_output.Date.dt.year, dss_output.Date.dt.year + 1)

    #Read in shastabin_ data
    if shastabin_data!= "":
        df_shastabin = pd.read_excel(shastabin_data,index_col =0)
    else:
        df_shastabin = pd.DataFrame(columns = ['SHASTABIN_', 'Scenario']) #Empty dataframe (just a placeholder. Doesn't get filled.)
        df_shastabin.index.name = 'calendar_yr'

    if report_type in ["temperature"]:
        #If temperature or DSM2 data is being read, convert daily data to monthly by averaging
        #scenario = dss_output.loc[0, "Scenario"]
        #dss_output.drop(columns = ["Scenario"], inplace = True)

        monthly_data = dss_output.groupby(["Scenario", "Year", "Month"]).mean()
        monthly_data.reset_index(inplace=True)
        dss_output = monthly_data

        #Drop rows with flag value for missing data
        rows_to_drop = (dss_output[[field]] < -100).any(axis=1)
        dss_output = dss_output[~rows_to_drop]
    else:
        raise ValueError("Parsing calendar year is not implemented yet for this report type.")

        #dss_output["Scenario"] = scenario
    #If we want elevation, need to convert from storage
    if convert_to_elevation:
        #Convert storage to elevation
        df_elevations = storage_to_elevation(dss_output, field, storage_elevation_fn, orig_unit = orig_unit)
        #Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_elevations
    if convert_to_cl:
        #Convert EC (microsiemens/cm) to mg/L Cl using the regression relationship defined as eqn 2 in
        #https://www.waterboards.ca.gov/waterrights/water_issues/programs/bay_delta/california_waterfix/exhibits/docs/ccc_cccwa/CCC-SC_25.pdf
        df_cl = ec_to_cl(dss_output, field, orig_unit = orig_unit)
        #Replace the dss_output dataframe and continue formatting the tables.
        dss_output = df_cl

    # Create df for each alternative/run and reformat
    run_dfs = []
    for run in runs:
        if run == "NAA":
            run_df = dss_output.loc[dss_output["Scenario"] == "Baseline"]
        else:
            run_df = dss_output.loc[dss_output["Scenario"] == run]

        #If the run has the shasta action in it (and has variable of "SHASTABIN_" in the CalSim DV file), then add a column for the shastabin.
        if run in df_shastabin.Scenario.values:
            run_df['Shastabin'] = run_df.apply(lambda l: df_shastabin.loc[(df_shastabin.index == l.Year)&(df_shastabin.Scenario == run)].SHASTABIN_.item(), axis = 1)

        run_df["month_name"] = " "

        #Add abbrievated month name to df for tables and plotting later
        for index, row in run_df.iterrows():
            run_df.loc[index, "month_name"] = calendar.month_abbr[int(row["Month"])]
        #Drop unneeded columns
        run_df.drop(columns=["Month", 'Date'], inplace=True)
        run_dfs.append(run_df)

    #Transpose dfs to be in correct format for tables
    t_dfs = []
    for run_df in run_dfs:
        scenario = run_df.Scenario.unique()[0]
        run_df.drop(columns = ["Scenario"], inplace = True)
        transposed_df = pd.DataFrame(
            columns=["Year", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"])
        #One row for each WY consisting of a column for each monthly EC value
        for year in np.unique(run_df["Year"]):
            year_t = run_df.loc[run_df["Year"] == year]
            year_t.set_index("month_name", inplace=True)
            year_t = year_t.transpose()
            year_t.insert(0, "Year", year)
            if scenario in df_shastabin.Scenario.values:
                year_t.insert(1, "SHASTABIN_", df_shastabin.loc[(df_shastabin.index== year)&(df_shastabin.Scenario == scenario)].SHASTABIN_.item())
            else:
                year_t.insert(1, "SHASTABIN_", np.nan)
            year_t.reset_index(drop=True, inplace=True)
            #Add each year as new row to df
            transposed_df = pd.concat([transposed_df, year_t.iloc[1:2]], axis=0, ignore_index=True)

        t_dfs.append(transposed_df)

    return t_dfs

def create_exceedance_tables(t_dfs, wy_flags_path, use_wytype, report_type, use_calendar_yr = False):
    """
    Creates exceedance tables formatted for appendix report from transposed DSSReader Output

    Parameters
    ----------
    t_dfs: list of dataframes
        Dataframe outputs from DSSReader that have been transposed to be formatted for table
    wy_flags_path: str
        excel file with the water year types
    use_wytype: str
        water year type to use.
        "40-30-30" to use WYT_SAC_
        "60-20-20" to use WYT_SJR_
        "TRIN" to use WYT_TRIN_
    report_type: str
        type of report (Calsim, temperature, etc)
    use_calendar_yr: bool
        Optional. Default is false. True indicates months should be sorted into calendar year instead of water year
        when calculating the water year type averages. False indicated that months will be sorted into water years when
        determining water year type.


    Returns
    ----
    exc_tables: list of pandas DataFrames containing 1,10,20,..., 90,99% exceedance probability data and the WYType summary statistics
    exc_probs: pandas dataseries of exceedance probabilities that are represented in the exc_tables
    fig_tables: list of pandas DataFrames of full list of exceedance probabilities and corresponding values. Used for plotting.
    il_num_years: list of pandas DataFrames that record how many years of data are available for each month

    """
    exc_tables = []
    fig_tables = []
    il_num_years = []
    wy_list = t_dfs[0].WY.tolist()
    for table in t_dfs:
        table = table.drop(columns = ["WY"]).copy()
        table = table.apply(lambda x: x.sort_values(ascending=False).values)
        #Remove first and last rows
        #table.iloc[::-1, ::1]
        #Rank ECs from 1 to 100 - No longer using this, since this produces inaccurate exceedance probabilities if simulation periods don't start Oct and end in Sept.
        #table.insert(0, "Rank", range(1,table.shape[0] + 1))
        ##Calculate exceedance probability and add column to table
        #table.insert(1, "Exc Prob", (table["Rank"]) / (table.shape[0] + 1) * 100) # m/(N+1)

        # Create dataframe for 10%, 20%, 30%, etc. exceedance values by linearly interpolating between the table values for each month
        table_interp = pd.DataFrame(index = range(10,100, 10))
        table_interp.index.name = 'Exc Prob'
        #Create dataframe for 1,2,3..., 99% exceedance values by linearly interpolating between the table values for each month. Used for plotting figures.
        table_all = pd.DataFrame(index = range(1,101,1))
        table_all.index.name = 'Exc Prob'

        for m_name in table.columns[-12:]:
            #Subset the table data column corresponding to this month.
            df_month = table[[m_name]]
            #Remove any null entries
            df_month.dropna(inplace = True)
            #Calculate the rank for the remaining (non-null) entries
            df_month['Rank'] = range(1,len(df_month)+1)
            df_month['Exc Prob'] = df_month["Rank"]/ (df_month.shape[0] + 1) * 100 # m/(N+1)

            #Interpolate to get the 10,20,...,90% exc prob values and place in dataframe.
            exceedance_values = np.interp(table_interp.index.values, df_month['Exc Prob'].values, df_month[m_name].values)
            table_interp[m_name]= exceedance_values

            # Interpolate to get the 1,2,3,4..., 99% exc prob values and place in dataframe.
            exceedance_values_all = np.interp(table_all.index.values, df_month['Exc Prob'].values,
                                              df_month[m_name].values)
            df_exceedance_values_all = pd.DataFrame(exceedance_values_all, index = table_all.index.values, columns = [m_name])
            #Only include values for exp probabilities between min(m/(N+1)) and max(m/(N+1))
            df_exceedance_values_all = df_exceedance_values_all.loc[(df_exceedance_values_all.index>= df_month['Exc Prob'].min()) &(df_exceedance_values_all.index<=df_month['Exc Prob'].max())]
            table_all[m_name]=  df_exceedance_values_all
        #Reset index for the table_all
        table_all.reset_index(inplace = True, drop = False)

        #Add the lowest and highest exceedance probability rows
        table_interp.loc[table_all.iloc[0]['Exc Prob']] = table_all.iloc[0][table_all.columns[-12:]].values
        table_interp.loc[table_all.iloc[-1]['Exc Prob']] = table_all.iloc[-1][table_all.columns[-12:]].values

        #Sort by exceedance probability.
        table_interp.sort_index(inplace = True)
        table_interp.reset_index(drop = False, inplace = True)

        #Store the exceedance table
        exc_tables.append(table_interp)
        #Store the exc probability table for plotting the figures
        fig_tables.append(table_all.copy(deep = True))

        #Calculate the sample size used to calculate statistics in each month
        df_num_years = table.count(axis = 0)
        il_num_years.append(df_num_years)

    #Calculate full simulation period average for each run and format to be added to exceedance table as one row
    stats_dfs = []
    for run in t_dfs:
        period_ave = run.drop(columns = ['WY']).mean(axis=0)
        stats_df = pd.DataFrame(period_ave)
        stats_df = stats_df.transpose()

        stats_df["Exc Prob"] = ["Full Simulation Period Average"]

        stats_dfs.append(stats_df)

    #Read in water year typing flags
    wy_flags_all = pd.read_excel(wy_flags_path, index_col = 0)
    #Subset to only the wytype that is specified by use_wytype
    wy_flags = wy_flags_all[[use_wytype]]

    if use_wytype == 'TRIN': #Names for Trinity WYType
        year_types = ["Extremely Wet", "Wet", "Normal", "Dry", "Critically Dry"]
    else: #Names for the Sacramento and SJR WYType
        year_types = ["Wet", "Above Normal", "Below Normal", "Dry", "Critical"]
    # make a copy of exc probabilities to use with figures after deleting from tables df
    exc_probs = exc_tables[0]["Exc Prob"]

    #Create empty dataframe to store percentages for each of the wytypes
    wytype_percents = pd.DataFrame(index = range(1, 6), columns = ['percentage'])

    # calculate wet, above normal, dry, etc water years (sum for year type/ count of year type)
    for table_index in range(len(t_dfs)):

        t_dfs[table_index].set_index('WY', inplace = True)
        t_dfs[table_index]["flag"] = wy_flags[use_wytype]
        if use_calendar_yr:
            df_monthly_ts = pd.melt(t_dfs[table_index].reset_index(drop=False), id_vars='WY',
                    value_vars=t_dfs[table_index].columns[:-1])
            df_monthly_ts['i_month']= df_monthly_ts.apply(lambda l: datetime.strptime(l.variable+ "-01-1900","%b-%d-%Y" ).month,axis = 1)
            df_monthly_ts['dates'] = df_monthly_ts.apply(lambda l: datetime(l.WY -1 , l.i_month, 1) if l.i_month>=10 else datetime(l.WY, l.i_month, 1), axis = 1)
            df_monthly_ts['calendar_yr'] = df_monthly_ts.apply(lambda l: l.WY -1  if l.i_month>=10 else l.WY, axis = 1)

            #Calendar year dataframe (rows are calendar years, columns are months, values are monthly values
            t_dfs_calendar_yr = df_monthly_ts.pivot(columns = 'variable', index = 'calendar_yr', values = 'value')
            #Reorder the columns so months are in order.
            t_dfs_calendar_yr = t_dfs_calendar_yr [['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']]

            #Grab the wy that is associated with each calendar year and add as a column. This means that calendar yr 1980 will be assigned the flag for wytype associated with wy1980.
            t_dfs_calendar_yr['flag'] = wy_flags[use_wytype]

        exc_probs_i = exc_tables[table_index]["Exc Prob"]
        month_vals = {}
        # Also add full sim period average as a row in exceedance table
        exc_tables[table_index] = pd.concat([exc_tables[table_index], stats_dfs[table_index].iloc[0:1]], ignore_index=True)

        #Iterate through each type of year (wet, above normal, etc) to compute sums
        for year_type in range(len(year_types)):

            # Calculate the percentage of total water years that have this wytype
            if use_calendar_yr:
                d_percent_wytype = round(len(t_dfs_calendar_yr.loc[t_dfs_calendar_yr['flag'] == year_type + 1]) / len(t_dfs_calendar_yr) * 100, 1)
            else:
                d_percent_wytype = round(len(t_dfs[table_index].loc[t_dfs[table_index]['flag'] == year_type + 1]) / len(t_dfs[table_index]) * 100, 1)
            wytype_percents.at[year_type+1, 'percentage'] = d_percent_wytype
            for month in t_dfs[table_index].columns[:-1]:
                #Flags are 1 - 5 to specify which type of year
                #Calculate mean of months classified as current year type based on flag

                ## Using water years
                if not use_calendar_yr:
                    month_vals[month] = [t_dfs[table_index].loc[t_dfs[table_index]['flag'] == (year_type + 1), month].mean()]
                else:
                    # Using calendar years
                    month_vals[month] =[t_dfs_calendar_yr.loc[t_dfs_calendar_yr['flag'] == (year_type + 1), month].mean()]




            month_vals["Exc Prob"] = year_types[year_type]

            exc_tables[table_index] = pd.concat([exc_tables[table_index], pd.DataFrame.from_dict(month_vals)], ignore_index=True)


        #Create list of desired row labels
        row_labels = [f"{round(value)}% Exceedance" for value in exc_probs_i.values]
        row_labels.append('Full Simulation Period Average')
        wy_type_labels = [f"{year_types[i]} Years ({wytype_percents.loc[i+1].item():.0f}%)" if wytype_percents.loc[i+1].item() == int(wytype_percents.loc[i+1].item())  else
                          f"{year_types[i]} Years ({wytype_percents.loc[i + 1].item():.1f}%)" for i in range(len(year_types))]
        row_labels.extend(wy_type_labels)

        #Remove extra columns
        exc_tables[table_index].drop(columns=["Exc Prob"], inplace=True)

        #Round table values
        if report_type == 'temperature':
            exc_tables[table_index] = exc_tables[table_index].astype(float)#.round(1)
        else:
            exc_tables[table_index] = exc_tables[table_index].astype(float)#.round(0)

        # Add row labels for report tables in first column
        exc_tables[table_index].insert(0, "Statistic", row_labels)

        # Move new header names to first row
        exc_tables[table_index].index = exc_tables[table_index].index + 1  # shifting index
        exc_tables[table_index] = exc_tables[table_index].sort_index()

    return exc_tables, exc_probs, fig_tables, il_num_years


def make_rows_bold(*rows):
    """
    Makes text in specified table rows bold.

    Parameters
    ----------
    rows: row attributes from docx table object
        1 or more rows that will be converted to bold text (ex table.rows[0])

    """

    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell's border.
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )

     Parameters
    ----------
    cell: cell attribute from docx table object
        1 cell that will be converted to bold text (ex table.rows[0].cells[0])

    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_color(cell, color):
    """
    Sets the cell color of a table
    Parameters
    ----------
    cell:  cell attribute from docx table object
        1 cell that will be have the color set
    color: str
        Color code to set the cell color

    Returns
    -------
    None
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def change_table_font_size(document, font_size):
    """
    Changes the font size of all text in all tables within a document.

    Parameters
    ----------
    document: docx document object
        Document that will have font size adjusted
    font_size: int
        New font size

    """

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(font_size)


def add_commas_to_table(doc):
    """
    Adds commas to numbers in all tables of a docx document.

    Parameters
    ----------
    doc: docx document object
        Document that will have commas added to numeric values in all tables

    """

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        try:
                            # Check if the text is a number
                            number = float(run.text)
                            # Format the number with commas
                            formatted_number = f"{number:,}"
                            formatted_number = formatted_number.rsplit(".", 1)[0]
                            run.text = formatted_number
                        except ValueError:
                            # If the text is not a number, do nothing
                            pass

def format_decimals(doc):
    """
    Adds commas to numbers in all tables of a docx document.

    Parameters
    ----------
    doc: docx document object
        Document that will have commas added to numeric values in all tables

    """

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        try:
                            # Check if the text is a number
                            number = float(run.text)
                            # Format the number with commas
                            #formatted_number = f"{number:,}"
                            formatted_number = "{:.1f}".format(number)
                            #formatted_number = formatted_number.rsplit(".", 1)[0]
                            run.text = formatted_number
                        except ValueError:
                            # If the text is not a number, do nothing
                            pass
def change_orientation(doc, new_orientation):
    """
    Changes section orientation from portrait to landscape or vice versa

    Parameters
    ----------
    doc: docx document object
        Document that will have commas added to numeric values in all tables
    new_orientation: string
        Either "landscape" or "portrait" to indicate the desire page orientation for the new section

    Returns
    --------
    new_section: docx section

    """

    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    if new_orientation == "landscape":
        new_section.orientation = WD_ORIENT.LANDSCAPE
    else:
        new_section.orientation = WD_ORIENT.PORTRAIT

    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section

def format_table_basic(doc_table, table_df, doc):
    """
    Creates tables formatted for appendix report from exceedance tables

    Parameters
    ----------
    doc_table: docx table object
        Exceedance table to be formatted for report
    table_df: dataframe
        Dataframe containing data to go into report table
    doc: docx object
        Docx object containing table to be formatted
    """
    # Change font size to fit on page better
   # change_table_font_size(doc, 8)

    # add the header rows.
    for column_index in range(table_df.shape[1]):
        doc_table.cell(0, column_index).text = table_df.columns[column_index]

    # add the rest of the data frame
    for row_index in range(table_df.shape[0]):
        for column_index in range(table_df.shape[1]):
            # Round index to 1 decimal. Round all other values to nearest whole number

            if column_index == 0:
                doc_table.cell(row_index + 1, column_index).text = str(round(table_df.values[row_index, column_index],1))
            else:
                doc_table.cell(row_index + 1, column_index).text = str(round(table_df.values[row_index, column_index]))

    # Set table top and bottom borders
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')
    borders.append(bottom_border)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')
    borders.append(top_border)

    doc_table._tbl.tblPr.append(borders)

    # Make headers bold
    make_rows_bold(doc_table.rows[0])

    # Make first column bold
    bolding_columns = [0]
    for row in list(range(table_df.shape[0] + 1)):
        for column in bolding_columns:
            doc_table.rows[row].cells[column].paragraphs[0].runs[0].font.bold = True

    # Add borders under header
    for cell in doc_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})

        # Align values in center of cells
        for row in doc_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Decrease row spacing for table
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.45)  # 2 cm

        # Change font size to fit on page better
        change_table_font_size(doc, 8)

def format_table(doc_table, table_df, doc, report_type):
    """
    Creates tables formatted for appendix report from exceedance tables

    Parameters
    ----------
    doc_table: docx table object
        Exceedance table to be formatted for report
    table_df: dataframe
        Dataframe containing data to go into report table
    doc: docx object
        Docx object containing table to be formatted
    report_type: str
        type of report. (used to determine what decimal place to round to).

    Returns
    -------
    None

    """
    # Change font size to fit on page better
   # change_table_font_size(doc, 8)

    # add the header rows.
    for column_index in range(table_df.shape[1]):
        doc_table.cell(0, column_index).text = table_df.columns[column_index]

    # add the rest of the data frame
    for row_index in range(table_df.shape[0]):
        for column_index in range(table_df.shape[1]):
            if column_index == 0: #Add the exceedance percentage row labels to the table
                doc_table.cell(row_index + 1, column_index).text = str(table_df.values[row_index, column_index])
            else:
                #For all other columns, add the CalSim, temperature, or salinity values in the table. Round values.
                if report_type == 'temperature':
                    #For temperature tables, round values to nearest 10th place
                    doc_table.cell(row_index + 1, column_index).text = str(round(table_df.values[row_index, column_index],1))
                else:
                    #For CalSim or DSM2 tables, round values to the nearest integer
                    doc_table.cell(row_index + 1, column_index).text = str(round(table_df.values[row_index, column_index]))
    # Set table top and bottom borders
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')
    borders.append(bottom_border)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')
    borders.append(top_border)

    doc_table._tbl.tblPr.append(borders)

    # Make headers bold
    make_rows_bold(doc_table.rows[0])

    # Make first column bold
    bolding_columns = [0]
    for row in list(range(table_df.shape[0] + 1)):
        for column in bolding_columns:
            doc_table.rows[row].cells[column].paragraphs[0].runs[0].font.bold = True

    # Add superscript to Full Simulation Period Average cell
    script_cell = doc_table.cell(10, 0).paragraphs[0]
    run = script_cell.add_run("a")
    run.font.superscript = True

    # Add borders to middle row and under header
    for cell in doc_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})

    for cell in doc_table.rows[10].cells:
        set_cell_border(cell, bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})

    for cell in doc_table.rows[10].cells:
        set_cell_border(cell, top={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})

    # Widen margins of table
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Widen cell size in first column
    for cell in doc_table.columns[0].cells:
        cell.width = Inches(3.4)

    if report_type == "temperature":
        #format numbers to one decimal pt
        format_decimals(doc)
    else:
        # Add commas to values in table
        add_commas_to_table(doc)
        #Commas won't be needed for temperature values

    # Align values in center of cells
    for row in doc_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Decrease row spacing for table
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.45)  # 2 cm

    # Change font size to fit on page better
    change_table_font_size(doc, 8)


def format_table_supply(doc_table, df_table, doc, comparison, il_page_breaks):
    """
    Creates table for water supply data

    Parameters
    ----------
    doc_table: docx table object
        Table to be formatted for report
    df_table: dataframe
        Dataframe containing data to go into report table
    doc: docx object
        Docx object containing table to be formatted
    comparison: list
        List of the comparison names
    il_page_breaks: list
        Rows that need a page break header

    Returns
    -------
    none
    """

    doc_table.autofit = False
    # set consistent borders over whole table
    for row in doc_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                            bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                            start={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                            end={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Decrease row spacing for table
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Inches(0.25)

    # Create the Header Row
    doc_table.cell(0, 0).merge(doc_table.cell(0, 3))
    doc_table.cell(0, 0).text = 'Water Supply Reliability'
    doc_table.cell(0, 4).text = comparison[1]
    doc_table.cell(0, 5).text = comparison[0]
    doc_table.cell(0, 6).text = comparison[1] + ' minus ' + comparison[0]
    doc_table.rows[0].height = Inches(0.7)
    doc_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Make headers bold
    make_rows_bold(doc_table.rows[0])

    curr_row = 1

    # Loops through each section and subsetion and add into table
    for section_name in df_table.columns.get_level_values(0).unique():

        # if we hit a page break, recreate the header
        if curr_row in il_page_breaks:

            # add row to bottom and format it
            row = doc_table.add_row()
            for cell in row.cells:
                set_cell_border(cell, top={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                start={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                end={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Decrease row spacing for table
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(0.25)

            # Create the Header Row
            doc_table.cell(curr_row, 0).merge(doc_table.cell(curr_row, 3))
            doc_table.cell(curr_row, 0).text = 'Water Supply Reliability'
            doc_table.cell(curr_row, 4).text = comparison[1]
            doc_table.cell(curr_row, 5).text = comparison[0]
            doc_table.cell(curr_row, 6).text = comparison[1] + ' minus ' + comparison[0]
            doc_table.rows[curr_row].height = Inches(0.7)
            doc_table.rows[curr_row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            # Make headers bold
            make_rows_bold(doc_table.rows[curr_row])
            curr_row += 1

        # Create the section
        doc_table.cell(curr_row, 0).merge(doc_table.cell(curr_row, 6))
        doc_table.cell(curr_row, 0).text = section_name.upper()

        # Make headers bold
        make_rows_bold(doc_table.rows[curr_row])

        # Make the background grey
        set_cell_color(doc_table.cell(curr_row, 0), "#E8E8E8")

        curr_row += 1

        # go through the subsections
        for sub_section in df_table[section_name].columns:

            # if we hit a page break, recreate the header
            if curr_row in il_page_breaks:
                row = doc_table.add_row()
                for cell in row.cells:
                    set_cell_border(cell, top={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                    bottom={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                    start={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"},
                                    end={"sz": 7, "color": "#000a00", "space": 0.5, "val": "single"})
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Decrease row spacing for table
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Inches(0.25)

                # Create the Header Row
                doc_table.cell(curr_row, 0).merge(doc_table.cell(curr_row, 3))
                doc_table.cell(curr_row, 0).text = 'Water Supply Reliability'
                doc_table.cell(curr_row, 4).text = comparison[1]
                doc_table.cell(curr_row, 5).text = comparison[0]
                doc_table.cell(curr_row, 6).text = comparison[1] + ' minus ' + comparison[0]
                doc_table.rows[curr_row].height = Inches(0.7)
                doc_table.rows[curr_row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

                # Make headers bold
                make_rows_bold(doc_table.rows[curr_row])
                curr_row += 1

            # add the name of the subsection
            doc_table.cell(curr_row, 0).merge(doc_table.cell(curr_row + 1, 0))
            doc_table.cell(curr_row, 0).text = sub_section

            # add description
            doc_table.cell(curr_row, 1).merge(doc_table.cell(curr_row + 1, 1))
            doc_table.cell(curr_row, 1).text = df_table.loc['Description', (section_name, sub_section)]

            # Add units
            doc_table.cell(curr_row, 2).merge(doc_table.cell(curr_row+1, 2))
            doc_table.cell(curr_row, 2).text = df_table.loc['Units', (section_name, sub_section)]

            # Add long term numbers
            doc_table.cell(curr_row, 3).text = 'Long Term'
            doc_table.cell(curr_row, 4).text = str(round(df_table.loc[(comparison[1], 'Long Term'), (section_name, sub_section)]))
            doc_table.cell(curr_row, 5).text = str(round(df_table.loc[(comparison[0], 'Long Term'), (section_name, sub_section)]))
            doc_table.cell(curr_row, 6).text = str(round(df_table.loc[(comparison[1], 'Long Term'), (section_name, sub_section)]- df_table.loc[(comparison[0], 'Long Term'), (section_name, sub_section)]))

            curr_row += 1
            # add dry and crit numbers
            doc_table.cell(curr_row, 3).text = 'Dry and Critical'
            doc_table.cell(curr_row, 4).text = str(round(df_table.loc[(comparison[1], 'Dry and Critical'), (section_name, sub_section)]))
            doc_table.cell(curr_row, 5).text = str(round(df_table.loc[(comparison[0], 'Dry and Critical'), (section_name, sub_section)]))
            doc_table.cell(curr_row, 6).text = str(
                round(df_table.loc[(comparison[1], 'Dry and Critical'), (section_name, sub_section)] - df_table.loc[(comparison[0], 'Dry and Critical'), (section_name, sub_section)]))

            curr_row += 1

    # Formatting for table
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')
    borders.append(bottom_border)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')
    borders.append(top_border)

    doc_table._tbl.tblPr.append(borders)

    # Add commas to values in table
    add_commas_to_table(doc)

    # Adjust the width of the columns
    for cell in doc_table.columns[0].cells:
        cell.width = Inches(1.45)
    for cell in doc_table.columns[1].cells:
        cell.width = Inches(3.2)
    for cell in doc_table.columns[2].cells:
        cell.width = Inches(.75)
    for cell in doc_table.columns[3].cells:
        cell.width = Inches(1.2)
    for cell in doc_table.columns[4].cells:
        cell.width = Inches(0.8)
    for cell in doc_table.columns[5].cells:
        cell.width = Inches(0.8)
    for cell in doc_table.columns[6].cells:
        cell.width = Inches(0.8)

    # Change font size to fit on page better
    change_table_font_size(doc, 10)

def create_mixed_compliance_month_plots (location, dfs_calendaryr, fig_value, month, month_directory, alts, line_styles, line_colors, compliance_dict ):
    if not os.path.exists(month_directory):
        os.makedirs(month_directory)

    # # Read in SHASTABIN_ data
    # if shastabin_data != "":
    #     df_shastabin = pd.read_excel(shastabin_data)

    #Create figures
    df_month_alts = pd.DataFrame(columns = alts)
    fig, axs = plt.subplots(figsize=(10, 5), linewidth=3, edgecolor="black")
    for fig_index in range(len(dfs_calendaryr)):
        # Dataset for this alt
        df_alt_data = dfs_calendaryr[fig_index].copy(deep = True)
        df_alt_data.set_index('Year', inplace = True)

        #Subset to only the month of interest
        df_month = df_alt_data[[ 'SHASTABIN_', month]]

        #Now calculate exceedance values using this month's data
        df_month.sort_values(by = month, inplace = True, ascending = False)
        df_month.dropna(subset = [month], inplace = True)
        df_month['Rank'] = range(1, len(df_month) + 1)
        df_month['Exc Prob'] = df_month["Rank"] / (df_month.shape[0] + 1) * 100  # m/(N+1)

        #Create a column to indicate whether or not this location is used for compliance under this alt (This depends on Shastabin value)
        #Shastabin_ == 1 or 2 means compliance location is at Sac Rv at AIRPORT RD.
        #Shastabin_ == 3 or 4 means compliance location is  Sac Rv blw Clear Creek.
        #Shastabin_ == 5 or 6 means compliance location is at Sac Rv at HWY 44
        df_month['compliance'] = df_month.apply(lambda l: False if np.isnan(l.SHASTABIN_) else (True if compliance_dict[l.SHASTABIN_] == location else False), axis = 1)

        df_month_alts[alts[fig_index]] = df_month[[month, 'Exc Prob']].reset_index(drop = False).set_index("Exc Prob")['Year']

        #Percentage array from 0 to 100 (used for xtick labels)
        percentages = range(0, 101, 10)
        percentage_labels = [f"{int(i)}%" for i in percentages]

        #Plot the exceedance.
        axs.plot(df_month['Exc Prob'], df_month[month], color=line_colors[fig_index],
                 linestyle=line_styles[fig_index], label=alts[fig_index])

        #Add markers on the exceedance plot if this location is a compliance location for that year.
        if len(df_month.dropna(subset= ['SHASTABIN_']))>0:
            axs.plot(df_month.loc[df_month.compliance]['Exc Prob'], df_month.loc[df_month.compliance][month], color = line_colors[fig_index], linestyle = 'none', marker = 'o', markersize = 3,  label = alts[fig_index] + ' - Compliance Location Years')

        #Add annotations
        # for ind, row in df_month.loc[df_month.compliance].iterrows():
        #     if fig_index == 1:
        #         axs.annotate(text = str(ind), xy = (row['Exc Prob'], row[month]), ha = 'left', va = 'top', rotation = -45,
        #                      fontsize = 4, color =line_colors[fig_index] )
        #     elif fig_index == 2:
        #         axs.annotate(text=str(ind), xy=(row['Exc Prob'], row[month]), ha='right', va='bottom', rotation=-45,
        #                      fontsize=4, color=line_colors[fig_index])

        #Format axes
        axs.set_xticks(percentages)
        axs.set_xticklabels(percentage_labels)
        axs.set_ylabel(fig_value)
        axs.set_xlabel("Exceedance Probability")

        # Save this parameter to orient the legend correctly
        axbox = axs.get_position()

        # Add gridlines
        plt.grid(color='gray', linestyle='--', linewidth=0.8)

        # Add a legend
        plt.legend(loc='center', ncol=3, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08], fontsize = 10)

    # Add month number at beginning so that figures can be easily inserted in CY order to document later
    month_number = str(strptime(month, '%b').tm_mon)

    # flip x-axis
    axs.invert_xaxis()

    # Add leading zeros to month numbers
    if len(month_number) < 2:
        month_number = str(0) + month_number
    # Save figure to month directory
    plt.savefig(month_directory + "/" + month_number + "_" + month + "_monthly_exceedance" + ".png", dpi = 300)
    plt.close()
    return df_month_alts

def create_month_plot(dfs, fig_value, month, month_directory, alts, line_styles, line_colors, report_type=''):
    """
    Generates and saves individual month plots

    Parameters
    ----------
    dfs: list of dataframes
        List of dataframes with monthly values as one of the columns
    fig_value: str
        y-axis label
    month: string
        Current month to be evaluated
    month_directory: string
        Directory to save month plots in
    alts: list of strings
        Names of the runs being compared in report
    line_styles: list of strings
        Styles for lines on plots
    line_colors: list of strings
        Colors for lines on plots
    report_type: str
        Type of report, only really matters if its water supply

    Returns
    --------
    None

    """
    # Check for/create directory to store monthly exceedance plots
    if not os.path.exists(month_directory):
        os.makedirs(month_directory)

    # define size and borders
    if report_type == 'water supply':
        fig, axs = plt.subplots(figsize=(9, 5.5), linewidth=1, edgecolor="black")
    else:
        fig, axs = plt.subplots(figsize=(10, 5), linewidth=3, edgecolor="black")

    for fig_index in range(len(dfs)):
        # Dataset for this alt
        df_alt_data = dfs[fig_index].copy(deep=True)

        # Subset to only the month of interest
        df_month = df_alt_data[[month]]

        # Now calculate exceedance values using this month's data
        df_month.sort_values(by=month, inplace=True, ascending=False)
        df_month.dropna(subset=[month], inplace=True)
        df_month['Rank'] = range(1, len(df_month) + 1)
        df_month['Exc Prob'] = df_month["Rank"] / (df_month.shape[0] + 1) * 100  # m/(N+1)

        # plot exceedance probability vs monthly EC
        percentages = range(0, 101, 10)
        percentage_labels = [f"{int(i)}%" for i in percentages]

        axs.plot(df_month['Exc Prob'].values, df_month[month].values, color=line_colors[fig_index],
                 linestyle=line_styles[fig_index], label=alts[fig_index])
        axs.set_xticks(percentages)
        axs.set_xticklabels(percentage_labels)


        axs.set_ylabel(fig_value)
        axs.set_xlabel("Exceedance Probability")

        # Save this parameter to orient the legend correctly
        axbox = axs.get_position()

        # Add gridlines
        plt.grid(color='gray', linestyle='--', linewidth=0.8)

        # Add a legend
        plt.legend(loc='center', ncol=4, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08])

    if report_type != 'water supply':
        # Add month number at beginning so that figures can be easily inserted in CY order to document later
        month_number = str(strptime(month, '%b').tm_mon)
        # Add leading zeros to month numbers
        if len(month_number) < 2:
            month_number = str(0) + month_number

    # flip x-axis
    axs.invert_xaxis()

    if report_type == 'water supply':
        # Save figure to directory
        plt.savefig(month_directory + "/" + month + ".png")
    else:
        # Save figure to month directory
        plt.savefig(month_directory + "/" + month_number + "_" + month + "_monthly_exceedance" + ".png")

    plt.close()

def create_annual_exceedance_plot(df_annual, fig_value, yr_directory, alts, line_styles, line_colors):
    """
    Generates and saves individual month plots

    Parameters
    ----------
    df_annual: pandas dataframe
        Dataframe with WYs as the index and annual average values for each of the runs as values. Columns are run names.
    yr_directory: string
        Directory to save the annual exceedance plot to.
    alts: list of strings
        Names of the runs being compared in report
    line_styles: list of strings
        Styles for lines on plots
    line_colors: list of strings
        Colors for lines on plots
    """
    # Check for/create directory to store monthly exceedance plots
    if not os.path.exists(yr_directory):
        os.makedirs(yr_directory)

    fig, axs = plt.subplots(figsize=(10, 5), linewidth=3, edgecolor="black")
    for fig_index, altname in enumerate(df_annual.columns):
        # Dataset for this alt
        df_alt_data = df_annual[[altname]].copy(deep=True)

        # Now calculate exceedance values using this month's data
        df_alt_data.sort_values(by=altname, inplace=True, ascending=False)
        df_alt_data.dropna(subset=altname, inplace=True)
        df_alt_data['Rank'] = range(1, len(df_alt_data) + 1)
        df_alt_data['Exc Prob'] = df_alt_data["Rank"] / (df_alt_data.shape[0] + 1) * 100  # m/(N+1)

        # plot exceedance probability vs monthly EC
        percentages = range(0, 101, 10)
        percentage_labels = [f"{int(i)}%" for i in percentages]

        axs.plot(df_alt_data['Exc Prob'].values, df_alt_data[altname].values, color=line_colors[fig_index],
                 linestyle=line_styles[fig_index], label=alts[fig_index])

    axs.set_xticks(percentages)
    axs.set_xticklabels(percentage_labels)

    axs.set_ylabel(fig_value)
    axs.set_xlabel("Exceedance Probability")

    # Save this parameter to orient the legend correctly
    axbox = axs.get_position()

    # Add gridlines
    plt.grid(color='gray', linestyle='--', linewidth=0.8)

    # Add a legend
    plt.legend(loc='center', ncol=4, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08])

    # flip x-axis
    axs.invert_xaxis()

    # Save figure to month directory
    plt.savefig(os.path.join(yr_directory, "annualavg_exceedance.png"))
    plt.close()

def create_stat_plot(stat_fig_dfs, fig_value, stat, stat_directory, alts, line_styles, line_colors):
    """
    Generates and saves individual month plots

    Parameters
    ----------
    stat_fig_dfs: list of dataframes
        Dataframes with average values by year type
    fig_value: str
        plot ylabel
    stat: string
        Current type of year to be evaluated
    stat_directory: string
        Directory to save stat plots in
    alts: list of strings
        Names of the runs being compared in report
    line_styles: list of strings
        Styles for lines on plots
    line_colors: list of strings
        Colors for lines on plots
    Returns
    ----------
    None
    """
    if not os.path.exists(stat_directory):
        os.makedirs(stat_directory)

    fig, axs = plt.subplots(figsize=(10, 5), linewidth=3, edgecolor="black")
    for fig_index in range(len(stat_fig_dfs)):
        if stat == "Full Simulation Period":
            axs.plot(stat_fig_dfs[fig_index]["month"], stat_fig_dfs[fig_index]["Full Simulation Period Average"], color=line_colors[fig_index],
                     linestyle=line_styles[fig_index])
        else:
            axs.plot(stat_fig_dfs[fig_index]["month"], stat_fig_dfs[fig_index][stat], color=line_colors[fig_index],
                     linestyle=line_styles[fig_index])

        # Save this to position legend correctly
        axbox = axs.get_position()

        axs.set_ylabel(fig_value)

        # Add gridlines
        plt.grid(color='gray', linestyle='--', linewidth=0.8)
        # Add legend
        plt.legend(labels=alts, loc='center', ncol=4, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08])

    # Save stat fig to directory
    plt.savefig(stat_directory + "/" + stat[:5] + "_exceedance" + ".png")
    plt.close()

def order_elevation_storage_fields(fields):
    """
    Generates list of tuples where each tuple is (input field, storage or elevation), based on the fields provided.
    This is used to preprocess the fields for the storage and elevation appendix, since some fields have both storage
    and elevation, while others just have storage (Ex: S_SLUIS_CVP).

    Note that this function will check for whether the field is in the list or not. If it isn't it'll raise an error
    telling the user they need to update the master list.

    Parameters
    ----------
    fields: list of str
        List of the fields being included in this appendix. This function is only intended to take reservoir fields as
        inputs.
    """
    # List of all the location and elevation fields in the desired order
    master_list = [("S_TRNTY", 'Storage'),
                   ("S_TRNTY", 'Elevation'),
                   ("S_SHSTA", 'Storage'),
                   ("S_SHSTA", 'Elevation'),
                   ("S_OROVL", 'Storage'),
                   ("S_OROVL", 'Elevation'),
                   ("S_FOLSM", 'Storage'),
                   ("S_FOLSM", 'Elevation'),
                   ("S_SLUIS", 'Storage'),
                   ("S_SLUIS", 'Elevation'),
                   ("S_SLUIS_CVP", 'Storage'),
                   ("S_SLUIS_SWP", 'Storage'),
                   ("S_MELON", "Storage"),
                   ("S_MELON", "Elevation"),
                   ("S_MLRTN", "Storage"),
                   ("S_MLRTN", "Elevation"),
                    ] #List of all the location and elevation fields in the desired order

    #Subset the list based on the fields.
    subset_list = [ordered_field for ordered_field in master_list if ordered_field[0] in fields]

    #Check to make sure that there's no new field names that aren't in the master list
    fields_in_master = [ordered_field[0] for ordered_field in master_list]
    for field in fields:
        if field not in fields_in_master:
            #If any field is not included in the master_list, then throw an error telling the user to add it to the master list. Otherwise, it will be excluded from the generated appendix.
            raise ValueError(f"Need to update master list with new field {field}.")

    return subset_list

def create_water_supply_appendix(alts, appendix_prefix, dss_path, doc_name, new_doc, wy_flags_path, template, s_supply_formulas):
    """
    Creates the water supply appendix. Creates all the tables and plot and puts them into a doccument.

    Parameters
    ----------
    alts: list
        List of alternatives to include in the appendix
    appendix_prefix: str
        Prefix to use in the appendix. Typically, would be ' F.2.4'
    dss_path: str
        Path to the excel file with the DSS contents
    doc_name: str
        Name for temporary document
    new_doc: str
        Name for the final document
    wy_flags_path: str
        Path to the water year type flags file
    template: str
        Path to the template doc
    s_supply_formulas: str
        path to the water supply formulas excel file

    Returns
    -------
    None
    """

    # compare every run to the baseline run
    comparisons = [["NAA", alt] for alt in alts]
    # Remove first comparison that is NAA and NAA
    comparisons.pop(0)

    """
    For each field, there are:
        - 3 comparison tables per alternative. (Ex: If you have 6 alternatives, you will have 18 tables total for S_TRNTY)
        - 10 exceedance plots. 
    """

    # two tables per comparison
    num_tables = 3 * len(comparisons)
    # always 10 tables
    num_figures = 10
    # Alt Text strings, in order for tables
    alt_text_tables = ["Alt text table example" for t in range(0, num_tables)]

    # Alt text strings, order for figures
    alt_text_figures = ["This figure shows data also presented in data tables in this file." for f in range(0, num_figures)]

    # Create an instance of a word document
    # Open the word document template. This template has the heading style 2 formatted with numbering to allow the figures
    # to inherit the heading numbering.
    doc = docx.Document(template)
    doc.add_heading(f"Attachment{appendix_prefix}", level=1)  # Add Heading 1 (Attachment XXX)

    # Add caption style for Figure captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.name = 'Times New Roman'

    # Add caption style for Table captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    # calculate fields
    dfs, df_exceedances = calculate_supply_fields(dss_path, s_supply_formulas, wy_flags_path)

    for comparison_index, comparison in enumerate(comparisons):
        if comparison_index == 0:
            doc.add_page_break()
            change_orientation(doc, "landscape")

        # Add heading for first table
        tab_title_prefix = "Table " + appendix_prefix + "-"
        add_caption_water_supply(doc, "Table", tab_title_prefix, "CalSim 3 Water Summary Report, by Region and Type, Long-Term Average and Dry and Critical Year Averages",
                                 custom_style="Table Caption")

        # create table
        df_curr_table = dfs.loc[(comparison, ['Long Term', 'Dry and Critical']), :]
        df_curr_table.loc['Description', dfs.loc['Description', :].columns] = dfs.loc['Description', :].values
        df_curr_table.loc['Units', dfs.loc['Units', :].columns] = dfs.loc['Units', :].values

        # first do the region table
        region_table = df_curr_table[['Sacramento River Hydrologic Region', 'San Joaquin River Hydrologic Region (not including Friant-Kern and Madera Canal water users)',
                                      'San Francisco Bay Hydrologic Region', 'Central Coast Hydrologic Region', 'Tulare Lake Hydrologic Region (not including Friant-Kern Canal water users)',
                                      'South Lahontan Hydrologic Region', 'South Coast Hydrologic Region', 'Total For All Regions']]
        t = doc.add_table(2 * region_table.shape[1] + len(region_table.columns.get_level_values(0).unique()) + 1, 7)
        format_table_supply(t, region_table, doc, comparison, [21, 45])

        footnote1 = doc.add_paragraph()
        run = footnote1.add_run(
            'CVP = Central Valley Project; SWP = State Water Project; M&I = municipal and industrial; Ag = Agricultural; FRSA = Feather River Service Allocation;  TAF = thousand acre-feet.')
        run.font.size = Pt(9)
        footnote1.paragraph_format.space_before = Pt(1)
        footnote1.paragraph_format.space_after = Pt(1)

        footnote2 = doc.add_paragraph()
        run = footnote2.add_run(
            'Long-term average is the average quantity for the period of October 1921–September 2021. Dry and critical year average is the average quantity for the combination of the State Water Resources Control Board D-1641 40-30-30 dry and critical dry years for the period of October 1921–September 2021.')
        run.font.size = Pt(9)
        footnote2.paragraph_format.space_before = Pt(1)
        footnote2.paragraph_format.space_after = Pt(1)

        # next we do the north and south table
        # these are split only to get the headers to look good, functionally they are one table
        doc.add_page_break()
        add_caption_water_supply(doc, "Table", tab_title_prefix, "CalSim 3 Water Supply Summary Report, by Type, Long-Term Average and Dry and Critical Year Averages",
                                 custom_style="Table Caption")

        north_table = df_curr_table[['North of Delta', 'Total CVP North of Delta', 'Total SWP North of Delta', 'Total North of Delta']]
        t = doc.add_table(2 * north_table.shape[1] + len(north_table.columns.get_level_values(0).unique()) + 1, 7)
        format_table_supply(t, north_table, doc, comparison, [])

        doc.add_page_break()
        south_table = df_curr_table[['South of Delta', 'Total CVP South of Delta', 'Total SWP South of Delta', 'Total South of Delta']]
        t = doc.add_table(2 * south_table.shape[1] + len(south_table.columns.get_level_values(0).unique()) + 1, 7)
        format_table_supply(t, south_table, doc, comparison, [])

        footnote1 = doc.add_paragraph()
        run = footnote1.add_run(
            'CVP = Central Valley Project; SWP = State Water Project; M&I = municipal and industrial; Ag = Agricultural; FRSA = Feather River Service Allocation;  TAF = thousand acre-feet.')
        run.font.size = Pt(9)
        footnote1.paragraph_format.space_before = Pt(1)
        footnote1.paragraph_format.space_after = Pt(1)

        footnote2 = doc.add_paragraph()
        run = footnote2.add_run(
            'Long-term average is the average quantity for the period of October 1921–September 2021. Dry and critical year average is the average quantity for the combination of the State Water Resources Control Board D-1641 40-30-30 dry and critical dry years for the period of October 1921–September 2021.')
        run.font.size = Pt(9)
        footnote2.paragraph_format.space_before = Pt(1)
        footnote2.paragraph_format.space_after = Pt(1)

        doc.add_page_break()

    # create plots
    # Check for/create directory to save plots
    plot_directory = "supply_plots"

    if os.path.exists(plot_directory):
        # If the directory already exists, clear it out (Wytype names are different for trinity vs sjr and sac, so it
        # can cause issues if there's old results.
        shutil.rmtree(plot_directory, ignore_errors=True)

    # WYType Labels to use in stat plot titles. (Corresponds to the "Statistics" column value for the last 6 rows in the exceedance tables)
    fields = df_exceedances.columns
    df_exceedance_list = [df_exceedances.loc[scenario] for scenario in alts]
    fig_value = 'Average Volume (TAF)'
    line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
    line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

    # Iterate through each stat and plot month abbreivated name by EC in current type of year
    for field in fields:
        create_month_plot(df_exceedance_list, fig_value, field, plot_directory, alts, line_styles, line_colors, 'water supply')

        # Center figures in middle of page by adding some new lines above
        p = doc.add_paragraph()
        run = p.add_run()

        # Add figure as a picture
        o_fig = doc.add_picture(plot_directory + "/" + field + ".png")

        # Generate fig title
        fig_title_prefix = "Figure " + appendix_prefix + "-"
        fig_title = field

        # Add title below figure
        add_caption_water_supply(doc, "Figure", fig_title_prefix, fig_title, custom_style="Figure Caption")

        # if we are on the last plot we don't need a page break
        if field == fields[-1]:
            continue
        else:
            doc.add_page_break()

    # Save docx object to word doc
    doc.save(doc_name)

    ##### Use Python to Run VBS Script that adds alt text to table in saved docx file #######

    # Format alt text for all tables as one string to be passed to vbs
    alt_text_string_tables = ("+").join(alt_text_tables)
    alt_text_string_tables = alt_text_string_tables.replace(" ", "_")

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text_figures)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    # Run vbs script
    # Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    # This will fail if Microsoft Word has document open in the background
    # try opening Task Manager and Ending MS Word Background Task, then rerun

    try:  # Call the vbs script for table and figure alt text
        result = subprocess.call(
            "cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + alt_text_string_tables + " " + str(num_tables) + " " + alt_text_string_figures + " " + str(num_figures))

    except:  # If you have too many figures in the document, the above subprocess call will fail. Use workaround where the alt text for tables and figures are called separately.
        # Call script to add table alt text first
        result = subprocess.call(
            "cscript.exe add_alt_text.vbs " + doc_name + " " + doc_name.replace("temp.docx", 'temp2.docx') + " " + alt_text_string_tables + " " + str(num_tables) + " " + "xx" + " " + str(0))
        # Then call script to add figure alt text
        result = subprocess.call(
            "cscript.exe add_alt_text.vbs " + doc_name.replace("temp.docx", 'temp2.docx') + " " + new_doc + " " + "xxx" + " " + str(0) + " " + alt_text_string_figures + " " + str(num_figures))

    # Remove temporary doc if process ran successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")
    else:
        # Instructions on how to finish formatting numbered captions.
        print("After running this script, \n1. Open Word file and Ctrl+A to select all. Then F9 to update caption numbering.\
        \n2. For the Heading 2 Numbering, you may have to adjust it to match the appendix_prefix variable (Ex: 'F.2.2') by right clicking and selecting 'Adjust List Indents'. \nThen modify the numbering to match appendix_prefix under 'Enter formatting for number:'")


def create_appendix(report_type, alts, fields, appendix_prefix, dss_path, doc_name, new_doc, wy_flags_path, template, location_cw_path, use_calendar_yr=False, use_lumped_table_captions=False, storage_elevation_table='', compliance_fields=[], compliance_dict={}, shastabin_data_path=''):
    """
    Create a CalSim, temperature (HEC-5Q), or salinity (DSM2) appendix. Creates the tables and plots and puts them into a doccument.

    Parameters
    ----------
    report_type: str
        Report type. Could be "flow", "elevation', "diversion" (CalSim appendices), "temperature" (HEC-5Q appendix), "EC", "Cl", "Position" (salinity/DSM2 appendices).
    alts: list
        List of alternatives to include in the appendix
    fields: list
        List of fields of to include in the appendix
    appendix_prefix: str
        Prefix to use in the appendix. Typically, would be ' F.2.4'
    dss_path: str
        Path to the excel file with the DSS contents
    doc_name: str
        Name for temporary document
    new_doc: str
        Name for the final document
    wy_flags_path: str
        Path to the water year type flags file
    template: str
        Path to the template doc
    location_cw_path: str
        Path to the crosswalk file
    use_calendar_yr: bool
        Flag for if years should be grouped by calendar year (True) or water year (False)
    use_lumped_table_captions: bool
        Specify whether you want the table captions lumped or not
    storage_elevation_table: str
        Path to the storage elevation table, only used for elevation
    compliance_fields: list
        List of compliance locations, only for temperature appendix
    compliance_dict: dict
        Dictionary for which shasta bin values correspond to which compliance location
    shastabin_data_path: str
        Path to shasta bin values file

    Returns
    -------
    None
    """

    # Read location from crosswalk based on field later
    if report_type == 'elevation':
        # If the report_type is elevation, then order the fields in a specific order. (Ex: S_Trinity storages, S_Trinity elevations, etc).
        fields = order_elevation_storage_fields(
            fields)  # Returns a list of tuples with the type of field (elevation or storage). Ex: [("S_TRNTY", 'Storage'), ("S_TRNTY", 'Elevation'), ("S_SHSTA", 'Storage'),  ("S_SHSTA", 'Elevation')]
    elif report_type in ['EC', 'Position', 'Cl']:
        fields = [(field, report_type) for field in fields]
    locations = get_locations(location_cw_path, fields)  # Get location names for each field
    location_params = get_locations_params(location_cw_path, fields)  # Get the field parameter for each field (Ex: "Storage", "Elevation", "Diversion", "Delivery")
    locations_wytypes = get_location_wytypes(location_cw_path, fields)  # Get the wytype to use with each field.

    # compare every run to the baseline run
    comparisons = [["NAA", alt] for alt in alts]
    # Remove first comparison that is NAA and NAA
    comparisons.pop(0)

    """
    For each field, there are:
        - 3 comparison tables per alternative. (Ex: If you have 6 alternatives, you will have 18 tables total for S_TRNTY)
        - 12 monthly exceedance plots. 
        - full simulation period statistics plots (1 long-term avg plot and 5 plots of averages for different wy types.)
    """

    if report_type == 'temperature':
        included_compliance_fields = [f for f in fields if f in compliance_fields]
        # Each comparison will have 3 tables and will be included for every field/location
        num_tables = (len(comparisons) * 3) * len(fields) + len(included_compliance_fields) * 12
        # Include a figure for each month plus 6 full simulation period statistics plots
        num_figures = (12 + 6) * len(fields)
    else:
        # Each comparison will have 3 tables and will be included for every field/location
        num_tables = (len(comparisons) * 3) * len(fields)
        # Include a figure for each month plus 6 full simulation period statistics plots
        num_figures = (12 + 6) * len(fields)

    # Alt Text strings, in order for tables
    alt_text_tables = ["Alt text table example" for t in range(0, num_tables)]

    # Alt text strings, order for figures
    alt_text_figures = ["This figure shows data also presented in data tables in this file." for f in range(0, num_figures)]

    # Create an instance of a word document
    # Open the word document template. This template has the heading style 2 formatted with numbering to allow the figures
    # to inherit the heading numbering.
    doc = docx.Document(template)
    doc.add_heading(f"Attachment{appendix_prefix}", level=1)  # Add Heading 1 (Attachment XXX)

    # Add caption style for Figure captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.name = 'Times New Roman'

    # Add caption style for Table captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    for field_index, location in enumerate(fields):
        if field_index ==0:
            doc.add_page_break()

        #Add location heading in default word heading 2 style. This allows the figure numbering to inherit the heading 2 numbering.
        doc.add_heading(locations[field_index], level=2)

        ##### Read DSSReader output ########
        if report_type == 'elevation':
            #For the elevation report
            if location[1] == "Storage":
                dfs = parse_dssReader_output(dss_path, alts, location[0], report_type)
            elif location[1] == 'Elevation':
                # converted to elevation (ft) based on the storage elevation relationship from res_info.table in the CalSim 3 wresl code.
                dfs = parse_dssReader_output(dss_path, alts, location[0], report_type, convert_to_elevation= True, orig_unit = 'TAF', storage_elevation_fn = storage_elevation_table)
        elif report_type == 'Cl':
            dfs  = parse_dssReader_output(dss_path, alts, location[0], report_type, convert_to_cl= True, orig_unit = 'uS/cm')
        elif report_type in ["EC", "Position"]:
            dfs = parse_dssReader_output(dss_path, alts, location[0], report_type)
        else:
            dfs = parse_dssReader_output(dss_path, alts, location, report_type)

        if location in compliance_fields:
            dfs_calendaryr = parse_dssReader_calendaryr(dss_path, alts, location, report_type, shastabin_data = shastabin_data_path, use_calendar_yr=use_calendar_yr)

        # Get table value name depending on type of report
        if report_type == "flow":
            unit = 'cfs'
            table_value = "Monthly Flow (cfs)"
        elif report_type == "elevation":
            if location[1] == 'Elevation':
                unit = 'feet'
                table_value = "End of Month Elevation (feet)"
            elif location[1] == 'Storage':
                unit = 'TAF'
                table_value = 'End of Month Storage (TAF)'
        elif report_type == 'diversion':
            unit = 'cfs'
            table_value = f"Monthly {location_params[field_index]} (cfs)"
        elif report_type == 'temperature':
            unit = 'DEG-F'
            table_value = f"Monthly Temperature (DEG-F)"
        elif report_type == 'EC':
            unit = "UMHOS/CM"
            table_value = f"Monthly EC (UMHOS/CM)"
        elif report_type == 'Cl':
            unit = "mg/L"
            table_value = r"Monthly Cl (mg/L)"
        elif report_type == 'Position':
            unit = "KM"
            table_value = r"Monthly Position (KM)"
        else:
            raise ValueError(f"No report type for {report_type} is available. Needs to be implemented.")

        # Get figure value name depending on type of report. This is used in the stat figure captions.
        fig_value = f"Average {location_params[field_index]} ({unit})"

        #Create Exceedance Tables from DSS Reader output
        e_dfs, exc_prob, fig_dfs,il_num_years= create_exceedance_tables(dfs, wy_flags_path, locations_wytypes[field_index], report_type, use_calendar_yr = use_calendar_yr)

        ##### Use docx package to create a document with formatted table objects and save to Word .docx file ###########

        ## Add a table for each run in each comparison for the current field to the doc
        for comparison_index, scenario in enumerate(comparisons):

            #Then third table for each comparison should be first alt minus second alt listed
            comparison_tables = []
            for alt in scenario:
                #Get exceedance tables for each of the runs in the current comparison
                comparison_tables.append(e_dfs[alts.index(alt)])
            #Add one more table for second alt minus the baseline
            comparison_tables.append(comparison_tables[1].iloc[:, 1:] - comparison_tables[0].iloc[:, 1:])
            #Add the labels column back into the differenced table
            comparison_tables[-1].insert(0, "Statistic", comparison_tables[0]["Statistic"])

            #Set up Comparison labels to be used in table titles
            comparison_table_labels = ["NAA", scenario[1], scenario[1] + " Minus " + "NAA"]

            for comp_table_index, full_table in enumerate(comparison_tables):
                #Subset the statistics table to exclude the lowest and highest probability of exceedance (usually 1% and 99% exceedance)
                table = full_table.loc[~full_table.Statistic.isin([f'{round(exc_prob.iloc[0])}% Exceedance', f'{round(exc_prob.iloc[-1])}% Exceedance'])].copy(deep = True)
                table.reset_index(inplace = True,drop = True)

                table_letter = chr(ord('a') + comp_table_index)
                #table_number = str(comparison_index + 1) #Track the table number you are currently using.

                # Generate table title
                table_title_prefix = "Table" + appendix_prefix + "-"

                table_title = locations[field_index] + ", " + comparison_table_labels[comp_table_index] + ", " + table_value


                #table_title = "Table " + appendix_prefix + "-" + str(field_index + 1) + "-" + str(comparison_index + 1) + chr(ord('a') + comp_table_index)  +". " + locations[field_index] + ", " + comparison_table_labels[comp_table_index] + ", " + table_value

                # Add caption above table
                if not use_lumped_table_captions:
                    if table_letter == 'a': #If this is the first table of the 3 comparison tables, then use the next sequential table number + the letter a
                        add_caption_byfield(doc, "Table", table_title_prefix, table_letter +". " + table_title, custom_style="Table Caption")
                    else: #If this is not the first table of the 3 comparison tables, use the
                        add_caption_byfield(doc, 'Table', table_title_prefix, table_letter +". " + table_title, custom_style='Table Caption', use_prev_number= True)

                # add a table to the end and create a reference variable
                # extra row is so we can add the header row
                t = doc.add_table(table.shape[0] + 1, table.shape[1])
                #Format table for report
                format_table(t, table, doc, report_type)

            #Get the number of years of simulation record from the full exceedance probability/values dataframe for each of the naa and the alt you are comparing
            il_sample_sizes =[]
            for alt in scenario:
                #Averaging the number of samples we have for each month
                # gives you approximation of the full period of record length in years.
                il_sample_sizes.append(np.mean(il_num_years[alts.index(alt)]).tolist())

            #Determine the period of record footnote to include.
            #If the NAA and alternative you are comparing to have different sample sizes, use this footnote.
            if len(np.unique(il_sample_sizes))!=1:
                s_por_footnote  = f"{scenario[0]} Statistics based on approximately {round(il_sample_sizes[0], 1)}-year simulation period. {scenario[1]} statistics based on approximately {round(il_sample_sizes[1], 1)}-year simulation period."
            #If the NAA and alternative you are comparing to have the same sample size and it is a whole number of years, then use this footnote.
            elif il_sample_sizes[0] == int(il_sample_sizes[0]):
                s_por_footnote = f" Based on the {int(il_sample_sizes[0])}-year simulation period."
            # If the NAA and alternative you are comparing to have the same sample size and it includes a fraction of a year, then use this footnote.
            else:
                s_por_footnote = f" Based on the {round(il_sample_sizes[0],1)}-year simulation period."

            # Add footnotes to the final table
            if comp_table_index == (len(comparison_tables) - 1):
                # Add footnotes at end of table
                footnote0 = doc.add_paragraph()
                run = footnote0.add_run("a")
                run.font.superscript = True
                run1 = footnote0.add_run(s_por_footnote)
                run1.font.size = Pt(9)
                footnote0.paragraph_format.space_after = Pt(1)

                #Add footnote specifying hydrology
                footnote1 = doc.add_paragraph()
                run = footnote1.add_run('* All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
                run.font.size = Pt(9)
                footnote1.paragraph_format.space_before = Pt(1)
                footnote1.paragraph_format.space_after = Pt(1)

                #Add footnote specifying what WY type this field's table uses.
                footnote2 = doc.add_paragraph()
                if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:
                    run = footnote2.add_run(
                    f'* Water Year Types defined by the {locations_wytypes[field_index]} Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
                else:
                    run = footnote2.add_run(f"* Water Year Types defined by the Trinity Water Year Hydrologic Classification.")
                run.font.size = Pt(9)
                footnote2.paragraph_format.space_before = Pt(1)
                footnote2.paragraph_format.space_after = Pt(1)

                #Add footnote for water year type sorting method.
                footnote3 = doc.add_paragraph()
                if not use_calendar_yr:
                    run = footnote3.add_run('* Water Year Types results are displayed with water year – year type sorting.')
                else:
                    run = footnote3.add_run('* Water Year Types results are displayed with calendar year – year type sorting.')
                run.font.size = Pt(9)
                footnote3.paragraph_format.space_before = Pt(1)
            if comparison_index!=0:
                doc.add_page_break() #Add page break after the a,b,c comparison tables.

        #####Create Monthly EC and full simulation period statistic plots, save locally as images#####

        #Individual Month Plots tables are in fig_dfs

       #Format percent exceedances for labels
        exc_percents = [str(round(x)).split(".")[0] + "%" for x in exc_prob.values]
        ##Remove simulation period statistic rows
        # for fig_index in range(len(fig_dfs)):
        #     fig_dfs[fig_index] = fig_dfs[fig_index][:-6]
        #     #Add formatted exceedance probability percents back to dfs
        #     fig_dfs[fig_index]["exc_prob"] = exc_percents

        #Can plot up to 8 scenarios, these lines prepare linestyle and color
        line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
        line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

        # Flip doc to landscape orientation for images
        change_orientation(doc, "landscape")

        #Iterate through the dfs and create a figure for each month
        #Save month plots to directory
        month_directory = "month_plots"

        if os.path.exists(month_directory):
            # If the directory already exists, clear it out to prevent using any old figures by accident from previous field/alternative.
            shutil.rmtree(month_directory, ignore_errors=True)

        monthly_ranked_dfs = {}
        for month in fig_dfs[0].columns[1:]:
            if location in compliance_fields:
                #for compliance fields, make exceedance plots with the compliance years marked with a marker.
                df_month_alts  = create_mixed_compliance_month_plots(location, dfs_calendaryr, fig_value, month, month_directory, alts, line_styles, line_colors, compliance_dict)
                monthly_ranked_dfs[month] = df_month_alts
            else:

                #Create monthly plot. For compliance locations, a red marker will be plotted for the
                create_month_plot(dfs, fig_value, month, month_directory, alts, line_styles, line_colors)

        ##Simulation Period Statistic Plots###
        stat_fig_dfs = copy.deepcopy(e_dfs)

        for stat_fig_index in range(len(stat_fig_dfs)):
            #keep only simulation period statistic rows
            stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index][-6:]
            #Transpose to plot all months at once
            stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index].transpose()
            #Drop first row
            stat_fig_dfs[stat_fig_index].rename(columns=stat_fig_dfs[stat_fig_index].iloc[0], inplace=True)
            stat_fig_dfs[stat_fig_index].drop(stat_fig_dfs[stat_fig_index].index[0], inplace=True)
            #Add abbreviated month name column
            stat_fig_dfs[stat_fig_index]["month"] = ["Oct", "Nov", "Dec", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug",
                                        "Sep"]

        #Check for/create directory to save stat plots
        stat_directory = "stat_plots"

        if os.path.exists(stat_directory):
            # If the directory already exists, clear it out (Wytype names are different for trinity vs sjr and sac, so it
            # can cause issues if there's old results.
            shutil.rmtree(stat_directory, ignore_errors=True)

        #WYType Labels to use in stat plot titles. (Corresponds to the "Statistics" column value for the last 6 rows in the exceedance tables)
        stats = e_dfs[0].Statistic[-6:].values.tolist()
        # stats = ["Full Simulation Period", "Wet Water Years (28%)", "Above Normal Water Years (14%)",
        #          "Below Normal Water Years (18%)",
        #          "Dry Water Years (24%)", 'Critical Water Years (16%)']

        #Iterate through each stat and plot month abbreivated name by EC in current type of year
        for stat in stats:
            create_stat_plot(stat_fig_dfs, fig_value, stat, stat_directory, alts, line_styles, line_colors)

        ##Add saved figures to docx object as images####

        #Get saved month plots, in order from Oct - Sept.
        month_plots = [rf"{str(m).zfill(2)}_{datetime.strptime(str(m), '%m').strftime('%b')}_monthly_exceedance.png" for m in [10,11,12,1,2,3,4,5,6,7,8,9]]

        #Iterate through each monthly figure in the month plots directory
        for month_index, file in enumerate(month_plots):
            # Center figures in middle of page by adding some new lines above
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()
            run.add_break()

            #Add figure as a picture
            o_fig = doc.add_picture(month_directory + "/" + file)

            # Add captions below figure
            f = doc.add_paragraph()
            run = f.add_run(
                '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
            run.font.size = Pt(9)
            f.paragraph_format.space_before = Pt(1)
            f.paragraph_format.space_after = Pt(1)

            # Generate fig title
            fig_title_value = location_params[field_index]
            fig_title_prefix = "Figure " + appendix_prefix + "-"
            fig_title = locations[field_index] + ", " + datetime.strptime(file.split("_", 2)[1],
                                                                                   '%b').strftime(
                '%B') + " " + fig_title_value
            # Add title below figure
            add_caption_byfield(doc, "Figure", fig_title_prefix, fig_title, custom_style = "Figure Caption")

            #Add page break after each figure
            doc.add_page_break()

            #After each figure, add a table of the yearly value rankings (monthly_ranked_dfs)
            if location in compliance_fields:
                #Change orientation to be portrait for the tables.
                change_orientation(doc, "portrait")
                month_name = file.split("_", 2)[1]
                table = monthly_ranked_dfs[month_name].reset_index(drop = False)
                table.rename(columns = {'Exc Prob': "Exceedance Probability (%)"}, inplace = True)
                month_str = datetime.strptime(file.split("_", 2)[1],'%b').strftime('%B')
                if use_calendar_yr:
                    table_title = f"Calendar Year corresponding to {month_str} exceedance values at {locations[field_index]}, for each alternative."
                else:
                    table_title = f"Water Year corresponding to {month_str} exceedance values at {locations[field_index]}, for each alternative."
                #Add table caption
                add_caption_byfield(doc, "Table", table_title_prefix, table_title,
                                    custom_style="Table Caption")

                #Add table to document
                t = doc.add_table(table.shape[0] + 1, table.shape[1])

                # Format table for report
                format_table_basic(t, table, doc)
                doc.add_page_break()

                # Flip orientation back to landscape for the rest of the plots
                change_orientation(doc, "landscape")

        # Add stats plots as well
        #Set the statistics plot titles
        if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:  #For Sac or SJR WYType
            stat_titles = ["Long Term", "Wet Year", "Above Normal Year", "Below Normal Year", "Dry Year", 'Critical Year']
        else: #For Trinity WYType
            stat_titles = ["Long Term", "Extremely Wet Year", "Wet Year", "Normal Year", "Dry Year", "Critically Dry Year"]

        for stat_plot_index, stat_title in enumerate(stat_titles):
            # Center figures in middle of page by adding some new lines above
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()
            run.add_break()

            #Add stat figure as image to document
            file = stat_title[:5] + "_Exceedance.png" if stat_title!= 'Long Term' else "Full _exceedance.png"
            doc.add_picture(stat_directory + "/" + file)

            # Add footnotes below figure about water year type definition
            caption0 = doc.add_paragraph()
            if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:
                run = caption0.add_run(
                    f'*As defined by the {locations_wytypes[field_index]} Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
            else:
                run = caption0.add_run(
                    f"*As defined by the Trinity Water Year Hydrologic Classification.")
            run.font.size = Pt(9)
            caption0.paragraph_format.space_before = Pt(1)
            caption0.paragraph_format.space_after = Pt(1)

            # Add footnote for what wy type sorting is used.
            caption1 = doc.add_paragraph()
            if not use_calendar_yr:
                run = caption1.add_run('*These results are displayed with water year - year type sorting.')
            else:
                run = caption1.add_run('*These results are displayed with calendar year - year type sorting.')
            run.font.size = Pt(9)
            caption1.paragraph_format.space_before = Pt(1)
            caption1.paragraph_format.space_after = Pt(1)

            #Add footnotes below figure about climate change scenario
            caption2 = doc.add_paragraph()
            run = caption2.add_run(
                '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
            run.font.size = Pt(9)
            caption2.paragraph_format.space_before = Pt(1)

            # Generate fig title
            fig_title_prefix = "Figure " + appendix_prefix + "-"
            fig_title = locations[field_index] + ", " +  stat_title + " " + fig_value
            #Add fig title as the figure caption below figure.
            add_caption_byfield(doc, "Figure", fig_title_prefix, fig_title, custom_style="Figure Caption")

            #No need for the page break if it's the final plot of the document
            if stat_plot_index == (len(stat_titles) - 1) and field_index == (len(fields) - 1):
                continue
            else:
                doc.add_page_break()

            #Flip orientation back to portrait for the next group of tables
            if stat_plot_index == (len(stat_titles) - 1):
                # Flip doc to landscape orientation for images
                change_orientation(doc, "portrait")

    # Save docx object to word doc
    doc.save(doc_name)

    ##### Use Python to Run VBS Script that adds alt text to table in saved docx file #######

    # Format alt text for all tables as one string to be passed to vbs
    alt_text_string_tables = ("+").join(alt_text_tables)
    alt_text_string_tables = alt_text_string_tables.replace(" ", "_")

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text_figures)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    #Run vbs script
    #Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    #This will fail if Microsoft Word has document open in the background
    #try opening Task Manager and Ending MS Word Background Task, then rerun

    try: #Call the vbs script for table and figure alt text
        result = subprocess.call("cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + alt_text_string_tables +  " " + str(num_tables) + " " + alt_text_string_figures + " " + str(num_figures))

    except: #If you have too many figures in the document, the above subprocess call will fail. Use workaround where the alt text for tables and figures are called separately.
        #Call script to add table alt text first
        result = subprocess.call("cscript.exe add_alt_text.vbs " + doc_name + " " + doc_name.replace("temp.docx", 'temp2.docx') + " " + alt_text_string_tables + " " + str( num_tables) + " " + "xx"+ " " + str(0))
        #Then call script to add figure alt text
        result = subprocess.call(
            "cscript.exe add_alt_text.vbs " + doc_name.replace("temp.docx", 'temp2.docx') + " " + new_doc + " " + "xxx" + " " + str(0) + " " + alt_text_string_figures + " " + str(num_figures))

    #Remove temporary doc if process ran successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")
    else:
        #Instructions on how to finish formatting numbered captions.
        print("After running this script, \n1. Open Word file and Ctrl+A to select all. Then F9 to update caption numbering.\
        \n2. For the Heading 2 Numbering, you may have to adjust it to match the appendix_prefix variable (Ex: 'F.2.2') by right clicking and selecting 'Adjust List Indents'. \nThen modify the numbering to match appendix_prefix under 'Enter formatting for number:'")


### The following function are meant to recreate what get_dsm2_comp_data_20240221.r from Jacobs does
### They are for the water quality compliance attachment

def get_stations():
    """
    Gets the stations for the compliance locations. These are the DSS b parts.
    Returns
    -------
    stations: list
    List of stations
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 0].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_compliance_locations():
    """
    Gets the location names for the compliance locations
    Returns
    -------
    stations: list
    List of location names
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 1].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_compliance_stats():
    """
    Gets the statistics for the compliance locations. (MIN or MAX or MEAN etc)
    Returns
    -------
    stations: list
    List of stats
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 2].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_sri_current_condition():
    """
    Gets the dataframe of sac river index data
    Returns
    -------
    df_sri: dataframe
    dataframe of sac river index data for each year
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    df_sri = None

    # Look for the SacRiverIndex.csv file
    for j in range(len(fnames)):
        if fnames[j] == "SacRiverIndex.csv":
            print("\nFound file:", fnames[j])
            df_sri = pd.read_csv(fpaths[j], header=0, sep=",")
            print("\n")
            break  # Exit loop once the file is found

    return df_sri


def get_wyts(s_wyt_path):
    """
    Gets the water year type for each year depending on the given file name
    Parameters
    ----------
    s_wyt_path: str
        Name of wyt file

    Returns
    -------
    wyts_df: dataframe
        dataframe of water year type for each year depending on the given file name
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    wyts_df = None

    # Look for the WYT_2022MED.csv file
    for j in range(len(fnames)):
        if fnames[j] == s_wyt_path:
            wyts_df = pd.read_csv(fpaths[j], header=0, sep=",")
            break  # Exit loop once the file is found

    return wyts_df


def get_specified_table(s_csv_name):
    """
    Reads in the specified csv file and returns a dataframe
    Parameters
    ----------
    s_csv_name: str
        Name of csv file

    Returns
    -------
    std_df: dataframe
        Dataframe of data from the specified csv file
    """
    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    std_df = None

    # Look for the D1641_AG.csv file
    for j in range(len(fnames)):
        if fnames[j] == s_csv_name:
            std_df = pd.read_csv(fpaths[j], header=0, sep=",")
            break  # Exit loop once the file is found

    return std_df


def get_dsm2_timeseries_data(file_path):
    """
    Function to do the main work for the water quality compliance appendix.
    This will read in the DSS file and all the needed compliance info and create two csvs with the output data.

    Parameters
    ----------
    file_path: str
        Path to the DSS file

    Returns
    -------
    None
    """

    # Construct the path to the model's directory
    file_path = os.path.join("..", "studies", file_path)

    input_model_name = os.path.basename(file_path).split('.')[0]

    print("Model Name:", input_model_name)

    # get the WYT data based on the hydrology
    if "2030" in input_model_name:
        df_wyt = get_wyts("WYT_2030.csv")
    elif "2070" in input_model_name:
        df_wyt = get_wyts("WYT_2070.csv")
    elif "2022" in input_model_name:
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_2022MED.csv")
    elif "Current" in input_model_name:
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_CurrentConditions.csv")
    else:
        print('Hydrology is undefined in study name, using current')
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_CurrentConditions.csv")

    # Load csv files for compliance standards
    df_D1641AG = get_specified_table("D1641_AG.csv")
    df_D1641FWS = get_specified_table("D1641_FWS.csv")
    df_D1641MI = get_specified_table("D1641_MI.csv")
    df_D1641MID = get_specified_table("D1641_MID.csv", )
    df_MI_Antioch = get_specified_table("MI_Antioch.csv")
    df_MI_Other = get_specified_table("MI_Other.csv")
    comp_files = [df_D1641AG,
                  df_D1641FWS,
                  df_D1641MI,
                  df_D1641MID,
                  df_MI_Antioch,
                  df_MI_Other]
    comp_names = ["D1641 AG",
                  "D1641 FWS",
                  "D1641 MI",
                  "D1641 MID",
                  "MI Antioch",
                  "MI Other"]

    with HecDss.Open(file_path) as input_file:
        # Retreive stations and locations
        out_statns = get_stations()
        out_locs = get_compliance_locations()
        out_stats = get_compliance_stats()

        outfile_location = "./water_qual_csvs"
        if not os.path.exists(outfile_location):
            os.makedirs(outfile_location)

        # Construct the output filename
        tsfilename = f"DSM2ComplianceData_{input_model_name}.csv"
        tsfilename = os.path.join(outfile_location, tsfilename)

        # Open the file for writing
        with open(tsfilename, 'w') as tsfile:
            # Write the header line
            header = "Var Name,Location,Var type,Date,ValueEC,ValueCl,Study Scenario,Study Type,UnitsEC,UnitsCl,D1641AG,D1641FWS,D1641MI,D1641MIDNumDays,D1641MIDThreshold,MIAntiochNumDays,MIAntiochThreshold,MIOther,SAC INDEX\n"

            tsfile.write(header)

        nd_flags = []

        # Loop through locations
        for station_index, station in enumerate(out_statns):
            print(station)

            cl_flag = 0
            nd_flag = 0
            def_stn_flag = 0
            sjr_fws_stn_flag = 0

            bpart = station

            # Set flags based on station name
            if bpart in ["SLCBN002", "SLSUS012"]:
                def_stn_flag = 1
            if bpart in ["RSAN018", "RSAN032", "RSAN037"]:
                sjr_fws_stn_flag = 1
            # Create search path
            cpart = f"EC-{out_stats[station_index]}"
            path = f"/*/{bpart}/{cpart}/*/1DAY/*/"
            path_list = input_file.getPathnameList(path, sort=1)

            if path_list == []:
                continue

            expanded_path = path_list[0].split('/')
            expanded_path[4] = ''
            path = "/".join(expanded_path)
            o_timeseries = input_file.read_ts(path)

            if o_timeseries.empty:
                continue

            dates = np.array(o_timeseries.pytimes)

            # mimikling the date shifting that the R code does
            # python reads the dates in one day later than R does so we subtract two days and then another for the years and months
            dates = dates + timedelta(days=-2)

            years = [(date + timedelta(days=-1)).strftime("%Y") for date in dates]
            months = [(date + timedelta(days=-1)).strftime("%m") for date in dates]
            days = [(date + timedelta(days=-1)).strftime("%d") for date in dates]

            values = np.round(o_timeseries.values, 5)

            unitsEC = ["mmhos/cm"] * len(values)
            unitsCl = ["mg/L"] * len(values)
            stations = [out_statns[station_index]] * len(values)
            locations = [out_locs[station_index]] * len(values)
            cpart = [cpart] * len(values)
            study_scenario = [input_model_name] * len(values)
            study_type = ["DSM2"] * len(values)

            # create this data frame to do some matching quicker than a loop
            df_wyt_sri = pd.DataFrame({'Year': years, 'Month': months})
            df_wyt_sri = df_wyt_sri.astype(int)

            # get the water year and the previous year
            df_wyt_sri['WY'] = np.where(df_wyt_sri['Month'] < 10, df_wyt_sri['Year'], df_wyt_sri['Year'] + 1)
            df_wyt_sri['Prev Year'] = df_wyt_sri['Year'] - 1

            # Merge with df_wyt on water year
            df_wyt_sri = df_wyt_sri.merge(df_wyt.rename(columns={df_wyt.columns[1]: 'wy_ind', df_wyt.columns[2]: 'wyt'}),
                                          left_on='WY', right_on='YEAR', how='left')

            # Merge to get current SRI
            df_wyt_sri = df_wyt_sri.merge(df_sri[[df_sri.columns[0], df_sri.columns[3]]].rename(columns={df_sri.columns[3]: 'curr_sri'}),
                                          left_on='Year', right_on='WaterYear', how='left')

            # Merge to get previous SRI
            df_wyt_sri = df_wyt_sri.merge(df_sri[[df_sri.columns[0], df_sri.columns[3]]].rename(columns={df_sri.columns[3]: 'prev_sri'}),
                                          left_on='Prev Year', right_on='WaterYear', how='left')

            # Extract final lists
            wyts = df_wyt_sri['wyt'].tolist()
            wy_inds = df_wyt_sri['wy_ind'].tolist()
            prev_sri = df_wyt_sri['prev_sri'].tolist()
            curr_sri = df_wyt_sri['curr_sri'].tolist()

            # Append the last wyt value to the list
            wyts.append(wyts[-1])
            wyts = wyts[1:]

            # Print unique values
            print("Unique WYT values:", list(set(wyts)))

            # Initialize list for compliance DataFrames
            loc_comp_dfs = [None] * len(comp_names)

            # Initialize compliance count list with zeros
            comp_count = [0] * len(comp_names)

            for compliance_index in range(len(comp_files)):
                comp_df = comp_files[compliance_index]

                # Filter rows where Var Name matches the current station
                loc_comp_df = comp_df[comp_df["Var Name"] == out_statns[station_index]]

                if not loc_comp_df.empty:
                    if out_statns[station_index] == "RSAN007":
                        if out_stats[station_index] == "MAX" and comp_names[compliance_index] == "MI Antioch":
                            print(comp_names[compliance_index])
                            comp_count[compliance_index] = 1
                            loc_comp_dfs[compliance_index] = loc_comp_df
                        elif out_stats[station_index] == "MEAN" and comp_names[compliance_index] != "MI Antioch":
                            print(comp_names[compliance_index])
                            comp_count[compliance_index] = 1
                            loc_comp_dfs[compliance_index] = loc_comp_df
                    else:
                        print(comp_names[compliance_index])
                        comp_count[compliance_index] = 1
                        loc_comp_dfs[compliance_index] = loc_comp_df

            # Loop through time series and set standard
            dates_copy = dates.copy()
            std_nms = []
            # std_ts_df = pd.DataFrame()
            count = 0
            for compliance_index in range(len(comp_count)):
                print('comp_count[compliance_index]: ' + str(comp_count[compliance_index]))
                if comp_count[compliance_index] < 1:
                    print('less than 1')
                    std_ts = [np.nan] * len(years)
                    std_nm = comp_names[compliance_index]

                    if std_nm in ["D1641 MID", "MI Antioch"]:
                        std_ts2 = std_ts

                        if count < 1:
                            std_ts_df = pd.DataFrame({
                                f"{std_nm}NumDays": std_ts,
                                f"{std_nm}Threshold": std_ts2
                            })
                            std_nms.append(str(std_nm))
                            count = 1
                        else:
                            std_ts_df[f"{std_nm}NumDays"] = std_ts
                            std_ts_df[f"{std_nm}Threshold"] = std_ts2
                            std_nms.append(str(std_nm))


                    else:
                        if count < 1:
                            std_ts_df = pd.DataFrame({std_nm: std_ts})
                            std_nms.append(str(std_nm))
                            count = 1
                        else:
                            std_ts_df[std_nm] = std_ts
                            std_nms.append(str(std_nm))

                    print("std_ts_df")
                    print(std_ts_df)
                    print("\n")
                else:
                    print("\ncomp_count[compliance_index] >= 1")
                    df_comp = loc_comp_dfs[compliance_index]
                    print(df_comp)
                    std_nm = comp_names[compliance_index]

                    std_ts = []
                    std_ts2 = []

                    if not pd.isna(df_comp["NumDays"].iloc[0]):
                        print("NumDays")
                        cl_flag = 1
                        nd_flag = 1

                        for year_index in range(len(years)):
                            if year_index == 0:
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                NumDays = wyt_data["NumDays"].iloc[0]
                                Threshold = wyt_data["Val1"].iloc[0]
                                NumDays_next = NumDays
                                Threshold_next = Threshold

                                if int(months[year_index]) < 10:
                                    wy_prev = years[year_index]
                                    yr_prev = years[year_index]
                                else:
                                    wy_prev = str(int(years[year_index]) + 1)
                                    yr_prev = years[year_index]

                            yr = years[year_index]
                            wy = years[year_index] if int(months[year_index]) < 10 else str(int(years[year_index]) + 1)

                            if wy != wy_prev:
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                NumDays_next = wyt_data["NumDays"].iloc[0] if not wyt_data.empty else np.nan
                                Threshold_next = wyt_data["Val1"].iloc[0] if not wyt_data.empty else np.nan
                                wy_prev = wy

                            if yr != yr_prev:
                                NumDays = NumDays_next
                                Threshold = Threshold_next
                                yr_prev = yr

                            std_ts.append(NumDays)
                            std_ts2.append(Threshold)

                        # Remove first placeholder if needed (mimicking R's append + [-1])
                        std_ts = std_ts[1:] if len(std_ts) > len(years) else std_ts
                        std_ts2 = std_ts2[1:] if len(std_ts2) > len(years) else std_ts2

                        # Add to std_ts_df
                        if count < 1:
                            std_ts_df = pd.DataFrame({
                                f"{std_nm}NumDays": std_ts,
                                f"{std_nm}Threshold": std_ts2
                            })
                            std_nms.append(std_nm)
                            count = 1
                        else:
                            std_ts_df[f"{std_nm}NumDays"] = std_ts
                            std_ts_df[f"{std_nm}Threshold"] = std_ts2
                            std_nms.append(std_nm)

                    else:
                        if std_nm == "D1641 MI":
                            cl_flag = 1
                        for year_index in range(len(years)):
                            yr_count = 0
                            if year_index == 0:
                                # Reset deficiency flag
                                def_flag = 0

                                # Filter compliance data for the current water year index
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]

                                # Parse start and end dates for each compliance window

                                start_1 = datetime.strptime(wyt_data['Start Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 1'].iloc[0]) else None
                                end_1 = datetime.strptime(wyt_data['End Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 1'].iloc[0]) else None
                                start_2 = datetime.strptime(wyt_data['Start Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 2'].iloc[0]) else None
                                end_2 = datetime.strptime(wyt_data['End Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 2'].iloc[0]) else None
                                start_3 = datetime.strptime(wyt_data['Start Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 3'].iloc[0]) else None
                                end_3 = datetime.strptime(wyt_data['End Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 3'].iloc[0]) else None
                                start_4 = datetime.strptime(wyt_data['Start Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 4'].iloc[0]) else None
                                end_4 = datetime.strptime(wyt_data['End Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 4'].iloc[0]) else None
                                start_5 = datetime.strptime(wyt_data['Start Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 5'].iloc[0]) else None
                                end_5 = datetime.strptime(wyt_data['End Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 5'].iloc[0]) else None
                                start_6 = datetime.strptime(wyt_data['Start Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 6'].iloc[0]) else None
                                end_6 = datetime.strptime(wyt_data['End Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 6'].iloc[0]) else None

                                # Store previous year and water year index values
                                year_prev = int(years[year_index])
                                wy_ind_prev = int(wy_inds[year_index])
                                wy_ind_prev2 = int(wy_inds[year_index])
                            year = years[year_index]
                            if int(year) != year_prev:
                                print(f"\nYear: {years[year_index]}")
                                yr_count += 1
                                def_flag = 0

                                # Get new WYT data for this year
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                print(f"WY index: {wy_inds[year_index]}")
                                print(f"Prev Sac River index: {prev_sri[year_index]}\n")

                                # Deficiency logic
                                if wy_inds[year_index] == 5 and wy_ind_prev >= 4:
                                    def_flag = 1
                                elif wy_inds[year_index] == 4 and prev_sri[year_index] < 11.35:
                                    def_flag = 1
                                elif wy_inds[year_index] == 4 and wy_ind_prev >= 4 and wy_ind_prev2 == 5 and yr_count > 1:
                                    def_flag = 1

                                # SJR FWS logic
                                sjr_fws_flag = 1 if wy_inds[year_index] == 4 and curr_sri[year_index] < 8.1 else 0

                                # Set compliance windows
                                if def_flag == 1 and def_stn_flag == 1:
                                    print("Deficiency")
                                    start_1 = datetime.strptime(f"1-Jan", "%d-%b")
                                    end_1 = datetime.strptime(f"31-Mar", "%d-%b")
                                    start_2 = datetime.strptime(f"1-Apr", "%d-%b")
                                    end_2 = datetime.strptime(f"30-Apr", "%d-%b")
                                    start_3 = datetime.strptime(f"1-May", "%d-%b")
                                    end_3 = datetime.strptime(f"31-May", "%d-%b")
                                    start_4 = datetime.strptime(f"1-Oct", "%d-%b")
                                    end_4 = datetime.strptime(f"31-Oct", "%d-%b")
                                    start_5 = datetime.strptime(f"1-Nov", "%d-%b")
                                    end_5 = datetime.strptime(f"30-Nov", "%d-%b")
                                    start_6 = datetime.strptime(f"1-Dec", "%d-%b")
                                    end_6 = datetime.strptime(f"31-Dec", "%d-%b")
                                else:

                                    start_1 = datetime.strptime(wyt_data['Start Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 1'].iloc[0]) else None
                                    end_1 = datetime.strptime(wyt_data['End Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 1'].iloc[0]) else None
                                    start_2 = datetime.strptime(wyt_data['Start Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 2'].iloc[0]) else None
                                    end_2 = datetime.strptime(wyt_data['End Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 2'].iloc[0]) else None
                                    start_3 = datetime.strptime(wyt_data['Start Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 3'].iloc[0]) else None
                                    end_3 = datetime.strptime(wyt_data['End Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 3'].iloc[0]) else None
                                    start_4 = datetime.strptime(wyt_data['Start Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 4'].iloc[0]) else None
                                    end_4 = datetime.strptime(wyt_data['End Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 4'].iloc[0]) else None
                                    start_5 = datetime.strptime(wyt_data['Start Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 5'].iloc[0]) else None
                                    end_5 = datetime.strptime(wyt_data['End Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 5'].iloc[0]) else None
                                    start_6 = datetime.strptime(wyt_data['Start Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 6'].iloc[0]) else None
                                    end_6 = datetime.strptime(wyt_data['End Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 6'].iloc[0]) else None

                                # Adjust end_1 if SJR FWS condition is met
                                if sjr_fws_flag == 1 and sjr_fws_stn_flag == 1:
                                    end_1 = datetime.strptime(f"30-Apr", "%d-%b")

                                # Update previous year/index trackers
                                year_prev = int(years[year_index])
                                wy_ind_prev2 = wy_ind_prev
                                wy_ind_prev = int(wy_inds[year_index])

                            # Add 8 hours to start
                            start_plus = timedelta(hours=8)

                            # Add one day for days marked 12-31
                            end_plus = timedelta(hours=32)

                            # Current date being evaluated
                            now = dates_copy[year_index]

                            val_found = 0
                            reg_val = None

                            if start_1 is not None:
                                start_1 = start_1.replace(year=int(years[year_index]))
                                end_1 = end_1.replace(year=int(years[year_index]))
                                if now >= start_1 + start_plus and now < end_1 + end_plus:
                                    reg_val = wyt_data["Val1"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 15.6
                                    val_found = 1
                                end_1 = end_1.replace(year=2017)
                            if start_2 is not None and val_found == 0:
                                start_2 = start_2.replace(year=int(years[year_index]))
                                end_2 = end_2.replace(year=int(years[year_index]))
                                if now >= start_2 + start_plus and now < end_2 + end_plus:
                                    reg_val = wyt_data["Val2"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 14.0
                                    val_found = 1
                                end_2 = end_2.replace(year=2017)

                            if start_3 is not None and val_found == 0:
                                start_3 = start_3.replace(year=int(years[year_index]))
                                end_3 = end_3.replace(year=int(years[year_index]))
                                if now >= start_3 + start_plus and now < end_3 + end_plus:
                                    reg_val = wyt_data["Val3"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 12.5
                                    val_found = 1
                                end_3 = end_3.replace(year=2017)

                            if start_4 is not None and val_found == 0:
                                start_4 = start_4.replace(year=int(years[year_index]))
                                end_4 = end_4.replace(year=int(years[year_index]))
                                if now >= start_4 + start_plus and now < end_4 + end_plus:
                                    reg_val = wyt_data["Val4"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 19.0
                                    val_found = 1
                                end_4 = end_4.replace(year=2017)

                            if start_5 is not None and val_found == 0:
                                start_5 = start_5.replace(year=int(years[year_index]))
                                end_5 = end_5.replace(year=int(years[year_index]))
                                if now >= start_5 + start_plus and now < end_5 + end_plus:
                                    reg_val = wyt_data["Val5"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 16.5
                                end_5 = end_5.replace(year=2017)

                            if start_6 is not None and val_found == 0:
                                start_6 = start_6.replace(year=int(years[year_index]))
                                end_6 = end_6.replace(year=int(years[year_index]))
                                if now >= start_6 + start_plus and now < end_6 + end_plus:
                                    reg_val = wyt_data["Val6"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 15.6
                                    val_found = 1
                                end_6 = end_6.replace(year=2017)

                            std_ts.append(reg_val)
                        std_ts = std_ts[1:]
                        std_ts.append(std_ts[-1])

                        if count < 1:
                            std_ts_df = pd.DataFrame({std_nm: std_ts})
                            std_nms = [str(std_nm)]
                            count = 1
                        else:
                            std_ts_df[std_nm] = std_ts
                            std_nms.append(str(std_nm))

            ThirtyDay_Stns = ["RSAN112", "RSAN072", "OLDR_MIDR", "ROLD059"]
            Monthly_Stns = ["RSAC081", "SLMZU025", "SLMZU011", "SLCBN002", "SLSUS012", "CHWST000", "CHDMC004"]
            new_EC_vals = []

            if out_statns[station_index] in ThirtyDay_Stns:
                for k in range(len(values)):
                    if k < 29:
                        new_val = np.mean(values[:k + 1])
                    else:
                        sub_vals = []
                        for k2 in range(0, 30):
                            val = values[k - k2]
                            if val > -1000:
                                sub_vals.append(val)

                        new_val = np.mean(sub_vals)
                    new_val = new_val / 1000
                    new_EC_vals.append(new_val)
            elif out_statns[station_index] in Monthly_Stns:

                daily_df = pd.DataFrame({'Date': dates, 'Value': values})
                daily_df['Month'] = [date.month for date in dates]
                daily_df['Year'] = [date.year for date in dates]

                month_df = daily_df.groupby(['Year', 'Month'], as_index=False).mean()

                month_df['Date'] = [datetime(1921, 3, 1) + pd.DateOffset(months=i) + timedelta(days=-1) for i in range(len(month_df))]
                for k in range(len(values)):
                    m = daily_df['Month'][k]
                    y = daily_df['Year'][k]
                    sub_df = month_df[month_df['Month'] == m]
                    sub_df = sub_df[sub_df['Year'] == y]
                    new_val = sub_df['Value'].iloc[0] / 1000
                    new_EC_vals.append(new_val)
            else:  # 14 days
                for k in range(len(values)):
                    if k < 13:
                        new_val = np.mean(values[:k + 1])
                    else:
                        sub_vals = []
                        for k2 in range(0, 14):
                            val = values[k - k2]
                            if val > -1000:
                                sub_vals.append(val)

                        new_val = np.mean(sub_vals)
                    new_val = new_val / 1000
                    new_EC_vals.append(new_val)

            if cl_flag > 0:
                new_Cl_vals = []
                for val in values:
                    if val < -1000:
                        new_val = None
                    elif val <= 280:
                        new_val = val * 0.15 - 12
                    else:
                        new_val = val * 0.285 - 50
                    new_Cl_vals.append(new_val)
            else:
                new_Cl_vals = [None] * len(values)

            nd_flags.append(nd_flag)

            df = pd.DataFrame({
                "Var Name": stations,
                "Location": locations,
                "Var type": cpart,
                "Date": dates,
                "ValueEC": new_EC_vals,
                "ValueCl": new_Cl_vals,
                "Study Scenario": study_scenario,
                "Study Type": study_type,
                "UnitsEC": unitsEC,
                "UnitsCl": unitsCl
            })
            df = pd.concat([df, std_ts_df], axis=1)

            # Add water year type column
            df["WYT"] = wyts

            # Filter to start in WY 1921 (i.e., Oct 1921 or later)
            df = df[(df["Date"].dt.year > 1921) | ((df["Date"].dt.year == 1921) & (df["Date"].dt.month > 9))]

            # Remove rows with invalid EC values
            df = df[df["ValueEC"] > -1000000]

            # Remove rows where all values are NA
            df = df.dropna(how='all')

            # Rename columns to match final output
            df.columns = [
                "Var Name", "Location", "Var type", "Date", "ValueEC", "ValueCl",
                "Study Scenario", "Study Type", "UnitsEC", "UnitsCl",
                "D1641AG", "D1641FWS", "D1641MI", "D1641MIDNumDays", "D1641MIDThreshold",
                "MIAntiochNumDays", "MIAntiochThreshold", "MIOther", "SAC INDEX"
            ]

            # Write to file (no header, no row names, comma-separated)
            df.to_csv(tsfilename, index=False, header=False, mode='a')  # quoting=3 disables quoting (like quote = FALSE in R)

    print(f"Done with creating {os.path.basename(tsfilename)}")

    # Read the output ts file
    df_pre = pd.read_csv(tsfilename, skiprows=1, header=None)
    df_pre_names = pd.read_csv(tsfilename, nrows=1, header=None)

    # Assign column names from the first row
    df_pre.columns = df_pre_names.iloc[0]

    # Calculate differences
    df_pre['AG_diff'] = df_pre['ValueEC'] - df_pre['D1641AG']
    df_pre['FWS_diff'] = df_pre['ValueEC'] - df_pre['D1641FWS']
    df_pre['MI_diff'] = df_pre['ValueCl'] - df_pre['D1641MI']

    ndblw_arr = []
    ndvio_arr = []
    ndblw_arrA = []
    ndvio_arrA = []

    for ind in range(len(out_statns)):
        stn = out_statns[ind]
        v_type = f"EC-{out_stats[ind]}"

        df_stn2 = df_pre[df_pre['Var Name'] == stn]
        df_stn = df_stn2[df_stn2['Var type'] == v_type]

        if df_stn.empty:
            continue

        MIDT = df_stn['D1641MIDThreshold'].values
        AntiochT = df_stn['MIAntiochThreshold'].values

        if not (pd.isna(MIDT[0]) or not pd.isna(AntiochT[0])):
            dt = pd.to_datetime(df_stn['Date'])
            years = [date.year for date in dt]
            months = [date.month for date in dt]
            Clval = df_stn['ValueCl'].values

            if out_stats[ind] == 'MEAN':
                MIDND = df_stn['D1641MIDNumDays'].values
                itr = 0
                yrtot = 0
                for k in range(len(df_stn)):
                    if k == 0:
                        yr_prev = years[k]
                    yr = years[k]

                    if yr != yr_prev:
                        yrtot += itr
                        ndblw_arr[-1] = yrtot
                        if np.isnan(yrtot) or np.isnan(MIDND[k - 1]):
                            ndvio_arr[-1] = 0
                        elif yrtot < MIDND[k - 1]:
                            ndvio_arr[-1] = 1
                        else:
                            ndvio_arr[-1] = 0
                        ndblw_arr.append(np.nan)
                        ndvio_arr.append(0)
                        yrtot = 0
                        itr = 0
                        yr_prev = yr
                    elif k == len(df_stn) - 1:
                        yrtot += itr
                        ndblw_arr.append(yrtot)
                        if np.isnan(yrtot) or np.isnan(MIDND[k]):
                            ndvio_arr.append(0)
                        elif yrtot < MIDND[k]:
                            ndvio_arr.append(1)
                        else:
                            ndvio_arr.append(0)
                    else:
                        ndblw_arr.append(np.nan)
                        ndvio_arr.append(0)

                    if np.isnan(MIDT[k]) or np.isnan(Clval[k]):
                        continue
                    elif Clval[k] < MIDT[k]:
                        itr += 1
                    elif itr >= 14:
                        yrtot += itr
                        itr = 0
                    else:
                        itr = 0

                ndblw_add = [np.nan] * len(df_stn)
                ndblw_arrA.extend(ndblw_add)
                ndvio_arrA.extend(ndblw_add)

            else:
                AntiochND = df_stn['MIAntiochNumDays'].values
                itrA = 0
                for k in range(len(df_stn)):
                    if k == 0:
                        yr_prev = years[k]
                    yr = years[k]

                    if yr != yr_prev:
                        ndblw_arrA[-1] = itrA
                        if np.isnan(itrA) or np.isnan(AntiochND[k - 1]):
                            ndvio_arrA[-1] = 0
                        elif itrA < AntiochND[k - 1]:
                            ndvio_arrA[-1] = 1
                        else:
                            ndvio_arrA[-1] = 0
                        ndblw_arrA.append(np.nan)
                        ndvio_arrA.append(0)
                        itrA = 0
                        yr_prev = yr
                    elif k == len(df_stn) - 1:
                        ndblw_arrA.append(itrA)
                        if np.isnan(itrA) or np.isnan(AntiochND[k]):
                            ndvio_arrA.append(0)
                        elif itrA < AntiochND[k]:
                            ndvio_arrA.append(1)
                        else:
                            ndvio_arrA.append(0)
                    else:
                        ndblw_arrA.append(np.nan)
                        ndvio_arrA.append(0)

                    if np.isnan(AntiochT[k]) or np.isnan(Clval[k]):
                        continue
                    elif Clval[k] < AntiochT[k]:
                        itrA += 1

                ndblw_arr.extend([np.nan] * len(df_stn))
                ndvio_arr.extend([np.nan] * len(df_stn))

        else:
            ndblw_add = [np.nan] * len(df_stn)
            ndblw_arr.extend(ndblw_add)
            ndvio_arr.extend(ndblw_add)
            ndblw_arrA.extend(ndblw_add)
            ndvio_arrA.extend(ndblw_add)

    # Assign new columns to df_pre
    df_pre['NumDaysBlw'] = ndblw_arr
    df_pre['NumDaysViolated'] = ndvio_arr
    df_pre['AntiochNumDaysBlw'] = ndblw_arrA
    df_pre['AntiochNumDaysViolated'] = ndvio_arrA
    # Construct output filename
    tsfilenamediff = f"DSM2ComplianceDiffData_{input_model_name}.csv"
    tsfilenamediff = os.path.join(outfile_location, tsfilenamediff)

    # Define custom header
    custom_header = [
        "Var Name", "Location", "Var type", "Date", "ValueEC", "ValueCl", "Study Scenario", "Study Type",
        "UnitsEC", "UnitsCl", "D1641AG", "D1641FWS", "D1641MI", "D1641MIDNumDays", "D1641MIDThreshold",
        "MIAntiochNumDays", "MIAntiochThreshold", "MIOther", "SAC INDEX", "DiffAG", "DiffFWS", "DiffMI",
        "NumDaysBlw", "NumDaysViolated", "AntiochNumDaysBlw", "AntiochNumDaysViolated"
    ]

    # Write to CSV with custom header
    with open(tsfilenamediff, 'w') as f:
        f.write(','.join(custom_header) + '\n')
        df_pre.to_csv(f, index=False, header=False, na_rep='NA')


def percentile(dl_values, percent, key=lambda x: x):
    """
    Find the percentile of a list of values.

    Parameters
    ----------
    dl_values: list
        list of values to get the value from. Note this MUST BE already sorted.
    percent: float
        a float value from 0.0 to 1.0. The percentile we want to find
    key: function
        optional key function to compute value from each element of dl_values

    Returns
    -------
    the percentile of the values
    """

    if not dl_values:
        return None
    k = (len(dl_values) - 1) * percent
    f = floor(k)
    c = ceil(k)
    if f == c:
        return key(dl_values[int(k)])
    d0 = key(dl_values[int(f)]) * (c - k)
    d1 = key(dl_values[int(c)]) * (k - f)
    return d0 + d1


def combine_percentiles(scen_nm):
    """
    Calculates the percentiles for all of the compliance locations. For each compliance point ranks the values from the get_dsm2_timeseries_data output.
    Write out a _Percentiles.csv file for each location. Also writes a _Dates.csv for each and a compliance summary for the scenario
    Parameters
    ----------
    scen_nm: str
        name of the scenario

    Returns
    -------
    None
    """

    infn = "./water_qual_csvs/DSM2ComplianceDiffData_" + scen_nm + ".csv"
    outfn = "./water_qual_csvs/_ComplianceSummary/DSM2ComplianceSummary_" + scen_nm + ".csv"
    out_dir = "./water_qual_csvs/_Percentiles/"

    if not os.path.exists("./water_qual_csvs/_ComplianceSummary/"):
        os.mkdir("./water_qual_csvs/_ComplianceSummary/")

    if not os.path.exists(out_dir):
        os.mkdir(out_dir)

    with open(infn, "r") as inf:
        in_data = inf.readlines()

    outf = open(outfn, "w")

    print(in_data[0].rstrip().split(","))

    outf.write("VarName,DiffAG,DiffFWS,DiffMI,NumDaysBlw,NumDaysViolated,AntiochNumDaysBlw,AntiochNumDaysViolated,TotNumDays_D1641AG,TotNumDays_D1641FWS,TotNumDays_D1641MI\n")
    date_ind = in_data[0].rstrip().split(",").index("Date")
    wyt_ind = in_data[0].rstrip().split(",").index("SAC INDEX")
    diff_ag_ind = in_data[0].rstrip().split(",").index("DiffAG")
    diff_fws_ind = in_data[0].rstrip().split(",").index("DiffFWS")
    diff_mi_ind = in_data[0].rstrip().split(",").index("DiffMI")
    mi_numdaysblw_ind = in_data[0].rstrip().split(",").index("NumDaysBlw")
    mi_numdaysstandard_ind = in_data[0].rstrip().split(",").index("D1641MIDNumDays")
    mi_numdaysvio_ind = in_data[0].rstrip().split(",").index("NumDaysViolated")
    antioch_numdaysblw_ind = in_data[0].rstrip().split(",").index("AntiochNumDaysBlw")
    antioch_numdaysvio_ind = in_data[0].rstrip().split(",").index("AntiochNumDaysViolated")

    for i, x in enumerate(in_data):
        if in_data[i] == '\n': continue
        if i == 0: continue
        if i == 1:
            diff_ag = 0
            diff_ag_arr = []
            diff_ag_date_arr = []
            diff_ag_wyt_arr = []
            diff_fws = 0
            diff_fws_arr = []
            diff_fws_date_arr = []
            diff_fws_wyt_arr = []
            diff_mi = 0
            diff_mi_arr = []
            diff_mi_date_arr = []
            diff_mi_wyt_arr = []
            mi_numdaysblw_arr = []
            mi_numdaysblw = 0
            mi_numdaysvio = 0
            antioch_numdaysblw = 0
            antioch_numdaysvio = 0
            VarNamePrev = in_data[i].split(",")[0]
            VarName = in_data[i].split(",")[0]
            TotNumDays_D1641AG = 0
            TotNumDays_D1641FWS = 0
            TotNumDays_D1641MI = 0
        if i == (len(in_data) - 1):
            if in_data[i].rstrip().split(",")[diff_ag_ind] != "NA":
                TotNumDays_D1641AG += 1
                diff_ag_arr.append(float(in_data[i].rstrip().split(",")[diff_ag_ind]))
                if float(in_data[i].rstrip().split(",")[diff_ag_ind]) > 0:
                    diff_ag_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_ag_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_ag += 1
            if in_data[i].rstrip().split(",")[diff_fws_ind] != "NA":
                TotNumDays_D1641FWS += 1
                diff_fws_arr.append(float(in_data[i].rstrip().split(",")[diff_fws_ind]))
                if float(in_data[i].rstrip().split(",")[diff_fws_ind]) > 0:
                    diff_fws_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_fws_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_fws += 1
            if in_data[i].rstrip().split(",")[diff_mi_ind] != "NA":
                TotNumDays_D1641MI += 1
                diff_mi_arr.append(float(in_data[i].rstrip().split(",")[diff_mi_ind]))
                if float(in_data[i].rstrip().split(",")[diff_mi_ind]) > 0:
                    diff_mi_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_mi_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_mi += 1
            if in_data[i].rstrip().split(",")[mi_numdaysblw_ind] != "NA":
                # if float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) > 0:
                mi_numdaysblw += 1
                mi_numdaysblw_arr.append(float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) - float(in_data[i].rstrip().split(",")[mi_numdaysstandard_ind]))
            if in_data[i].rstrip().split(",")[mi_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[mi_numdaysvio_ind]) > 0: mi_numdaysvio += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysblw_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysblw_ind]) > 0: antioch_numdaysblw += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysvio_ind]) > 0: antioch_numdaysvio += 1
            print(VarNamePrev, diff_ag, diff_fws, diff_mi, mi_numdaysblw, mi_numdaysvio, antioch_numdaysblw, antioch_numdaysvio, TotNumDays_D1641AG, TotNumDays_D1641FWS, TotNumDays_D1641MI)
            outf.write(VarNamePrev + "," + str(diff_ag) + "," + str(diff_fws) + "," + str(diff_mi) + "," + str(mi_numdaysblw) + "," + str(mi_numdaysvio) + "," + str(antioch_numdaysblw) + "," + str(
                antioch_numdaysvio) + "," + str(TotNumDays_D1641AG) + "," + str(TotNumDays_D1641FWS) + "," + str(TotNumDays_D1641MI) + "\n")
            if len(diff_ag_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Dates.csv", "w") as date_file:
                    for j in range(len(diff_ag_date_arr)): date_file.write(str(diff_ag_date_arr[j]) + "," + str(diff_ag_wyt_arr[j]) + "\n")
                diff_ag_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_ag_arr, float(j + 1) / 100.0)) + "\n")
            if len(diff_fws_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Dates.csv", "w") as date_file:
                    for j in range(len(diff_fws_date_arr)): date_file.write(str(diff_fws_date_arr[j]) + "," + str(diff_fws_wyt_arr[j]) + "\n")
                diff_fws_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_fws_arr, float(j + 1) / 100.0)) + "\n")
            if len(diff_mi_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Dates.csv", "w") as date_file:
                    for j in range(len(diff_mi_date_arr)): date_file.write(str(diff_mi_date_arr[j]) + "," + str(diff_mi_wyt_arr[j]) + "\n")
                diff_mi_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_mi_arr, float(j + 1) / 100.0)) + "\n")
            if len(mi_numdaysblw_arr) > 0:
                mi_numdaysblw_arr = [float(i) for i in mi_numdaysblw_arr]
                mi_numdaysblw_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Days_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(mi_numdaysblw_arr, float(j + 1) / 100.0)) + "\n")

        else:
            VarName = in_data[i].split(",")[0]
            if VarName != VarNamePrev:
                print(VarName)
                print(VarNamePrev, diff_ag, diff_fws, diff_mi, mi_numdaysblw, mi_numdaysvio, antioch_numdaysblw, antioch_numdaysvio, TotNumDays_D1641AG, TotNumDays_D1641FWS, TotNumDays_D1641MI)
                outf.write(
                    VarNamePrev + "," + str(diff_ag) + "," + str(diff_fws) + "," + str(diff_mi) + "," + str(mi_numdaysblw) + "," + str(mi_numdaysvio) + "," + str(antioch_numdaysblw) + "," + str(
                        antioch_numdaysvio) + "," + str(TotNumDays_D1641AG) + "," + str(TotNumDays_D1641FWS) + "," + str(TotNumDays_D1641MI) + "\n")
                print("agg_arr")
                if len(diff_ag_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Dates.csv", "w") as date_file:
                        for j in range(len(diff_ag_date_arr)): date_file.write(str(diff_ag_date_arr[j]) + "," + str(diff_ag_wyt_arr[j]) + "\n")
                    diff_ag_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_ag_arr, float(j + 1) / 100.0)) + "\n")
                print("fws_arr")
                if len(diff_fws_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Dates.csv", "w") as date_file:
                        for j in range(len(diff_fws_date_arr)): date_file.write(str(diff_fws_date_arr[j]) + "," + str(diff_fws_wyt_arr[j]) + "\n")
                    diff_fws_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_fws_arr, float(j + 1) / 100.0)) + "\n")
                print("mi_arr")
                if len(diff_mi_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Dates.csv", "w") as date_file:
                        for j in range(len(diff_mi_date_arr)): date_file.write(str(diff_mi_date_arr[j]) + "," + str(diff_mi_wyt_arr[j]) + "\n")
                    diff_mi_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_mi_arr, float(j + 1) / 100.0)) + "\n")
                print("mid_arr")
                print(mi_numdaysblw_arr)
                if len(mi_numdaysblw_arr) > 0:
                    mi_numdaysblw_arr = [float(i) for i in mi_numdaysblw_arr]
                    print(mi_numdaysblw_arr)
                    mi_numdaysblw_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Days_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(mi_numdaysblw_arr, float(j + 1) / 100.0)) + "\n")

                diff_ag = 0
                diff_ag_arr = []
                diff_ag_date_arr = []
                diff_ag_wyt_arr = []
                diff_fws = 0
                diff_fws_arr = []
                diff_fws_date_arr = []
                diff_fws_wyt_arr = []
                diff_mi = 0
                diff_mi_arr = []
                diff_mi_date_arr = []
                diff_mi_wyt_arr = []
                mi_numdaysblw_arr = []
                mi_numdaysblw = 0
                mi_numdaysvio = 0
                antioch_numdaysblw = 0
                antioch_numdaysvio = 0
                TotNumDays_D1641AG = 0
                TotNumDays_D1641FWS = 0
                TotNumDays_D1641MI = 0
            if in_data[i].rstrip().split(",")[diff_ag_ind] != "NA":
                TotNumDays_D1641AG += 1
                diff_ag_arr.append(float(in_data[i].rstrip().split(",")[diff_ag_ind]))
                if float(in_data[i].rstrip().split(",")[diff_ag_ind]) > 0:
                    diff_ag_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_ag_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_ag += 1
            if in_data[i].rstrip().split(",")[diff_fws_ind] != "NA":
                TotNumDays_D1641FWS += 1
                diff_fws_arr.append(float(in_data[i].rstrip().split(",")[diff_fws_ind]))
                if float(in_data[i].rstrip().split(",")[diff_fws_ind]) > 0:
                    diff_fws_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_fws_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_fws += 1
            if in_data[i].rstrip().split(",")[diff_mi_ind] != "NA":
                TotNumDays_D1641MI += 1
                diff_mi_arr.append(float(in_data[i].rstrip().split(",")[diff_mi_ind]))
                if float(in_data[i].rstrip().split(",")[diff_mi_ind]) > 0:
                    diff_mi_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_mi_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_mi += 1
            if in_data[i].rstrip().split(",")[mi_numdaysblw_ind] != "NA":
                # if float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) > 0:
                mi_numdaysblw += 1
                mi_numdaysblw_arr.append(float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) - float(in_data[i].rstrip().split(",")[mi_numdaysstandard_ind]))
            if in_data[i].rstrip().split(",")[mi_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[mi_numdaysvio_ind]) > 0: mi_numdaysvio += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysblw_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysblw_ind]) > 0: antioch_numdaysblw += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysvio_ind]) > 0: antioch_numdaysvio += 1
            VarNamePrev = VarName

    outf.close()


def combine_all_runs(studies, percentile_files):
    """
    Combines the outputs from the combine_percentiles function's _Percentile cvs into one csv file that is used for the plots.
    Parameters
    ----------
    studies: list
        The different studies/alternatives to combine
    percentile_files:
        All of the csv files ending in _Percentile. These have the percentile data

    Returns
    -------
    final_data_frame: dataframe
        The combined dataframe
    """

    data = []
    flag = 0
    header = []

    for study in studies:
        for file in percentile_files:
            if study in file:
                with open("./water_qual_csvs/_Percentiles/" + file, "r") as inf:
                    in_data = inf.readlines()
                if "AG_" in file:
                    comp = "AG"
                elif "FWS_" in file:
                    comp = "FWS"
                elif "MI_Percentile" in file:
                    comp = "MI"
                elif "Days" in file:
                    comp = "MI_Days"
                if flag == 0:
                    header.append("Percentile")
                    header.append(study + "_" + file.split("_")[0] + "_" + comp)
                    col = []
                    for line in in_data:
                        col.append(line.rstrip().split(",")[0])
                    data.append(col)
                    col = []
                    for line in in_data:
                        col.append(line.rstrip().split(",")[1])
                    data.append(col)
                    flag = 1
                else:
                    if file.split("_")[0] == "OLDR" and file.split("_")[1] == "MIDR":
                        header.append(study + "_" + file.split("_")[0] + "_" + file.split("_")[1] + "_" + comp)
                        col = []
                        for line in in_data:
                            col.append(line.rstrip().split(",")[1])
                        data.append(col)
                    else:
                        header.append(study + "_" + file.split("_")[0] + "_" + comp)
                        col = []
                        for line in in_data:
                            col.append(line.rstrip().split(",")[1])
                        data.append(col)

    final_data_frame = pd.DataFrame(data).transpose()
    final_data_frame.columns = header
    final_data_frame = final_data_frame.apply(pd.to_numeric)

    return final_data_frame


def create_water_qual_plot(df_percentiles, fig_value, plot_directory, alts, line_styles, line_colors, d_ymin, d_ymax):
    """
    Creates the plot of the probability of compliance. Write the plot to a file but returns the path.
    Parameters
    ----------
    df_percentiles: dataframe
        Data frame of all of the data
    fig_value: str
        Value to plot
    plot_directory: str
        path to the folder to hold the plots
    alts: dict
        Dictionary of alternatives and names to show on plot
    line_styles: list
        List of line styles to show on plot
    line_colors: list
        List of line colors to show on plot
    d_ymin: float
        y axis minimum
    d_ymax: float
        y axis maximum

    Returns
    -------
    Path to the plot
    """
    # Check for/create directory to store monthly exceedance plots
    if not os.path.exists(plot_directory):
        os.makedirs(plot_directory)

    # define size and borders
    fig, axs = plt.subplots(figsize=(9, 5.5), linewidth=3, edgecolor="black")

    for fig_index, display_name in enumerate(alts):

        model_name = alts[display_name].split('.')[0]
        # Dataset for this alt
        df_alt_data = df_percentiles[['Percentile', model_name + '_' + fig_value]]

        # plot exceedance probability vs monthly EC
        percentages = range(0, 101, 10)
        percentage_labels = [f"{int(i)}%" for i in percentages]

        axs.plot(df_alt_data['Percentile'].values, df_alt_data[model_name + '_' + fig_value].values, color=line_colors[fig_index],
                 linestyle=line_styles[fig_index], label=display_name)
        axs.set_xticks(percentages)
        axs.set_xticklabels(percentage_labels)

        # set the Y axis depending on if its chloride or EC
        if fig_value.split('_')[-1] == 'MI':
            axs.set_ylabel("Difference in Chloride (Scenario minus Standard) (mg/L)")
        else:
            axs.set_ylabel("Difference in EC (Scenario minus Standard) (mmhos/cm)")
        axs.set_xlabel("Probability of Compliance (%)")

        # set the y limits
        axs.set_ylim(d_ymin, d_ymax)

        # Save this parameter to orient the legend correctly
        axbox = axs.get_position()

        # Add gridlines
        plt.grid(color='gray', linestyle='--', linewidth=0.8)

        # Add a legend
        plt.legend(loc='center', ncol=4, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08])

    # flip x-axis
    axs.invert_xaxis()

    plt.savefig(plot_directory + "/" + fig_value + "_exceedance" + ".png")

    plt.close()

    return plot_directory + "/" + fig_value + "_exceedance" + ".png"


def get_wq_location_data():
    """
    Gets the table of location data for the plots
    Returns
    -------
    crosswalk: dataframe
        Dataframe containing all locations, names, and y axis limits
    """

    crosswalk = pd.read_excel("../inputs/location_code_crosswalk_water_quality.xlsx")
    crosswalk.drop(columns="Model", inplace=True)
    return crosswalk


def create_compliance_appendix(scenario_names, template, doc_name, new_doc):
    """
    Creates the water quality compliance appendix. Creates the plots and puts them in a doccument
    Parameters
    ----------
    scenario_names: dict
        Dictionary of scenario names and files
    template: str
        Path to the template doc
    doc_name: str
        Name for temporary document
    new_doc: str
        Name for the final document

    Returns
    -------
    None
    """

    # Set working directory to the script's location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)

    # List all dss files in studies folder
    studies_path = os.path.abspath(os.path.join(script_dir, "../studies"))
    path_names = [d for d in os.listdir(studies_path)
                  if os.path.isfile(os.path.join(studies_path, d)) and d.endswith(".dss") and d in scenario_names.values()]

    print("Model directories found:")
    print(path_names)

    # Loop through each model directory and call the processing function
    for model_path in path_names:
        print(f"Processing model: {model_path}")
        get_dsm2_timeseries_data(model_path)

    # get the study names
    studies = [study.split(".")[0] for study in path_names]

    # loop through and call the combine percentiles function
    for study_name in studies:
        combine_percentiles(study_name)

    # get the percentile files that were created
    percentile_files = []

    for file in os.listdir("./water_qual_csvs/_Percentiles/"):
        if "Percentile" in file: percentile_files.append(file)

    # call the function to combine them
    final_data_frame = combine_all_runs(studies, percentile_files)

    # get the data for the different plots
    df_location_info = get_wq_location_data()

    # this will hold the alt text for each figure
    alt_text = []

    # set up the document
    appendix_prefix = " F.2.8"
    doc = docx.Document(template)
    doc.add_heading(f"Attachment{appendix_prefix}", level=1)

    # Add caption style for Figure captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.name = 'Times New Roman'

    # add the set up for the plots. A folder to hold them and line colors and styles
    s_plot_directory = "./wq_plots"
    line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
    line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

    # loop though the rows which coorespond to a plot
    for index, location in df_location_info.iterrows():

        # create the plot and capture the path
        s_plot_path = create_water_qual_plot(final_data_frame, location['VarName'], s_plot_directory,
                                             scenario_names, line_styles, line_colors,
                                             location['Ymin'], location['Ymax'])

        # Generate the caption to match the previous doccumentation
        s_fig_caption = 'D1641 ' + location['VarName'].split('_')[-1] + ' ' + location['Location (Title)'] + ' Compliance Exceedance Plot'

        # change to landscape to fit the images
        if index == 0:
            change_orientation(doc, "landscape")

        # Add figure as a picture
        o_fig = doc.add_picture(s_plot_path)

        # Generate fig title
        fig_title_prefix = "Figure " + appendix_prefix + "-"

        # Add title below figure
        add_caption_water_supply(doc, "Figure", fig_title_prefix, s_fig_caption, custom_style="Figure Caption")

        # Add to alt text
        alt_text.append(s_fig_caption)

        if index != len(df_location_info) - 1:
            doc.add_page_break()

    # save the doc to the temporary name
    doc.save(doc_name)

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    # Run vbs script
    # Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    # This will fail if Microsoft Word has document open in the background
    # try opening Task Manager and Ending MS Word Background Task, then rerun

    # Call the vbs script for figure alt text
    result = subprocess.call(
        "cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + "x" + " " + str(0) + " " + alt_text_string_figures + " " + str(len(df_location_info)))

    # check if it worked successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")
    else:
        # Instructions on how to finish formatting numbered captions.
        print("After running this script, \n1. Open Word file and Ctrl+A to select all. Then F9 to update caption numbering.")
