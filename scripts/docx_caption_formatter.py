"""
Functions used to format Figure and Table captions to be Word captions that can auto update.

"""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def add_caption_byfield(document, tab_or_figure, prefix, caption, custom_style='Caption', use_prev_number=False):
    """
    Adds a figure or table caption with auto-updating numbering, inheriting from Heading 2 number.

    Caption format is: {"Table" or "Figure"} {Heading 2 number}-{fig or table number}. {caption}
    Example: "Table 2.2-3. Some Caption"

    Parameters
    ----------
    document: python docx document object
    tab_or_figure: str
        "Figure" or "Table" depending on what type of caption you want to make
    prefix: str
        The caption prefix. (Ex: F.2.2. if you want your figures to be numbered F.2.2.1, F.2.2.2, ...)
    caption: str
        Caption contents
    custom_style: str
        python docx style name. Default is "Caption" (the default Word Caption style)
        Other options are "Table Caption" and "Figure Caption"
    use_prev_number: int
        Use the previous table number if this is True. This is used for the tables with multiple letters after the
        table number. (Ex: Table F.1.2a, F.1.2b, F.1.2c).


    Returns
    -------
    None
    """
    target = tab_or_figure  # Caption type. Options: "Table" or "Figure"

    # Set up word formatting for caption.
    # Start caption with the input prefix.
    paragraph = document.add_paragraph(f'{prefix}', style=custom_style)
    # Set up caption numbering that inherits heading 2 numbering.
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' STYLEREF 2  \s'#These captions inherit their formatting from Style Reference Heading 2.
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

    paragraph.add_run('-') #Add dash between heading 2 number and the figure/table number.

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    r.append(instrText)

    #Add table number to caption.
    run= paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    if use_prev_number == False:  #If no figure or table number to hardcode is provided, use the default sequential numbering
        instrText.text = f'SEQ {target}\* ARABIC \s 2' #2 is for inheriting the heading 2 label.
    else:
        instrText.text = f'SEQ {target}\* ARABIC \c ' #\c denotes using the previous table or figure number.
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

    # Add caption text
    run = paragraph.add_run(f' {caption}')
    run.font.size = Pt(12)
    return


def add_caption_water_supply(document, tab_or_figure, prefix, caption, custom_style='Caption'):
    """
    Adds a figure or table caption with auto-updating numbering, inheriting from Heading 2 number. For water supply version.

    Caption format is: {"Table" or "Figure"} {Heading 2 number}-{fig or table number}. {caption}
    Example: "Table 2.2-3. Some Caption"

    Parameters
    ----------
    document: python docx document object
    tab_or_figure: str
        "Figure" or "Table" depending on what type of caption you want to make
    prefix: str
        The caption prefix. (Ex: F.2.2. if you want your figures to be numbered F.2.2.1, F.2.2.2, ...)
    caption: str
        Caption contents
    custom_style: str
        python docx style name. Default is "Caption" (the default Word Caption style)
        Other options are "Table Caption" and "Figure Caption"
    use_prev_number: int
        Use the previous table number if this is True. This is used for the tables with multiple letters after the
        table number. (Ex: Table F.1.2a, F.1.2b, F.1.2c).


    Returns
    -------
    None
    """
    target = tab_or_figure  # Caption type. Options: "Table" or "Figure"

    # Set up word formatting for caption.
    # Start caption with the input prefix.
    paragraph = document.add_paragraph(f'{prefix}', style=custom_style)
    # Set up caption numbering that inherits heading 2 numbering.
    run = paragraph.add_run()
    r = run._r

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    r.append(instrText)

    #Add table number to caption.
    run= paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = f'SEQ {target}\* ARABIC \s 2' #2 is for inheriting the heading 2 label.
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

    # Add caption text
    run = paragraph.add_run(f'. {caption}')
    run.font.size = Pt(12)
    return
