### This file is meant to recreate what get_dsm2_comp_data_20240221.r from Jacobs does
### to get the full appendix, more steps must be followed. See Jacobs Attachment_F2-08_Scripts/ReadMe.txt
### This for the water quality compliance attachment




import os
import re
import pandas as pd
import datetime
from pydsstools.heclib.dss import HecDss
import numpy as np
import math
import matplotlib.pyplot as plt
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx_caption_formatter import add_caption_water_supply
from EISAppendixGen_functions import *
import subprocess


def get_stations():
    """
    Gets the stations for the compliance locations. These are the DSS b parts.
    Returns
    -------
    stations: list
    List of stations
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 0].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_locations():
    """
    Gets the location names for the compliance locations
    Returns
    -------
    stations: list
    List of location names
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 1].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_stats():
    """
    Gets the statistics for the compliance locations. (MIN or MAX or MEAN etc)
    Returns
    -------
    stations: list
    List of stats
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    stations_df = None

    # Look for the DSM2ComplianceLocations.csv file
    for j in range(len(fnames)):
        if fnames[j] == "DSM2ComplianceLocations.csv":
            stations_df = pd.read_csv(fpaths[j], header=None, sep=",")
            break  # Exit loop once the file is found

    if stations_df is not None:
        stations = stations_df.iloc[:, 2].astype(str).tolist()
        return stations
    else:
        return []  # Return empty list if file not found


def get_sri_current_condition():
    """
    Gets the dataframe of sac river index data
    Returns
    -------
    df_sri: dataframe
    dataframe of sac river index data for each year
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    df_sri = None

    # Look for the SacRiverIndex.csv file
    for j in range(len(fnames)):
        if fnames[j] == "SacRiverIndex.csv":
            print("\nFound file:", fnames[j])
            df_sri = pd.read_csv(fpaths[j], header=0, sep=",")
            print("\n")
            break  # Exit loop once the file is found

    return df_sri


def get_wyts(s_wyt_path):
    """
    Gets the water year type for each year depending on the given file name
    Parameters
    ----------
    s_wyt_path: str
        Name of wyt file

    Returns
    -------
    wyts_df: dataframe
        dataframe of water year type for each year depending on the given file name
    """

    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    wyts_df = None

    # Look for the WYT_2022MED.csv file
    for j in range(len(fnames)):
        if fnames[j] == s_wyt_path:
            wyts_df = pd.read_csv(fpaths[j], header=0, sep=",")
            break  # Exit loop once the file is found

    return wyts_df


def get_specified_table(s_csv_name):
    """
    Reads in the specified csv file and returns a dataframe
    Parameters
    ----------
    s_csv_name: str
        Name of csv file

    Returns
    -------
    std_df: dataframe
        Dataframe of data from the specified csv file
    """
    # Define the path to the stations directory
    stations_dir = "../inputs/stations"

    # Get full paths and file names
    fpaths = [os.path.join(stations_dir, fname) for fname in os.listdir(stations_dir)]
    fnames = os.listdir(stations_dir)

    std_df = None

    # Look for the D1641_AG.csv file
    for j in range(len(fnames)):
        if fnames[j] == s_csv_name:
            std_df = pd.read_csv(fpaths[j], header=0, sep=",")
            break  # Exit loop once the file is found

    return std_df


def get_dsm2_timeseries_data(file_path):
    """
    Function to do the main work for the water quality compliance appendix.
    This will read in the DSS file and all the needed compliance info and create two csvs with the output data.

    Parameters
    ----------
    file_path: str
        Path to the DSS file

    Returns
    -------
    None
    """

    # Construct the path to the model's directory
    file_path = os.path.join("..", "studies", file_path)

    input_model_name = os.path.basename(file_path).split('.')[0]

    print("Model Name:", input_model_name)


    # get the WYT data based on the hydrology
    if "2030" in input_model_name:
        df_wyt = get_wyts("WYT_2030.csv")
    elif "2070" in input_model_name:
        df_wyt = get_wyts("WYT_2070.csv")
    elif "2022" in input_model_name:
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_2022MED.csv")
    elif "Current" in input_model_name:
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_CurrentConditions.csv")
    else:
        print('Hydrology is undefined in study name, using current')
        df_sri = get_sri_current_condition()
        df_wyt = get_wyts("WYT_CurrentConditions.csv")

    # Load csv files for compliance standards
    df_D1641AG = get_specified_table("D1641_AG.csv")
    df_D1641FWS = get_specified_table("D1641_FWS.csv")
    df_D1641MI = get_specified_table("D1641_MI.csv")
    df_D1641MID = get_specified_table("D1641_MID.csv", )
    df_MI_Antioch = get_specified_table("MI_Antioch.csv")
    df_MI_Other  = get_specified_table("MI_Other.csv")
    comp_files = [df_D1641AG,
                       df_D1641FWS,
                       df_D1641MI,
                       df_D1641MID,
                       df_MI_Antioch,
                       df_MI_Other]
    comp_names = ["D1641 AG",
                    "D1641 FWS",
                    "D1641 MI",
                    "D1641 MID",
                    "MI Antioch",
                    "MI Other"]

    with HecDss.Open(file_path) as input_file:
        # Retreive stations and locations
        out_statns = get_stations()
        out_locs = get_locations()
        out_stats = get_stats()


        outfile_location = "./water_qual_csvs"
        if not os.path.exists(outfile_location):
            os.makedirs(outfile_location)

        # Construct the output filename
        tsfilename = f"DSM2ComplianceData_{input_model_name}.csv"
        tsfilename = os.path.join(outfile_location, tsfilename)

        # Open the file for writing
        with open(tsfilename, 'w') as tsfile:
            # Write the header line
            header = "Var Name,Location,Var type,Date,ValueEC,ValueCl,Study Scenario,Study Type,UnitsEC,UnitsCl,D1641AG,D1641FWS,D1641MI,D1641MIDNumDays,D1641MIDThreshold,MIAntiochNumDays,MIAntiochThreshold,MIOther,SAC INDEX\n"

            tsfile.write(header)

        nd_flags = []

        # Loop through locations
        for station_index, station in enumerate(out_statns):
            print(station)

            cl_flag = 0
            nd_flag = 0
            def_stn_flag = 0
            sjr_fws_stn_flag = 0

            bpart = station

            # Set flags based on station name
            if bpart in ["SLCBN002", "SLSUS012"]:
                def_stn_flag = 1
            if bpart in ["RSAN018", "RSAN032", "RSAN037"]:
                sjr_fws_stn_flag = 1
            # Create search path
            cpart = f"EC-{out_stats[station_index]}"
            path = f"/*/{bpart}/{cpart}/*/1DAY/*/"
            path_list = input_file.getPathnameList(path, sort=1)

            if path_list == []:
                continue

            expanded_path = path_list[0].split('/')
            expanded_path[4] = ''
            path = "/".join(expanded_path)
            o_timeseries = input_file.read_ts(path)

            if o_timeseries.empty:
                continue

            dates = np.array(o_timeseries.pytimes)

            # mimikling the date shifting that the R code does
            # python reads the dates in one day later than R does so we subtract two days and then another for the years and months
            dates = dates + datetime.timedelta(days=-2)

            years = [(date + datetime.timedelta(days=-1)).strftime("%Y") for date in dates]
            months = [(date + datetime.timedelta(days=-1)).strftime("%m") for date in dates]
            days = [(date + datetime.timedelta(days=-1)).strftime("%d") for date in dates]

            values = np.round(o_timeseries.values, 5)

            unitsEC = ["mmhos/cm"] * len(values)
            unitsCl = ["mg/L"] * len(values)
            stations = [out_statns[station_index]] * len(values)
            locations = [out_locs[station_index]] * len(values)
            cpart = [cpart] * len(values)
            study_scenario = [input_model_name] * len(values)
            study_type = ["DSM2"] * len(values)

            # create this data frame to do some matching quicker than a loop
            df_wyt_sri = pd.DataFrame({'Year': years, 'Month':months})
            df_wyt_sri = df_wyt_sri.astype(int)

            # get the water year and the previous year
            df_wyt_sri['WY'] = np.where(df_wyt_sri['Month'] < 10, df_wyt_sri['Year'], df_wyt_sri['Year']+1)
            df_wyt_sri['Prev Year'] = df_wyt_sri['Year'] - 1

            # Merge with df_wyt on water year
            df_wyt_sri = df_wyt_sri.merge(df_wyt.rename(columns={df_wyt.columns[1]: 'wy_ind', df_wyt.columns[2]: 'wyt'}),
                          left_on='WY', right_on='YEAR', how='left')

            # Merge to get current SRI
            df_wyt_sri = df_wyt_sri.merge(df_sri[[df_sri.columns[0], df_sri.columns[3]]].rename(columns={df_sri.columns[3]: 'curr_sri'}),
                          left_on='Year', right_on='WaterYear', how='left')

            # Merge to get previous SRI
            df_wyt_sri = df_wyt_sri.merge(df_sri[[df_sri.columns[0], df_sri.columns[3]]].rename(columns={df_sri.columns[3]: 'prev_sri'}),
                                          left_on='Prev Year', right_on='WaterYear', how='left')

            # Extract final lists
            wyts = df_wyt_sri['wyt'].tolist()
            wy_inds = df_wyt_sri['wy_ind'].tolist()
            prev_sri = df_wyt_sri['prev_sri'].tolist()
            curr_sri = df_wyt_sri['curr_sri'].tolist()

            # Append the last wyt value to the list
            wyts.append(wyts[-1])
            wyts = wyts[1:]

            # Print unique values
            print("Unique WYT values:", list(set(wyts)))

            # Initialize list for compliance DataFrames
            loc_comp_dfs = [None] *  len(comp_names)

            # Initialize compliance count list with zeros
            comp_count = [0] * len(comp_names)

            for compliance_index in range(len(comp_files)):
                comp_df = comp_files[compliance_index]

                # Filter rows where Var Name matches the current station
                loc_comp_df = comp_df[comp_df["Var Name"] == out_statns[station_index]]

                if not loc_comp_df.empty:
                    if out_statns[station_index] == "RSAN007":
                        if out_stats[station_index] == "MAX" and comp_names[compliance_index] == "MI Antioch":
                            print(comp_names[compliance_index])
                            comp_count[compliance_index] = 1
                            loc_comp_dfs[compliance_index] = loc_comp_df
                        elif out_stats[station_index] == "MEAN" and comp_names[compliance_index] != "MI Antioch":
                            print(comp_names[compliance_index])
                            comp_count[compliance_index] = 1
                            loc_comp_dfs[compliance_index] = loc_comp_df
                    else:
                        print(comp_names[compliance_index])
                        comp_count[compliance_index] = 1
                        loc_comp_dfs[compliance_index] = loc_comp_df

            # Loop through time series and set standard
            dates_copy = dates.copy()
            std_nms = []
            # std_ts_df = pd.DataFrame()
            count = 0
            for compliance_index in range(len(comp_count)):
                print('comp_count[compliance_index]: ' + str(comp_count[compliance_index]))
                if comp_count[compliance_index] < 1:
                    print('less than 1')
                    std_ts = [np.nan] * len(years)
                    std_nm = comp_names[compliance_index]

                    if std_nm in ["D1641 MID", "MI Antioch"]:
                        std_ts2 = std_ts

                        if count < 1:
                            std_ts_df = pd.DataFrame({
                                f"{std_nm}NumDays": std_ts,
                                f"{std_nm}Threshold": std_ts2
                            })
                            std_nms.append(str(std_nm))
                            count = 1
                        else:
                            std_ts_df[f"{std_nm}NumDays"] = std_ts
                            std_ts_df[f"{std_nm}Threshold"] = std_ts2
                            std_nms.append(str(std_nm))


                    else:
                        if count < 1:
                            std_ts_df = pd.DataFrame({std_nm: std_ts})
                            std_nms.append(str(std_nm))
                            count = 1
                        else:
                            std_ts_df[std_nm] = std_ts
                            std_nms.append(str(std_nm))

                    print("std_ts_df")
                    print(std_ts_df)
                    print("\n")
                else:
                    print("\ncomp_count[compliance_index] >= 1")
                    df_comp = loc_comp_dfs[compliance_index]
                    print(df_comp)
                    std_nm = comp_names[compliance_index]

                    std_ts = []
                    std_ts2 = []

                    if not pd.isna(df_comp["NumDays"].iloc[0]):
                        print("NumDays")
                        cl_flag = 1
                        nd_flag = 1

                        for year_index in range(len(years)):
                            if year_index == 0:
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                NumDays = wyt_data["NumDays"].iloc[0]
                                Threshold = wyt_data["Val1"].iloc[0]
                                NumDays_next = NumDays
                                Threshold_next = Threshold

                                if int(months[year_index]) < 10:
                                    wy_prev = years[year_index]
                                    yr_prev = years[year_index]
                                else:
                                    wy_prev = str(int(years[year_index]) + 1)
                                    yr_prev = years[year_index]

                            yr = years[year_index]
                            wy = years[year_index] if int(months[year_index]) < 10 else str(int(years[year_index]) + 1)

                            if wy != wy_prev:
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                NumDays_next = wyt_data["NumDays"].iloc[0] if not wyt_data.empty else np.nan
                                Threshold_next = wyt_data["Val1"].iloc[0] if not wyt_data.empty else np.nan
                                wy_prev = wy

                            if yr != yr_prev:
                                NumDays = NumDays_next
                                Threshold = Threshold_next
                                yr_prev = yr

                            std_ts.append(NumDays)
                            std_ts2.append(Threshold)

                        # Remove first placeholder if needed (mimicking R's append + [-1])
                        std_ts = std_ts[1:] if len(std_ts) > len(years) else std_ts
                        std_ts2 = std_ts2[1:] if len(std_ts2) > len(years) else std_ts2

                        # Add to std_ts_df
                        if count < 1:
                            std_ts_df = pd.DataFrame({
                                f"{std_nm}NumDays": std_ts,
                                f"{std_nm}Threshold": std_ts2
                            })
                            std_nms.append(std_nm)
                            count = 1
                        else:
                            std_ts_df[f"{std_nm}NumDays"] = std_ts
                            std_ts_df[f"{std_nm}Threshold"] = std_ts2
                            std_nms.append(std_nm)

                    else:
                        if std_nm == "D1641 MI":
                            cl_flag = 1
                        for year_index in range(len(years)):
                            yr_count = 0
                            if year_index == 0:
                                # Reset deficiency flag
                                def_flag = 0

                                # Filter compliance data for the current water year index
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]

                                # Parse start and end dates for each compliance window

                                start_1 = datetime.datetime.strptime(wyt_data['Start Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 1'].iloc[0]) else None
                                end_1 = datetime.datetime.strptime(wyt_data['End Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 1'].iloc[0]) else None
                                start_2 = datetime.datetime.strptime(wyt_data['Start Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 2'].iloc[0]) else None
                                end_2 = datetime.datetime.strptime(wyt_data['End Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 2'].iloc[0]) else None
                                start_3 = datetime.datetime.strptime(wyt_data['Start Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 3'].iloc[0]) else None
                                end_3 = datetime.datetime.strptime(wyt_data['End Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 3'].iloc[0]) else None
                                start_4 = datetime.datetime.strptime(wyt_data['Start Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 4'].iloc[0]) else None
                                end_4 = datetime.datetime.strptime(wyt_data['End Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 4'].iloc[0]) else None
                                start_5 = datetime.datetime.strptime(wyt_data['Start Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 5'].iloc[0]) else None
                                end_5 = datetime.datetime.strptime(wyt_data['End Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 5'].iloc[0]) else None
                                start_6 = datetime.datetime.strptime(wyt_data['Start Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 6'].iloc[0]) else None
                                end_6 = datetime.datetime.strptime(wyt_data['End Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 6'].iloc[0]) else None

                                # Store previous year and water year index values
                                year_prev = int(years[year_index])
                                wy_ind_prev = int(wy_inds[year_index])
                                wy_ind_prev2 = int(wy_inds[year_index])
                            year = years[year_index]
                            if int(year) != year_prev:
                                print(f"\nYear: {years[year_index]}")
                                yr_count += 1
                                def_flag = 0

                                # Get new WYT data for this year
                                wyt_data = df_comp[df_comp["Sac Index Val"] == wy_inds[year_index]]
                                print(f"WY index: {wy_inds[year_index]}")
                                print(f"Prev Sac River index: {prev_sri[year_index]}\n")

                                # Deficiency logic
                                if wy_inds[year_index] == 5 and wy_ind_prev >= 4:
                                    def_flag = 1
                                elif wy_inds[year_index] == 4 and prev_sri[year_index] < 11.35:
                                    def_flag = 1
                                elif wy_inds[year_index] == 4 and wy_ind_prev >= 4 and wy_ind_prev2 == 5 and yr_count > 1:
                                    def_flag = 1

                                # SJR FWS logic
                                sjr_fws_flag = 1 if wy_inds[year_index] == 4 and curr_sri[year_index] < 8.1 else 0

                                # Set compliance windows
                                if def_flag == 1 and def_stn_flag == 1:
                                    print("Deficiency")
                                    start_1 = datetime.datetime.strptime(f"1-Jan", "%d-%b")
                                    end_1 = datetime.datetime.strptime(f"31-Mar", "%d-%b")
                                    start_2 = datetime.datetime.strptime(f"1-Apr", "%d-%b")
                                    end_2 = datetime.datetime.strptime(f"30-Apr", "%d-%b")
                                    start_3 = datetime.datetime.strptime(f"1-May", "%d-%b")
                                    end_3 = datetime.datetime.strptime(f"31-May", "%d-%b")
                                    start_4 = datetime.datetime.strptime(f"1-Oct", "%d-%b")
                                    end_4 = datetime.datetime.strptime(f"31-Oct", "%d-%b")
                                    start_5 = datetime.datetime.strptime(f"1-Nov", "%d-%b")
                                    end_5 = datetime.datetime.strptime(f"30-Nov", "%d-%b")
                                    start_6 = datetime.datetime.strptime(f"1-Dec", "%d-%b")
                                    end_6 = datetime.datetime.strptime(f"31-Dec", "%d-%b")
                                else:

                                    start_1 = datetime.datetime.strptime(wyt_data['Start Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 1'].iloc[0]) else None
                                    end_1 = datetime.datetime.strptime(wyt_data['End Date 1'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 1'].iloc[0]) else None
                                    start_2 = datetime.datetime.strptime(wyt_data['Start Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 2'].iloc[0]) else None
                                    end_2 = datetime.datetime.strptime(wyt_data['End Date 2'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 2'].iloc[0]) else None
                                    start_3 = datetime.datetime.strptime(wyt_data['Start Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 3'].iloc[0]) else None
                                    end_3 = datetime.datetime.strptime(wyt_data['End Date 3'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 3'].iloc[0]) else None
                                    start_4 = datetime.datetime.strptime(wyt_data['Start Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 4'].iloc[0]) else None
                                    end_4 = datetime.datetime.strptime(wyt_data['End Date 4'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 4'].iloc[0]) else None
                                    start_5 = datetime.datetime.strptime(wyt_data['Start Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 5'].iloc[0]) else None
                                    end_5 = datetime.datetime.strptime(wyt_data['End Date 5'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 5'].iloc[0]) else None
                                    start_6 = datetime.datetime.strptime(wyt_data['Start Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['Start Date 6'].iloc[0]) else None
                                    end_6 = datetime.datetime.strptime(wyt_data['End Date 6'].iloc[0], "%d-%b") if pd.notna(wyt_data['End Date 6'].iloc[0]) else None

                                # Adjust end_1 if SJR FWS condition is met
                                if sjr_fws_flag == 1 and sjr_fws_stn_flag == 1:
                                    end_1 = datetime.datetime.strptime(f"30-Apr", "%d-%b")

                                # Update previous year/index trackers
                                year_prev = int(years[year_index])
                                wy_ind_prev2 = wy_ind_prev
                                wy_ind_prev = int(wy_inds[year_index])

                            # Add 8 hours to start
                            start_plus = datetime.timedelta(hours=8)

                            # Add one day for days marked 12-31
                            end_plus = datetime.timedelta(hours=32)

                            # Current date being evaluated
                            now = dates_copy[year_index]

                            val_found = 0
                            reg_val = None

                            if start_1 is not None:
                                start_1 = start_1.replace(year = int(years[year_index]))
                                end_1 = end_1.replace(year = int(years[year_index]))
                                if now >= start_1 + start_plus and now < end_1 + end_plus:
                                    reg_val = wyt_data["Val1"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 15.6
                                    val_found = 1
                                end_1 = end_1.replace(year=2017)
                            if start_2 is not None and val_found == 0:
                                start_2 = start_2.replace(year = int(years[year_index]))
                                end_2 = end_2.replace(year = int(years[year_index]))
                                if now >= start_2 + start_plus and now < end_2 + end_plus:
                                    reg_val = wyt_data["Val2"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 14.0
                                    val_found = 1
                                end_2 = end_2.replace(year=2017)

                            if start_3 is not None and val_found == 0:
                                start_3 = start_3.replace(year=int(years[year_index]))
                                end_3 = end_3.replace(year=int(years[year_index]))
                                if now >= start_3 + start_plus and now < end_3 + end_plus:
                                    reg_val = wyt_data["Val3"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 12.5
                                    val_found = 1
                                end_3 = end_3.replace(year=2017)

                            if start_4 is not None and val_found == 0:
                                start_4 = start_4.replace(year=int(years[year_index]))
                                end_4 = end_4.replace(year=int(years[year_index]))
                                if now >= start_4 + start_plus and now < end_4 + end_plus:
                                    reg_val = wyt_data["Val4"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 19.0
                                    val_found = 1
                                end_4 = end_4.replace(year=2017)

                            if start_5 is not None and val_found == 0:
                                start_5 = start_5.replace(year=int(years[year_index]))
                                end_5 = end_5.replace(year=int(years[year_index]))
                                if now >= start_5 + start_plus and now < end_5 + end_plus:
                                    reg_val = wyt_data["Val5"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 16.5
                                end_5 = end_5.replace(year=2017)

                            if start_6 is not None and val_found == 0:
                                start_6 = start_6.replace(year=int(years[year_index]))
                                end_6 = end_6.replace(year=int(years[year_index]))
                                if now >= start_6 + start_plus and now < end_6 + end_plus:
                                    reg_val = wyt_data["Val6"].iloc[0]
                                    if def_flag == 1 and def_stn_flag == 1:
                                        reg_val = 15.6
                                    val_found = 1
                                end_6 = end_6.replace(year=2017)

                            std_ts.append(reg_val)
                        std_ts = std_ts[1:]
                        std_ts.append(std_ts[-1])

                        if count < 1:
                            std_ts_df = pd.DataFrame({std_nm: std_ts})
                            std_nms = [str(std_nm)]
                            count = 1
                        else:
                            std_ts_df[std_nm] = std_ts
                            std_nms.append(str(std_nm))

            ThirtyDay_Stns = ["RSAN112", "RSAN072", "OLDR_MIDR", "ROLD059"]
            Monthly_Stns = ["RSAC081", "SLMZU025", "SLMZU011", "SLCBN002", "SLSUS012", "CHWST000", "CHDMC004"]
            new_EC_vals = []

            if out_statns[station_index] in ThirtyDay_Stns:
                for k in range(len(values)):
                    if k < 29:
                        new_val = np.mean(values[:k+1])
                    else:
                        sub_vals = []
                        for k2 in range(0, 30):
                            val = values[k-k2]
                            if val > -1000:
                                sub_vals.append(val)

                        new_val = np.mean(sub_vals)
                    new_val = new_val / 1000
                    new_EC_vals.append(new_val)
            elif out_statns[station_index] in Monthly_Stns:

                daily_df = pd.DataFrame({'Date': dates, 'Value': values})
                daily_df['Month'] = [date.month for date in dates]
                daily_df['Year'] = [date.year for date in dates]

                month_df = daily_df.groupby(['Year', 'Month'], as_index=False).mean()

                month_df['Date'] = [datetime.datetime(1921, 3, 1) + pd.DateOffset(months=i) + datetime.timedelta(days=-1) for i in range(len(month_df))]
                for k in range(len(values)):
                     m = daily_df['Month'][k]
                     y = daily_df['Year'][k]
                     sub_df = month_df[month_df['Month'] == m]
                     sub_df = sub_df[sub_df['Year'] == y]
                     new_val = sub_df['Value'].iloc[0] / 1000
                     new_EC_vals.append(new_val)
            else: # 14 days
                for k in range(len(values)):
                    if k < 13:
                        new_val = np.mean(values[:k+1])
                    else:
                        sub_vals = []
                        for k2 in range(0, 14):
                            val = values[k-k2]
                            if val > -1000:
                                sub_vals.append(val)

                        new_val = np.mean(sub_vals)
                    new_val = new_val / 1000
                    new_EC_vals.append(new_val)

            if cl_flag > 0:
                new_Cl_vals = []
                for val in values:
                    if val < -1000:
                        new_val = None
                    elif val <=280:
                        new_val = val*0.15 - 12
                    else:
                        new_val = val*0.285 - 50
                    new_Cl_vals.append(new_val)
            else:
                new_Cl_vals = [None] * len(values)

            nd_flags.append(nd_flag)

            df = pd.DataFrame({
                "Var Name": stations,
                "Location": locations,
                "Var type": cpart,
                "Date": dates,
                "ValueEC": new_EC_vals,
                "ValueCl": new_Cl_vals,
                "Study Scenario": study_scenario,
                "Study Type": study_type,
                "UnitsEC": unitsEC,
                "UnitsCl": unitsCl
            })
            df = pd.concat([df, std_ts_df], axis=1)

            # Add water year type column
            df["WYT"] = wyts

            # Filter to start in WY 1921 (i.e., Oct 1921 or later)
            df = df[(df["Date"].dt.year > 1921) | ((df["Date"].dt.year == 1921) & (df["Date"].dt.month > 9))]

            # Remove rows with invalid EC values
            df = df[df["ValueEC"] > -1000000]

            # Remove rows where all values are NA
            df = df.dropna(how='all')

            # Rename columns to match final output
            df.columns = [
                "Var Name", "Location", "Var type", "Date", "ValueEC", "ValueCl",
                "Study Scenario", "Study Type", "UnitsEC", "UnitsCl",
                "D1641AG", "D1641FWS", "D1641MI", "D1641MIDNumDays", "D1641MIDThreshold",
                "MIAntiochNumDays", "MIAntiochThreshold", "MIOther", "SAC INDEX"
            ]

            # Write to file (no header, no row names, comma-separated)
            df.to_csv(tsfilename, index=False, header=False, mode='a')  # quoting=3 disables quoting (like quote = FALSE in R)

    print(f"Done with creating {os.path.basename(tsfilename)}")

    # Read the output ts file
    df_pre = pd.read_csv(tsfilename, skiprows=1, header=None)
    df_pre_names = pd.read_csv(tsfilename, nrows=1, header=None)

    # Assign column names from the first row
    df_pre.columns = df_pre_names.iloc[0]

    # Calculate differences
    df_pre['AG_diff'] = df_pre['ValueEC'] - df_pre['D1641AG']
    df_pre['FWS_diff'] = df_pre['ValueEC'] - df_pre['D1641FWS']
    df_pre['MI_diff'] = df_pre['ValueCl'] - df_pre['D1641MI']

    ndblw_arr = []
    ndvio_arr = []
    ndblw_arrA = []
    ndvio_arrA = []

    for ind in range(len(out_statns)):
        stn = out_statns[ind]
        v_type = f"EC-{out_stats[ind]}"

        df_stn2 = df_pre[df_pre['Var Name'] == stn]
        df_stn = df_stn2[df_stn2['Var type'] == v_type]

        if df_stn.empty:
            continue

        MIDT = df_stn['D1641MIDThreshold'].values
        AntiochT = df_stn['MIAntiochThreshold'].values

        if not (pd.isna(MIDT[0]) or not pd.isna(AntiochT[0])):
            dt = pd.to_datetime(df_stn['Date'])
            years = [date.year for date in dt]
            months = [date.month for date in dt]
            Clval = df_stn['ValueCl'].values

            if out_stats[ind] == 'MEAN':
                MIDND = df_stn['D1641MIDNumDays'].values
                itr = 0
                yrtot = 0
                for k in range(len(df_stn)):
                    if k == 0:
                        yr_prev = years[k]
                    yr = years[k]

                    if yr != yr_prev:
                        yrtot += itr
                        ndblw_arr[-1] = yrtot
                        if np.isnan(yrtot) or np.isnan(MIDND[k - 1]):
                            ndvio_arr[-1] = 0
                        elif yrtot < MIDND[k - 1]:
                            ndvio_arr[-1] = 1
                        else:
                            ndvio_arr[-1] = 0
                        ndblw_arr.append(np.nan)
                        ndvio_arr.append(0)
                        yrtot = 0
                        itr = 0
                        yr_prev = yr
                    elif k == len(df_stn) - 1:
                        yrtot += itr
                        ndblw_arr.append(yrtot)
                        if np.isnan(yrtot) or np.isnan(MIDND[k]):
                            ndvio_arr.append(0)
                        elif yrtot < MIDND[k]:
                            ndvio_arr.append(1)
                        else:
                            ndvio_arr.append(0)
                    else:
                        ndblw_arr.append(np.nan)
                        ndvio_arr.append(0)

                    if np.isnan(MIDT[k]) or np.isnan(Clval[k]):
                        continue
                    elif Clval[k] < MIDT[k]:
                        itr += 1
                    elif itr >= 14:
                        yrtot += itr
                        itr = 0
                    else:
                        itr = 0

                ndblw_add = [np.nan] * len(df_stn)
                ndblw_arrA.extend(ndblw_add)
                ndvio_arrA.extend(ndblw_add)

            else:
                AntiochND = df_stn['MIAntiochNumDays'].values
                itrA = 0
                for k in range(len(df_stn)):
                    if k == 0:
                        yr_prev = years[k]
                    yr = years[k]

                    if yr != yr_prev:
                        ndblw_arrA[-1] = itrA
                        if np.isnan(itrA) or np.isnan(AntiochND[k - 1]):
                            ndvio_arrA[-1] = 0
                        elif itrA < AntiochND[k - 1]:
                            ndvio_arrA[-1] = 1
                        else:
                            ndvio_arrA[-1] = 0
                        ndblw_arrA.append(np.nan)
                        ndvio_arrA.append(0)
                        itrA = 0
                        yr_prev = yr
                    elif k == len(df_stn) - 1:
                        ndblw_arrA.append(itrA)
                        if np.isnan(itrA) or np.isnan(AntiochND[k]):
                            ndvio_arrA.append(0)
                        elif itrA < AntiochND[k]:
                            ndvio_arrA.append(1)
                        else:
                            ndvio_arrA.append(0)
                    else:
                        ndblw_arrA.append(np.nan)
                        ndvio_arrA.append(0)

                    if np.isnan(AntiochT[k]) or np.isnan(Clval[k]):
                        continue
                    elif Clval[k] < AntiochT[k]:
                        itrA += 1

                ndblw_arr.extend([np.nan] * len(df_stn))
                ndvio_arr.extend([np.nan] * len(df_stn))

        else:
            ndblw_add = [np.nan] * len(df_stn)
            ndblw_arr.extend(ndblw_add)
            ndvio_arr.extend(ndblw_add)
            ndblw_arrA.extend(ndblw_add)
            ndvio_arrA.extend(ndblw_add)

    # Assign new columns to df_pre
    df_pre['NumDaysBlw'] = ndblw_arr
    df_pre['NumDaysViolated'] = ndvio_arr
    df_pre['AntiochNumDaysBlw'] = ndblw_arrA
    df_pre['AntiochNumDaysViolated'] = ndvio_arrA
    # Construct output filename
    tsfilenamediff = f"DSM2ComplianceDiffData_{input_model_name}.csv"
    tsfilenamediff = os.path.join(outfile_location, tsfilenamediff)

    # Define custom header
    custom_header = [
        "Var Name", "Location", "Var type", "Date", "ValueEC", "ValueCl", "Study Scenario", "Study Type",
        "UnitsEC", "UnitsCl", "D1641AG", "D1641FWS", "D1641MI", "D1641MIDNumDays", "D1641MIDThreshold",
        "MIAntiochNumDays", "MIAntiochThreshold", "MIOther", "SAC INDEX", "DiffAG", "DiffFWS", "DiffMI",
        "NumDaysBlw", "NumDaysViolated", "AntiochNumDaysBlw", "AntiochNumDaysViolated"
    ]

    # Write to CSV with custom header
    with open(tsfilenamediff, 'w') as f:
        f.write(','.join(custom_header) + '\n')
        df_pre.to_csv(f, index=False, header=False, na_rep='NA')

def percentile(dl_values, percent, key=lambda x:x):
    """
    Find the percentile of a list of values.

    Parameters
    ----------
    dl_values: list
        list of values to get the value from. Note this MUST BE already sorted.
    percent: float
        a float value from 0.0 to 1.0. The percentile we want to find
    key: function
        optional key function to compute value from each element of dl_values

    Returns
    -------
    the percentile of the values
    """

    if not dl_values:
        return None
    k = (len(dl_values) - 1) * percent
    f = math.floor(k)
    c = math.ceil(k)
    if f == c:
        return key(dl_values[int(k)])
    d0 = key(dl_values[int(f)]) * (c - k)
    d1 = key(dl_values[int(c)]) * (k - f)
    return d0+d1


def combine_percentiles(scen_nm):
    """
    Calculates the percentiles for all of the compliance locations. For each compliance point ranks the values from the get_dsm2_timeseries_data output.
    Write out a _Percentiles.csv file for each location. Also writes a _Dates.csv for each and a compliance summary for the scenario
    Parameters
    ----------
    scen_nm: str
        name of the scenario

    Returns
    -------
    None
    """

    infn = "./water_qual_csvs/DSM2ComplianceDiffData_" + scen_nm + ".csv"
    outfn = "./water_qual_csvs/_ComplianceSummary/DSM2ComplianceSummary_" + scen_nm + ".csv"
    out_dir = "./water_qual_csvs/_Percentiles/"

    if not os.path.exists("./water_qual_csvs/_ComplianceSummary/"):
        os.mkdir("./water_qual_csvs/_ComplianceSummary/")

    if not os.path.exists(out_dir):
        os.mkdir(out_dir)

    with open(infn, "r") as inf:
        in_data = inf.readlines()

    outf = open(outfn, "w")

    print(in_data[0].rstrip().split(","))

    outf.write("VarName,DiffAG,DiffFWS,DiffMI,NumDaysBlw,NumDaysViolated,AntiochNumDaysBlw,AntiochNumDaysViolated,TotNumDays_D1641AG,TotNumDays_D1641FWS,TotNumDays_D1641MI\n")
    date_ind = in_data[0].rstrip().split(",").index("Date")
    wyt_ind = in_data[0].rstrip().split(",").index("SAC INDEX")
    diff_ag_ind = in_data[0].rstrip().split(",").index("DiffAG")
    diff_fws_ind = in_data[0].rstrip().split(",").index("DiffFWS")
    diff_mi_ind = in_data[0].rstrip().split(",").index("DiffMI")
    mi_numdaysblw_ind = in_data[0].rstrip().split(",").index("NumDaysBlw")
    mi_numdaysstandard_ind = in_data[0].rstrip().split(",").index("D1641MIDNumDays")
    mi_numdaysvio_ind = in_data[0].rstrip().split(",").index("NumDaysViolated")
    antioch_numdaysblw_ind = in_data[0].rstrip().split(",").index("AntiochNumDaysBlw")
    antioch_numdaysvio_ind = in_data[0].rstrip().split(",").index("AntiochNumDaysViolated")

    for i, x in enumerate(in_data):
        if in_data[i] == '\n': continue
        if i == 0: continue
        if i == 1:
            diff_ag = 0
            diff_ag_arr = []
            diff_ag_date_arr = []
            diff_ag_wyt_arr = []
            diff_fws = 0
            diff_fws_arr = []
            diff_fws_date_arr = []
            diff_fws_wyt_arr = []
            diff_mi = 0
            diff_mi_arr = []
            diff_mi_date_arr = []
            diff_mi_wyt_arr = []
            mi_numdaysblw_arr = []
            mi_numdaysblw = 0
            mi_numdaysvio = 0
            antioch_numdaysblw = 0
            antioch_numdaysvio = 0
            VarNamePrev = in_data[i].split(",")[0]
            VarName = in_data[i].split(",")[0]
            TotNumDays_D1641AG = 0
            TotNumDays_D1641FWS = 0
            TotNumDays_D1641MI = 0
        if i == (len(in_data) - 1):
            if in_data[i].rstrip().split(",")[diff_ag_ind] != "NA":
                TotNumDays_D1641AG += 1
                diff_ag_arr.append(float(in_data[i].rstrip().split(",")[diff_ag_ind]))
                if float(in_data[i].rstrip().split(",")[diff_ag_ind]) > 0:
                    diff_ag_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_ag_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_ag += 1
            if in_data[i].rstrip().split(",")[diff_fws_ind] != "NA":
                TotNumDays_D1641FWS += 1
                diff_fws_arr.append(float(in_data[i].rstrip().split(",")[diff_fws_ind]))
                if float(in_data[i].rstrip().split(",")[diff_fws_ind]) > 0:
                    diff_fws_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_fws_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_fws += 1
            if in_data[i].rstrip().split(",")[diff_mi_ind] != "NA":
                TotNumDays_D1641MI += 1
                diff_mi_arr.append(float(in_data[i].rstrip().split(",")[diff_mi_ind]))
                if float(in_data[i].rstrip().split(",")[diff_mi_ind]) > 0:
                    diff_mi_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_mi_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_mi += 1
            if in_data[i].rstrip().split(",")[mi_numdaysblw_ind] != "NA":
                # if float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) > 0:
                mi_numdaysblw += 1
                mi_numdaysblw_arr.append(float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) - float(in_data[i].rstrip().split(",")[mi_numdaysstandard_ind]))
            if in_data[i].rstrip().split(",")[mi_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[mi_numdaysvio_ind]) > 0: mi_numdaysvio += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysblw_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysblw_ind]) > 0: antioch_numdaysblw += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysvio_ind]) > 0: antioch_numdaysvio += 1
            print(VarNamePrev, diff_ag, diff_fws, diff_mi, mi_numdaysblw, mi_numdaysvio, antioch_numdaysblw, antioch_numdaysvio, TotNumDays_D1641AG, TotNumDays_D1641FWS, TotNumDays_D1641MI)
            outf.write(VarNamePrev + "," + str(diff_ag) + "," + str(diff_fws) + "," + str(diff_mi) + "," + str(mi_numdaysblw) + "," + str(mi_numdaysvio) + "," + str(antioch_numdaysblw) + "," + str(
                antioch_numdaysvio) + "," + str(TotNumDays_D1641AG) + "," + str(TotNumDays_D1641FWS) + "," + str(TotNumDays_D1641MI) + "\n")
            if len(diff_ag_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Dates.csv", "w") as date_file:
                    for j in range(len(diff_ag_date_arr)): date_file.write(str(diff_ag_date_arr[j]) + "," + str(diff_ag_wyt_arr[j]) + "\n")
                diff_ag_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_ag_arr, float(j + 1) / 100.0)) + "\n")
            if len(diff_fws_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Dates.csv", "w") as date_file:
                    for j in range(len(diff_fws_date_arr)): date_file.write(str(diff_fws_date_arr[j]) + "," + str(diff_fws_wyt_arr[j]) + "\n")
                diff_fws_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_fws_arr, float(j + 1) / 100.0)) + "\n")
            if len(diff_mi_arr) > 0:
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Dates.csv", "w") as date_file:
                    for j in range(len(diff_mi_date_arr)): date_file.write(str(diff_mi_date_arr[j]) + "," + str(diff_mi_wyt_arr[j]) + "\n")
                diff_mi_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(diff_mi_arr, float(j + 1) / 100.0)) + "\n")
            if len(mi_numdaysblw_arr) > 0:
                mi_numdaysblw_arr = [float(i) for i in mi_numdaysblw_arr]
                mi_numdaysblw_arr.sort()
                with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Days_Percentile.csv", "w") as percent_file:
                    for j in range(100):
                        percent_file.write(str(j + 1) + "," + str(percentile(mi_numdaysblw_arr, float(j + 1) / 100.0)) + "\n")

        else:
            VarName = in_data[i].split(",")[0]
            if VarName != VarNamePrev:
                print(VarName)
                print(VarNamePrev, diff_ag, diff_fws, diff_mi, mi_numdaysblw, mi_numdaysvio, antioch_numdaysblw, antioch_numdaysvio, TotNumDays_D1641AG, TotNumDays_D1641FWS, TotNumDays_D1641MI)
                outf.write(
                    VarNamePrev + "," + str(diff_ag) + "," + str(diff_fws) + "," + str(diff_mi) + "," + str(mi_numdaysblw) + "," + str(mi_numdaysvio) + "," + str(antioch_numdaysblw) + "," + str(
                        antioch_numdaysvio) + "," + str(TotNumDays_D1641AG) + "," + str(TotNumDays_D1641FWS) + "," + str(TotNumDays_D1641MI) + "\n")
                print("agg_arr")
                if len(diff_ag_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Dates.csv", "w") as date_file:
                        for j in range(len(diff_ag_date_arr)): date_file.write(str(diff_ag_date_arr[j]) + "," + str(diff_ag_wyt_arr[j]) + "\n")
                    diff_ag_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641AG_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_ag_arr, float(j + 1) / 100.0)) + "\n")
                print("fws_arr")
                if len(diff_fws_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Dates.csv", "w") as date_file:
                        for j in range(len(diff_fws_date_arr)): date_file.write(str(diff_fws_date_arr[j]) + "," + str(diff_fws_wyt_arr[j]) + "\n")
                    diff_fws_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641FWS_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_fws_arr, float(j + 1) / 100.0)) + "\n")
                print("mi_arr")
                if len(diff_mi_arr) > 0:
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Dates.csv", "w") as date_file:
                        for j in range(len(diff_mi_date_arr)): date_file.write(str(diff_mi_date_arr[j]) + "," + str(diff_mi_wyt_arr[j]) + "\n")
                    diff_mi_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(diff_mi_arr, float(j + 1) / 100.0)) + "\n")
                print("mid_arr")
                print(mi_numdaysblw_arr)
                if len(mi_numdaysblw_arr) > 0:
                    mi_numdaysblw_arr = [float(i) for i in mi_numdaysblw_arr]
                    print(mi_numdaysblw_arr)
                    mi_numdaysblw_arr.sort()
                    with open(out_dir + VarNamePrev + "_" + scen_nm + "_D1641MI_Days_Percentile.csv", "w") as percent_file:
                        for j in range(100):
                            percent_file.write(str(j + 1) + "," + str(percentile(mi_numdaysblw_arr, float(j + 1) / 100.0)) + "\n")

                diff_ag = 0
                diff_ag_arr = []
                diff_ag_date_arr = []
                diff_ag_wyt_arr = []
                diff_fws = 0
                diff_fws_arr = []
                diff_fws_date_arr = []
                diff_fws_wyt_arr = []
                diff_mi = 0
                diff_mi_arr = []
                diff_mi_date_arr = []
                diff_mi_wyt_arr = []
                mi_numdaysblw_arr = []
                mi_numdaysblw = 0
                mi_numdaysvio = 0
                antioch_numdaysblw = 0
                antioch_numdaysvio = 0
                TotNumDays_D1641AG = 0
                TotNumDays_D1641FWS = 0
                TotNumDays_D1641MI = 0
            if in_data[i].rstrip().split(",")[diff_ag_ind] != "NA":
                TotNumDays_D1641AG += 1
                diff_ag_arr.append(float(in_data[i].rstrip().split(",")[diff_ag_ind]))
                if float(in_data[i].rstrip().split(",")[diff_ag_ind]) > 0:
                    diff_ag_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_ag_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_ag += 1
            if in_data[i].rstrip().split(",")[diff_fws_ind] != "NA":
                TotNumDays_D1641FWS += 1
                diff_fws_arr.append(float(in_data[i].rstrip().split(",")[diff_fws_ind]))
                if float(in_data[i].rstrip().split(",")[diff_fws_ind]) > 0:
                    diff_fws_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_fws_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_fws += 1
            if in_data[i].rstrip().split(",")[diff_mi_ind] != "NA":
                TotNumDays_D1641MI += 1
                diff_mi_arr.append(float(in_data[i].rstrip().split(",")[diff_mi_ind]))
                if float(in_data[i].rstrip().split(",")[diff_mi_ind]) > 0:
                    diff_mi_date_arr.append(in_data[i].rstrip().split(",")[date_ind])
                    diff_mi_wyt_arr.append(in_data[i].rstrip().split(",")[wyt_ind])
                    diff_mi += 1
            if in_data[i].rstrip().split(",")[mi_numdaysblw_ind] != "NA":
                # if float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) > 0:
                mi_numdaysblw += 1
                mi_numdaysblw_arr.append(float(in_data[i].rstrip().split(",")[mi_numdaysblw_ind]) - float(in_data[i].rstrip().split(",")[mi_numdaysstandard_ind]))
            if in_data[i].rstrip().split(",")[mi_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[mi_numdaysvio_ind]) > 0: mi_numdaysvio += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysblw_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysblw_ind]) > 0: antioch_numdaysblw += 1
            if in_data[i].rstrip().split(",")[antioch_numdaysvio_ind] != "NA":
                if float(in_data[i].rstrip().split(",")[antioch_numdaysvio_ind]) > 0: antioch_numdaysvio += 1
            VarNamePrev = VarName

    outf.close()

def combine_all_runs(studies, percentile_files):
    """
    Combines the outputs from the combine_percentiles function's _Percentile cvs into one csv file that is used for the plots.
    Parameters
    ----------
    studies: list
        The different studies/alternatives to combine
    percentile_files:
        All of the csv files ending in _Percentile. These have the percentile data

    Returns
    -------
    final_data_frame: dataframe
        The combined dataframe
    """

    data = []
    flag = 0
    header = []

    for study in studies:
        for file in percentile_files:
            if study in file:
                with open("./water_qual_csvs/_Percentiles/" + file, "r") as inf:
                    in_data = inf.readlines()
                if "AG_" in file:
                    comp = "AG"
                elif "FWS_" in file:
                    comp = "FWS"
                elif "MI_Percentile" in file:
                    comp = "MI"
                elif "Days" in file:
                    comp = "MI_Days"
                if flag == 0:
                    header.append("Percentile")
                    header.append(study + "_" + file.split("_")[0] + "_" + comp)
                    col = []
                    for line in in_data:
                        col.append(line.rstrip().split(",")[0])
                    data.append(col)
                    col = []
                    for line in in_data:
                        col.append(line.rstrip().split(",")[1])
                    data.append(col)
                    flag = 1
                else:
                    if file.split("_")[0] == "OLDR" and file.split("_")[1] == "MIDR":
                        header.append(study + "_" + file.split("_")[0] + "_" + file.split("_")[1] + "_" + comp)
                        col = []
                        for line in in_data:
                            col.append(line.rstrip().split(",")[1])
                        data.append(col)
                    else:
                        header.append(study + "_" + file.split("_")[0] + "_" + comp)
                        col = []
                        for line in in_data:
                            col.append(line.rstrip().split(",")[1])
                        data.append(col)

    final_data_frame = pd.DataFrame(data).transpose()
    final_data_frame.columns = header
    final_data_frame = final_data_frame.apply(pd.to_numeric)

    return final_data_frame

def create_water_qual_plot(df_percentiles, fig_value, plot_directory, alts, line_styles, line_colors, d_ymin, d_ymax):
    """
    Creates the plot of the probability of compliance. Write the plot to a file but returns the path.
    Parameters
    ----------
    df_percentiles: dataframe
        Data frame of all of the data
    fig_value: str
        Value to plot
    plot_directory: str
        path to the folder to hold the plots
    alts: dict
        Dictionary of alternatives and names to show on plot
    line_styles: list
        List of line styles to show on plot
    line_colors: list
        List of line colors to show on plot
    d_ymin: float
        y axis minimum
    d_ymax: float
        y axis maximum

    Returns
    -------
    Path to the plot
    """
    # Check for/create directory to store monthly exceedance plots
    if not os.path.exists(plot_directory):
        os.makedirs(plot_directory)


    # define size and borders
    fig, axs = plt.subplots(figsize=(9, 5.5), linewidth=3, edgecolor="black")

    for fig_index, display_name in enumerate(alts):

        model_name = alts[display_name].split('.')[0]
        # Dataset for this alt
        df_alt_data = df_percentiles[['Percentile', model_name+'_'+fig_value]]

        # plot exceedance probability vs monthly EC
        percentages = range(0, 101, 10)
        percentage_labels = [f"{int(i)}%" for i in percentages]

        axs.plot(df_alt_data['Percentile'].values, df_alt_data[model_name+'_'+fig_value].values, color=line_colors[fig_index],
                 linestyle=line_styles[fig_index], label=display_name)
        axs.set_xticks(percentages)
        axs.set_xticklabels(percentage_labels)

        # set the Y axis depending on if its chloride or EC
        if fig_value.split('_')[-1] == 'MI':
            axs.set_ylabel("Difference in Chloride (Scenario minus Standard) (mg/L)")
        else:
            axs.set_ylabel("Difference in EC (Scenario minus Standard) (mmhos/cm)")
        axs.set_xlabel("Probability of Compliance (%)")

        # set the y limits
        axs.set_ylim(d_ymin, d_ymax)

        # Save this parameter to orient the legend correctly
        axbox = axs.get_position()

        # Add gridlines
        plt.grid(color='gray', linestyle='--', linewidth=0.8)

        # Add a legend
        plt.legend(loc='center', ncol=4, bbox_to_anchor=[axbox.x0 + 0.5 * axbox.width, 1.08])

    # flip x-axis
    axs.invert_xaxis()

    plt.savefig(plot_directory + "/" + fig_value + "_exceedance" + ".png")

    plt.close()

    return plot_directory + "/" + fig_value + "_exceedance" + ".png"

def get_wq_location_data():
    """
    Gets the table of location data for the plots
    Returns
    -------
    crosswalk: dataframe
        Dataframe containing all locations, names, and y axis limits
    """
    
    crosswalk = pd.read_excel("../inputs/location_code_crosswalk_water_quality.xlsx")
    crosswalk.drop(columns="Model", inplace=True)
    return crosswalk

if __name__ == '__main__':

    # this dictionary should hold the display name and the full DSS file for each alternative in the order you want them displayed
    # all of these dss files should be in the studies folder
    # note that the hydrology should be in the name (ex: '2022MED')
    # ex: {'NAA':"NAA_2022Med_090723_EC_p.dss",.... }
    scenario_names = {'NAA':"NAA_2022Med_090723_EC_p.dss",
                      "ALT1":"ALT1_2022Med_090923_EC_p.dss",
                      "Alt2woTUCPwoVA": "ALT2v1_woTUCP_2022Med_091324_EC_p.dss",
                      "Alt2wTUCPwoVA": "ALT2v1_wTUCP_2022Med_091324_EC_p.dss",
                      "Alt2woTUCPDeltaVA": "ALT2v2_woTUCP_2022Med_091324_EC_p.dss",
                      "Alt2woTUCPAllVA": "ALT2v3_woTUCP_2022Med_091324_EC_p.dss",
                      "ALT3": "ALT3_2022Med_101323_EC_p.dss",
                      "ALT4": "ALT4_2022MED_091624_EC_p.dss",
                      "Action 5": "ALT5_wTUCP_2022Med_052125_EC_p.dss"
    }

    # this is the template for the Word doc, generally doesn't need to change
    template = r"..\inputs\template_v2-fonts.docx"

    # name of the temporary document, needs to have no spaces so no OneDrive
    doc_name = rf"C:\Users\fnufferrodriguez\temp_appendix.docx"

    # Name of final word doc, needs to have no spaces so no OneDrive
    new_doc = rf"C:\Users\fnufferrodriguez\Attachment_2-08_Water_Quality_Compliance.docx"

    # if you want to change what plots are plotted, you can change them in inputs\location_code_crosswalk_water_quality.xlsx
    ####END OF USER INPUTS #######

    # Set working directory to the script's location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)

    # List all dss files in studies folder
    studies_path = os.path.abspath(os.path.join(script_dir, "../studies"))
    path_names = [d for d in os.listdir(studies_path)
                if os.path.isfile(os.path.join(studies_path, d)) and d.endswith(".dss")]

    print("Model directories found:")
    print(path_names)

    # Loop through each model directory and call the processing function
    for model_path in path_names:
        print(f"Processing model: {model_path}")
        get_dsm2_timeseries_data(model_path)

    # get the study names
    studies = [study.split(".")[0] for study in path_names]

    # loop through and call the combine percentiles function
    for study_name in studies:
        combine_percentiles(study_name)

    # get the percentile files that were created
    percentile_files = []

    for file in os.listdir("./water_qual_csvs/_Percentiles/"):
        if "Percentile" in file: percentile_files.append(file)

    # call the function to combine them
    final_data_frame = combine_all_runs(studies, percentile_files)

    # get the data for the different plots
    df_location_info = get_wq_location_data()

    # this will hold the alt text for each figure
    alt_text = []

    # set up the document
    appendix_prefix = " F.2.8"
    doc = docx.Document(template)
    doc.add_heading(f"Attachment{appendix_prefix}", level=1)

    # Add caption style for Figure captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.name = 'Times New Roman'

    # add the set up for the plots. A folder to hold them and line colors and styles
    s_plot_directory = "./wq_plots"
    line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
    line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

    # loop though the rows which coorespond to a plot
    for index, location in df_location_info.iterrows():

        # create the plot and capture the path
        s_plot_path = create_water_qual_plot(final_data_frame, location['VarName'], s_plot_directory,
                                             scenario_names, line_styles, line_colors,
                                             location['Ymin'], location['Ymax'])

        # Generate the caption to match the previous doccumentation
        s_fig_caption = 'D1641 ' + location['VarName'].split('_')[-1] + ' ' + location['Location (Title)'] + ' Compliance Exceedance Plot'

        # change to landscape to fit the images
        if index == 0:
            change_orientation(doc, "landscape")

        # Add figure as a picture
        o_fig = doc.add_picture(s_plot_path)

        # Generate fig title
        fig_title_prefix = "Figure " + appendix_prefix + "-"

        # Add title below figure
        add_caption_water_supply(doc, "Figure", fig_title_prefix, s_fig_caption, custom_style="Figure Caption")

        # Add to alt text
        alt_text.append(s_fig_caption)

        if index != len(df_location_info) - 1:
            doc.add_page_break()

    # save the doc to the temporary name
    doc.save(doc_name)

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    # Run vbs script
    # Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    # This will fail if Microsoft Word has document open in the background
    # try opening Task Manager and Ending MS Word Background Task, then rerun

    # Call the vbs script for figure alt text
    result = subprocess.call(
        "cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + "x" + " " + str(0) + " " + alt_text_string_figures + " " + str(len(df_location_info)))

    # check if it worked successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")
    else:
        # Instructions on how to finish formatting numbered captions.
        print("After running this script, \n1. Open Word file and Ctrl+A to select all. Then F9 to update caption numbering.")
