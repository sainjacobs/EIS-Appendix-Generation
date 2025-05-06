from EISAppendixGen_functions import get_locations, parse_dssReader_output, create_exceedance_tables, format_table, create_month_plot, create_stat_plot, change_orientation
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import copy
import os
import datetime
import pandas as pd

if __name__ == "__main__":

    ###USER INPUTS BELOW#####

    #Fields to use from DSS Reader
    #fields = ["C_LWSTN", "C_CLR011"]#, "C_KSWCK", "C_SAC257", "C_SAC240", "C_SAC201", "C_SAC120", "C_FTR059", "C_FTR003"]
    fields = ["BLW CLEAR CREEK"]
    #Scenarios to compare
    #alts = ["NAA", "ALT1", "Alt2woTUCPwoVA", "Alt2woTUCPDeltaVA"]#, "Alt2woTUCPAllVA", "Alt2wTUCPwoVA", "ALT3", "Alt4"]
    alts = ["NAA", "ALT2"]

    report_type = "temperature"


    #TODO Add selection of units (elevation, temperature, provide both cfs and taf?)
    #Get more info from crosswalk, units, description, etc
    #Add salinity and temperature - could break into separate scripts

    # Prefix for tables and figures in appendix
    appendix_prefix = " F.2.2"

    # Path to file with location code crosswalk
    location_cw_path = "C:/Users/emadonna/eis-appendix-generation/inputs/location_code_crosswalk_Temp.xlsx"
    #Path to file with DSSReader output
    #Use output from DSS reader in desired units (CFS or TAF)
    dss_path = "C:/Users/emadonna/eis-appendix-generation/inputs/DSS_contents.xlsx"
    #Path to file with WY Typing data
    wy_flags_path = "C:/Users/emadonna/eis-appendix-generation/inputs/wy_flags.xlsx"

    # Windows command prompt can't save to OneDrive bc of the space in the file path, save locally instead
    # Pass absolute paths to VBS
    #Name of intermediate word doc - update parent directory
    doc_name = "C:/Users/emadonna/eis-appendix-generation/appendix_temp_temperature.docx"
    #Name of final word doc
    new_doc = "C:/Users/emadonna/eis-appendix-generation/appendix_final_temperature.docx"

####END OF USER INPUTS #######

    # Read location from crosswalk based on field later
    locations = get_locations(location_cw_path, fields)

    #compare every run to the baseline run
    comparisons = [["NAA", alt] for alt in alts]
    #Remove first comparison that is NAA and NAA
    comparisons.pop(0)

    #Each comparison will have 3 tables and will be included for every field/location
    num_tables = (len(comparisons) * 3) * len(fields)
    #Include a figure for each month plus 6 full simulation period statistics plots
    num_figures = 12 + 6

    # Alt Text strings, in order for tables
    alt_text_tables = ["Alt text table example" for t in range(0,num_tables)]
    # Alt text strings, order for figures
    alt_text_figures = ["This figure shows data also presented in data tables in this file." for f in range(0,num_figures)]

    # Set table value name
    table_value = "Monthly Temperature (DEG-F)"

    #Set figure value name
    fig_value = "Monthly Temperature (DEG-F)"
    # Create an instance of a word document
    doc = docx.Document()

    #Convert DSS Output to monthly from daily for temp by averaging
    daily_data = pd.read_excel(dss_path)
    daily_data.drop(columns = ["Date", "Year", "Day", "DY"], inplace = True)
    monthly_data = daily_data.groupby(['Scenario', "WY", "Month"]).mean()
    monthly_data.reset_index(inplace = True)
    monthly_data.drop(columns = "Index", inplace = True)

    for field_index, location in enumerate(fields):

        ##### Read DSSReader output ########
        dfs = parse_dssReader_output(dss_path, alts, report_type, location)

        #Create Exceedance Tables from DSS Reader output
        e_dfs, exc_prob = create_exceedance_tables(dfs, wy_flags_path, report_type)

        ##### Use docx package to create a document with formatted table objects and save to Word .docx file ###########

        ## Add a table for each run in each comparison for the current field to the doc
        for comparison_index, scenario in enumerate(comparisons):

            #Then third table for each comparison should be first alt minus second alt listed
            comparison_tables = []
            for alt in scenario:
                #Get exceedance tables for each of the runs in the current comparison
                comparison_tables.append(e_dfs[alts.index(alt)])
            #Add one more table for second alt minus the baseline
            comparison_tables.append(comparison_tables[1].iloc[:, 1:] - comparison_tables[0].iloc[:, 1:])
            #Add the labels column back into the differenced table
            comparison_tables[-1].insert(0, "Statistic", comparison_tables[0]["Statistic"])

            #Set up Comparison labels to be used in table titles
            comparison_table_labels = ["NAA", scenario[1], scenario[1] + " Minus " + "NAA"]

            for comp_table_index, table in enumerate(comparison_tables):

                # Generate table title
                table_title = "Table " + appendix_prefix + "-" + str(field_index + 1) + "-" + str(comparison_index + 1) + chr(ord('a') + comp_table_index)  +". " + locations[field_index] + ", " + comparison_table_labels[comp_table_index] + ", " + table_value

                # Add caption above table
                p = doc.add_paragraph()
                run = p.add_run(table_title)
                run.font.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(1)

                # add a table to the end and create a reference variable
                # extra row is so we can add the header row
                t = doc.add_table(table.shape[0] + 1, table.shape[1])
                #Format table for report
                format_table(t, table, doc, report_type)

            # Add footnotes to the final table
            if comp_table_index == (len(comparison_tables) - 1):
                # Add footnotes at end of table
                footnote0 = doc.add_paragraph()
                run = footnote0.add_run("a")
                run.font.superscript = True
                run1 = footnote0.add_run(" Based on the 100-year simulation period.")
                run1.font.size = Pt(9)
                footnote0.paragraph_format.space_after = Pt(1)

                footnote1 = doc.add_paragraph()
                run = footnote1.add_run('* All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
                run.font.size = Pt(9)
                footnote1.paragraph_format.space_before = Pt(1)
                footnote1.paragraph_format.space_after = Pt(1)

                footnote2 = doc.add_paragraph()
                run = footnote2.add_run(
                    '* Water Year Types defined by the Sacramento Valley 40-30-30 Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
                run.font.size = Pt(9)
                footnote2.paragraph_format.space_before = Pt(1)
                footnote2.paragraph_format.space_after = Pt(1)

                footnote3 = doc.add_paragraph()
                run = footnote3.add_run('* Water Year Types results are displayed with calendar year – year type sorting.')
                run.font.size = Pt(9)
                footnote3.paragraph_format.space_before = Pt(1)

        #####Create Monthly EC and full simulation period statistic plots, save locally as images#####

        #Individual Month Plots
        fig_dfs = copy.deepcopy(e_dfs)

       #Format percent exceedances for labels
        exc_percents = [str(x).split(".")[0] + "%" for x in exc_prob.values]
        #Remove simulation period statistic rows
        for fig_index in range(len(fig_dfs)):
            fig_dfs[fig_index] = fig_dfs[fig_index][:-6]
            #Add formatted exceedance probability percents back to dfs
            fig_dfs[fig_index]["exc_prob"] = exc_percents

        #Can plot up to 8 scenarios, these lines prepare linestyle and color
        line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g"]
        line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-."]

        # Flip doc to landscape orientation for images
        change_orientation(doc, "landscape")

        #Iterate through the dfs and create a figure for each month
        #Save month plots to directory
        month_directory = "month_plots"
        for month in fig_dfs[0].columns[1:-1]:
            create_month_plot(fig_dfs, fig_value, month, month_directory, alts, line_styles, line_colors)

        ##Simulation Period Statistic Plots###
        stat_fig_dfs = copy.deepcopy(e_dfs)

        for stat_fig_index in range(len(stat_fig_dfs)):
            #keep only simulation period statistic rows
            stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index][-6:]
            #Transpose to plot all months at once
            stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index].transpose()
            #Drop first row
            stat_fig_dfs[stat_fig_index].rename(columns=stat_fig_dfs[stat_fig_index].iloc[0], inplace=True)
            stat_fig_dfs[stat_fig_index].drop(stat_fig_dfs[stat_fig_index].index[0], inplace=True)
            #Add abbreviated month name column
            stat_fig_dfs[stat_fig_index]["month"] = ["Oct", "Nov", "Dec", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug",
                                        "Sep"]

        #Check for/create directory to save stat plots
        stat_directory = "stat_plots"

        #Labels to use in stat plot titles
        stats = ["Full Simulation Period", "Wet Water Years (28%)", "Above Normal Water Years (14%)",
                 "Below Normal Water Years (18%)",
                 "Dry Water Years (24%)", 'Critical Water Years (16%)']

        #Iterate through each stat and plot month abbreivated name by EC in current type of year
        for stat in stats:
            create_stat_plot(stat_fig_dfs, fig_value, stat, stat_directory, alts, line_styles, line_colors)

        ##Add saved figures to docx object as images####

        #Get saved month and stat plots as lists
        month_plots = os.listdir(month_directory)
        stat_plots = os.listdir(stat_directory)

        #Iterate through each monthly figure in the month plots directory
        for month_index, file in enumerate(month_plots):
            # Center figures in middle of page by adding some new lines above
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()
            run.add_break()

            # Generate fig title
            fig_title = "Figure " + appendix_prefix + str(month_index + 1) + ". " + locations[field_index] + ", " + datetime.datetime.strptime(file.split("_", 2)[1], '%b').strftime('%B') + " " + fig_value

            # Add title above figure
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(fig_title)
            run.font.size = Pt(12)
            run.font.bold = True

            #Add figure as a picture
            doc.add_picture(month_directory + "/" + file)

            # Add captions below figure
            f = doc.add_paragraph()
            run = f.add_run(
                '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
            run.font.size = Pt(9)
            f.paragraph_format.space_before = Pt(1)
            f.paragraph_format.space_after = Pt(1)

            #Add page break after each figure
            doc.add_page_break()

        # Add stats plots as well

        stat_titles = ["Long Term", "Wet Year", "Above Normal Year", "Below Normal Year", "Dry Year", 'Critical Year']

        for stat_plot_index, file in enumerate(stat_plots):
            # Center figures in middle of page by adding some new lines above
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()
            run.add_break()

            # Generate fig title
            fig_title = "Figure " + appendix_prefix + str(stat_plot_index + 1) + ". " + locations[field_index] + ", " + stat_titles[stat_plot_index] + " Average Temperature"

            # Add title above figure
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(fig_title)
            run.font.size = Pt(12)
            run.font.bold = True

            #Add stat figure as image to document
            doc.add_picture(stat_directory + "/" + file)

            #TODO move captions to user inputs to double check

            # Add captions below figure
            caption0 = doc.add_paragraph()
            run = caption0.add_run(
                '*As defined by the Sacramento Valley 40-30-30 Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
            run.font.size = Pt(9)
            caption0.paragraph_format.space_before = Pt(1)
            caption0.paragraph_format.space_after = Pt(1)

            caption1 = doc.add_paragraph()
            run = caption1.add_run('*These results are displayed with calendar year - year type sorting.')
            run.font.size = Pt(9)
            caption1.paragraph_format.space_before = Pt(1)
            caption1.paragraph_format.space_after = Pt(1)

            caption2 = doc.add_paragraph()
            run = caption2.add_run(
                '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
            run.font.size = Pt(9)
            caption2.paragraph_format.space_before = Pt(1)

            #No need for the page break if it's the final plot of the document
            if stat_plot_index == (len(stat_plots) - 1) and field_index == (len(fields) - 1):
                continue
            else:
                doc.add_page_break()

            #Flip orientation back to portrait for the next group of tables
            if stat_plot_index == (len(stat_plots) - 1):
                # Flip doc to landscape orientation for images
                change_orientation(doc, "portrait")

    # Save docx object to word doc
    doc.save(doc_name)

    ##### Use Python to Run VBS Script that adds alt text to table in saved docx file #######

    # Format alt text for all tables as one string to be passed to vbs
    alt_text_string_tables = ("+").join(alt_text_tables)
    alt_text_string_tables = alt_text_string_tables.replace(" ", "_")

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text_figures)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    #Run vbs script
    #Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    #This will fail if Microsoft Word has document open in the background
    #try opening Task Manager and Ending MS Word Background Task, then rerun
    result = subprocess.call("cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + alt_text_string_tables +  " " + str(num_tables) + " " + alt_text_string_figures + " " + str(num_figures))

    #Remove temporary doc if process ran successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")