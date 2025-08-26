# from IPython.utils.text import date_format

from EISAppendixGen_functions import (get_locations, get_location_wytypes,get_locations_params, parse_dssReader_output, create_exceedance_tables, format_table, create_month_plot, create_stat_plot,
                                      change_orientation, order_elevation_storage_fields, calculate_supply_fields, format_table_supply)
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import copy
import os
import datetime
from docx_caption_formatter import add_caption_byfield, add_caption_water_supply
from docx.enum.style import WD_STYLE_TYPE
import shutil
import numpy as np

if __name__ == "__main__":

    ###USER INPUTS BELOW#####

    #Fields to use from DSS Reader

    # Use for running "elevations" report type, in desired order.
    fields = ["S_TRNTY","S_SHSTA","S_OROVL","S_FOLSM","S_SLUIS","S_SLUIS_CVP","S_SLUIS_SWP","S_MELON","S_MLRTN"]

    # #Use for running "flow" report type
    # fields = ["C_SAC048", "C_YBP020", "C_SAC007", "C_SJR070", "C_SJR070", "C_OMR014", "NDO", "C_SJR225", "C_SJR180",
    #         "C_SJR115", "C_STS004", "C_STS059", "C_KSWCK", "C_SAC257", "C_SAC240", "C_SAC201", "C_SAC120",
    #         "SP_SAC083_YBP037", "C_FTR059", "C_FTR003", "C_NTOMA", "C_AMR004", "C_LWSTN", "C_CLR011",
    #           ]

    # #Used for running "diversions" report type
    # fields = [ "D_LWSTN_CCT011","D_SAC240_TCC001"
    # ,"D_SAC207_GCC007","D_NTOMA_FSC003","D_MLRTN_FRK000","D_MLRTN_MDC006",
    # "D_SAC030_MOK014","TOTAL_EXP", "C_DMC003","C_CAA003_CVP","C_CAA003_SWP","D_DMC007_CAA009"]

    #Temperature
    # alts = ['NAA', "Action 5"]
    # fields = [
    #     "BLW LEWISTON",
    #     "WHISKEYTOWN",
    #     "IGO",
    #     "ABV SACRAMENTO",
    #     "BLW KESWICK",
    #     "BLW CLEAR CREEK",
    #     "BALLS FERRY",
    #     "JELLYS FERRY",
    #     "BEND BRIDGE",

    #     "RED_BLUFF",
    #     "RED BLUFF DAM",
    #     "HAMILTON CITY",
    #     "BLW NIMBUS(HAZEL AVE)",
    #      "WATT AVE",
    #      "ABV CONFLUENCE"
    # ]

    # Water supply fields, order doesn't matter
    alts = ['NAA', "ALT2v1"]
    # fields = ['D_SBP028_17S_PR', 'D_JBC002_17N_PR',
    #    'D_CRK005_17N_NR', 'D_SAC294_03_PA', 'D_CAA143_90_PA2',
    #    'D_WTPCSD_02_PA', 'D_DMC034_71_PA2', 'D_XCC025_72_PA',
    #    'D_WTPBLV_03_PU2', 'D_WTPCSD_02_PU', 'D_WKYTN_02_PU',
    #    'D_WTPJMS_03_PU1', 'D_FOLSM_WTPSJP_CVP', 'D_NFA016_WTPWAL_CVP',
    #    'D_WTPJJO_50_PU', 'D_WTPFTH_02_SU', 'D_WTPFTH_03_SU',
    #    'D_WTPBUK_03_SU', 'D_SAC296_02_SA', 'D_SAC289_03_SA',
    #    'D_GCC027_08N_SA2', 'D_GCC056_08S_SA2', 'D_SAC178_08N_SA1',
    #    'D_SAC159_08N_SA1', 'D_SAC159_08S_SA1', 'D_SAC121_08S_SA3',
    #    'D_SAC109_08S_SA3', 'D_SAC136_18_SA', 'D_SAC122_19_SA',
    #    'D_SAC115_19_SA', 'D_SAC099_19_SA', 'D_SAC091_19_SA',
    #    'D_MTC000_09_SA1', 'D_SAC082_22_SA1', 'D_SAC078_22_SA1',
    #    'D_SAC083_21_SA', 'D_SAC074_21_SA', 'D_MDOTA_73_XA',
    #    'D_DMC111_73_XA', 'D_MDOTA_64_XA', 'D_XCC010_72_XA2',
    #    'D_ARY010_72_XA1', 'D_XCC054_72_XA3', 'D_GCC027_08N_PR1',
    #    'D_MTC000_09_PR', 'D_GCC056_08S_PR', 'D_CBD037_08S_PR',
    #    'D_GCC039_08N_PR2', 'D_MDOTA_91_PR', 'D_ARY010_72_PR6',
    #    'D_XCC025_72_PR6', 'D_XCC054_72_PR5', 'D_VLW008_72_PR5',
    #    'D_ARY010_72_PR3', 'D_XCC033_72_PR4', 'D_ARY010_72_PR4',
    #    'D_XCC033_72_PR2', 'D_VLW008_72_PR1', 'D_CAA046_71_PA7_PAG',
    #    'D_THRMF_11_NU1_PMI', 'D_WTPCYC_16_PU', 'D_CSB038_OBISPO_PCO',
    #    'D_CSB103_BRBRA_PMI', 'D_CSB103_BRBRA_PCO', 'D_CSB103_BRBRA_PIN',
    #    'D_ESB324_AVEK_PCO', 'D_ESB324_AVEK_PIN', 'D_ESB433_MWDSC_PMI',
    #    'D_ESB433_MWDSC_PCO', 'D_ESB433_MWDSC_PIN', 'D_THRMA_WEC000',
    #    'D_THRMA_RVC000', 'D_FTR039_SEC009', 'D_THRMA_JBC000',
    #    'D_FTR021_16_SA', 'D_FTR014_16_SA', 'D_FTR018_15S_SA',
    #    'D_FTR018_16_SA', 'D_WTPMNR_13_NU1', 'D_OROVL_13_NU1',
    #    'D_FPT013_WTPVNY_CVP', 'D_WTPSAC_26S_PU4_CVP',
    #    'D_FPT013_FSC013_EBMUD', 'D_FPT013_FSC013_CCWD', 'D_CCL005_04_PA1',
    #    'D_TCC022_04_PA2', 'D_TCC036_07N_PA', 'D_TCC081_07S_PA',
    #    'D_CBD028_08S_PA', 'D_CBD049_08N_PA', 'D_KLR005_21_PA',
    #    'DG_08N_PR2', 'DG_12_NU1', 'DG_11_NU1', 'DG_16_PU',
    #    'D_FTR021_16_PA', 'DG_16_SA', 'DG_15S_SA', 'D_SVRWD_CSTLN_PCO',
    #    'D_SVRWD_CSTLN_PMI', 'D_PRRIS_MWDSC', 'D_PRRIS_MWDSC_PCO',
    #    'D_PRRIS_MWDSC_PIN', 'D_PRRIS_MWDSC_PMI', 'D_PYRMD_VNTRA_LCP',
    #    'D_PYRMD_VNTRA_PCO', 'D_PYRMD_VNTRA_PMI', 'D_CSTIC_VNTRA',
    #    'D_CSTIC_VNTRA_PCO', 'D_CSTIC_VNTRA_PIN', 'D_CSTIC_VNTRA_PMI',
    #    'D_BKR004_NBA009_NAPA_PMI', 'D_BKR004_NBA009_NAPA_PCO',
    #    'D_BKR004_NBA009_NAPA_PIN', 'D_BKR004_NBA009_SCWA_PMI',
    #    'D_BKR004_NBA009_SCWA_PCO', 'D_BKR004_NBA009_SCWA_PIN',
    #    'D_CCC019_CCWD', 'DG_13_NU1', 'D_MDOTA_90_PA1', 'D_CAA109_90_PA1',
    #    'D_CAA143_90_PA1', 'D_CAA155_90_PA1', 'D_CAA172_90_PA1',
    #    'D_MDOTA_91_PA', 'DG_73_XA', 'DG_64_XA', 'DG_72_XA2', 'DG_72_XA1',
    #    'DG_91_PR', 'DG_72_PR6', 'DG_72_PR5', 'DG_72_PR3', 'DG_63_PR3',
    #    'DG_72_PR4', 'D_DMC011_71_PA8', 'D_DMC030_71_PA1',
    #    'D_DMC034_71_PA3', 'D_DMC044_71_PA4', 'D_DMC044_71_PA5',
    #    'D_DMC064_71_PA6', 'D_DMC021_50_PA1', 'D_DMC070_73_PA1',
    #    'D_CAA087_73_PA1', 'D_DMC105_73_PA2', 'D_CAA109_73_PA3',
    #    'D_DMC091_73_PA3', 'DG_72_XA3', 'DG_72_PR2', 'DG_72_PR1',
    #    'D_PCH000_SBCWD_PA', 'D_PCH000_SCVWD_PA', 'D_PCH000_SBCWD_PU',
    #    'D_PCH000_SCVWD_PU', 'DG_11_SA1', 'DG_11_SA3',
    #    'D_FOLSM_WTPEDH_CVP', 'D_FOLSM_EDCOCA_CVP', 'D_FOLSM_PCWA_CVP',
    #    'D_FOLSM_WTPFOL_CVP', 'D_FOLSM_WTPRSV_CVP', 'D_CAA046_71_PA7_PIN',
    #    'D_CAA046_71_PA7_PCO', 'D_SBA009_ACFC_PMI', 'D_SBA009_ACFC_PCO',
    #    'D_SBA009_ACFC_PIN', 'D_SBA020_ACFC_PMI', 'D_SBA020_ACFC_PCO',
    #    'D_SBA029_ACWD_PMI', 'D_SBA029_ACWD_PCO', 'D_SBA029_ACWD_PIN',
    #    'D_SBA036_SCVWD_PMI', 'D_SBA036_SCVWD_PCO', 'D_SBA036_SCVWD_PIN',
    #    'D_CAA143_90_PU', 'D_CAA156_90_PU', 'D_CAA165_90_PU',
    #    'D_CAA173_EMPIRE_PAG', 'D_CAA173_EMPIRE_PCO',
    #    'D_CAA173_EMPIRE_PIN', 'D_CAA181_KINGS_PAG', 'D_CAA181_KINGS_PCO',
    #    'D_CAA181_KINGS_PIN', 'D_CAA183_TULARE_PAG', 'D_CAA183_TULARE_PCO',
    #    'D_CAA183_TULARE_PIN', 'D_CAA184_DUDLEY_PAG',
    #    'D_CAA184_DUDLEY_PCO', 'D_CAA184_DUDLEY_PIN', 'D_CAA194_KERN_PAG',
    #    'D_CAA194_KERN_PCO', 'D_CAA194_KERNA_PMI', 'D_CAA194_KERNB_PMI',
    #    'D_CAA238_CVPCV', 'D_CAA239_CVPRF', 'D_CAA242_KERN_PAG',
    #    'D_CAA242_KERN_PCO', 'D_CAA242_KERN_PIN', 'D_CAA279_KERN',
    #    'D_CSB015_KERN_BMWD_PAG', 'D_CSB015_KERN_BMWD_PCO',
    #    'D_CSB009_CLRTA1_DDWD_PAG', 'D_CSB009_CLRTA1_DDWD_PCO',
    #    'D_CSB009_CLRTA1_DDWD_PIN', 'D_CSB038_OBISPO_PCO',
    #    'D_CSB038_OBISPO_PIN', 'D_CSB038_OBISPO_PMI', 'D_CSB103_BRBRA_PCO',
    #    'D_CSB103_BRBRA_PIN', 'D_CSB103_BRBRA_PMI', 'D_ESB324_AVEK_PCO',
    #    'D_ESB324_AVEK_PIN', 'D_ESB324_AVEK_PMI', 'D_ESB347_PLMDL_PCO',
    #    'D_ESB347_PLMDL_PIN', 'D_ESB347_PLMDL_PMI', 'D_ESB355_LROCK_PCO',
    #    'D_ESB355_LROCK_PMI', 'D_ESB403_MOJVE_PCO', 'D_ESB403_MOJVE_PMI',
    #    'D_ESB407_CCHLA_PCO', 'D_ESB407_CCHLA_PIN', 'D_ESB407_CCHLA_PMI',
    #    'D_ESB408_DESRT_PCO', 'D_ESB408_DESRT_PIN', 'D_ESB408_DESRT_PMI',
    #    'D_ESB413_MWDSC_LCP', 'D_ESB413_MWDSC_PCO', 'D_ESB413_MWDSC_PIN',
    #    'D_ESB413_MWDSC_PMI', 'D_ESB414_BRDNO_LCP', 'D_ESB414_BRDNO_PCO',
    #    'D_ESB414_BRDNO_PIN', 'D_ESB414_BRDNO_PMI', 'D_ESB415_GABRL_LCP',
    #    'D_ESB415_GABRL_PCO', 'D_ESB415_GABRL_PIN', 'D_ESB415_GABRL_PMI',
    #    'D_ESB420_GRGNO_LCP', 'D_ESB420_GRGNO_PCO', 'D_ESB420_GRGNO_PIN',
    #    'D_ESB420_GRGNO_PMI', 'D_WSB031_MWDSC_LCP', 'D_WSB031_MWDSC_PCO',
    #    'D_WSB031_MWDSC_PIN', 'D_WSB031_MWDSC_PMI', 'D_WSB032_CLRTA2_LCP',
    #    'D_WSB032_CLRTA2_PCO', 'D_WSB032_CLRTA2_PMI', 'D_FSC015_60N_NA2',
    #    'D_FSC025_60N_PU_CVP', 'PERDV_CVPAG_S', 'PERDV_CVPAG_SYS',
    #    'PERDV_CVPEX_S', 'PERDV_CVPMI_S', 'PERDV_CVPMI_SYS',
    #    'PERDV_CVPRF_SYS', 'PERDV_CVPSC_SYS', 'PERDV_SWP_AG1',
    #    'PERDV_SWP_FSC', 'PERDV_SWP_MWD1', 'SWP_CO_TOTAL', 'SWP_IN_TOTAL',
    #    'SWP_TA_FEATH', 'SWP_TA_NBAY', 'SWP_TA_TOTAL', 'TOTAL_EXP',
    #    'D_MLRTN_FRK_C1', 'D_MLRTN_FRK_C2', 'D_MLRTN_MDC_C1',
    #    'D_MLRTN_MDC_C2', 'D_TCC111_07S_PA', 'D_SAC162_09_SA2']

    #alts = ['NAA', 'Alternative 1', 'Alternative 2a', 'Alternative 2b', 'Alternative 3', 'Alternative 4', 'Alternative 6', 'Alternative 7']

    # Scenarios to compare
    #alts = ["NAA", "Alt1", "Alt2a", 'Alt2b', 'Alt3', 'Alt4', 'Alt6', 'Alt7']
    #alts = ['NAA', "SACLTO_Alt4"]
    #Temperature test
    #fields = ["BLW CLEAR CREEK"]
    #alts = ["NAA", "NAA"]

    #Salinity Test
    #fields = ["SAC_DS_STMBTSL","RSAN007","RSAC075", "RSAC081"] #Test fields for EC DSM2 appendix
    #fields = ['ROLD024','RSAN007'] #Test fields for Cl DSM2 appendix
    # fields = ['X2']
    #alts = ["NAA", "ALT1"]
    """
    Specify whether report is "flow", "elevation', "diversion", or "water supply" (CalSim appendices), "temperature" (HEC-5Q appendix), 
    "EC", "Cl", "Position" (salinity/DSM2 appendices). 
    
    Note 1: "elevation" option also includes storages. "Position" is the X2 position.
    Note 2: Conversion from microSiemens/cm to mg/L Cl uses equation 2 of https://www.waterboards.ca.gov/waterrights/water_issues/programs/bay_delta/california_waterfix/exhibits/docs/ccc_cccwa/CCC-SC_25.pdf
    
    """
    report_type = "elevation"

    #For NAA vs alternative comparison tables, specify whether you want the table captions lumped or not.
    use_lumped_table_captions = False

    #TODO Add selection of units (elevation, temperature, provide both cfs and taf?)
    #Get more info from crosswalk, units, description, etc
    #Add salinity and temperature - could break into separate scripts

    # Prefix for tables and figures in appendix
    appendix_prefix = " F.2.1" #F.2.1 is elevation; F.2.2 is flow; F.2.3 is diversion; F.2.4 is water supply
                                #F.2.5 is DSM2-EC ; F.2.6 is DSM2-X2 (position); F.2.7 is DSM2 - Chloride; F.2.8 is DSM2

    # Path to file with location code crosswalk
    if report_type in ["flow", "elevation", "diversion", "water supply"]:
        #CalSim related appendices use the calsim crosswalk
        location_cw_path = r"..\inputs\location_code_crosswalk_CalSim.xlsx"
    elif report_type in ["EC", "Cl", "Position"]:
        #DSM2 related reports use the salinity crosswalk
        location_cw_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\location_code_crosswalk_salinity.xlsx"
    elif report_type == 'temperature':
        #Temperature related appendices use the temperature crosswalk
        location_cw_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\location_code_crosswalk_Temp.xlsx"

    s_supply_formulas = r"..\inputs\water_supply_formulas.xlsx"

    #Path to file with DSSReader output
    # for water supply, must be the _TAF output
    #Use output from DSS reader in desired units (CFS or TAF). Use TAF for elevation/storage and CFS for the flow and diversion appendices.
    #dss_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\DSS_contents_temperatureTest.xlsx" #Temperature test input
    #dss_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\DSS_contents_CFS.xlsx" #Trinity LTO flow/diversion input
    #dss_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\DSS_contents_TAF.xlsx" #Trinity LTO elevation/storage input
    #dss_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\DSS_contents_salinity_test.xlsx" #Salinity (sample from Sac LTO used for testing)
    #dss_path = r"C:\calsim_gits\eis-appendix-gen_upd\eis-appendix-generation\inputs\DSS_contents_TAF_SacLTOTest.xlsx" #TEST ONLY
    #dss_path = r"C:\Users\cyu\OneDrive - DOI\Documents\TemperatureModeling\temperature_outputs\appendixF\ForDSSReader_temperature_rename.xlsx"#Temperature, all alternatives (No action alternative was manually "renamed" to Baseline in excel.
    #dss_path = r"C:\Users\cyu\OneDrive - DOI\Documents\TemperatureModeling\temperature_outputs\appendixF\ForDSSReader_temperature_rename_alt7_June2017Removed.xlsx"
    dss_path = r"C:\Users\fnufferrodriguez\OneDrive - DOI\Desktop\calsim_dss_reader\DSS_contents.xlsx"
    # dss_path = r"C:\Users\cyu\sacLTO2021\DSS_contents_CombinedSacAmerican.xlsx" #Action 5 run.
    #Path to file with WY Typing data
    wy_flags_path = "..\inputs\wy_flags.xlsx"

    #Path to storage-elevation table data
    storage_elevation_table = r"..\inputs\storage_elevation_table.xlsx"

    # Windows command prompt can't save to OneDrive bc of the space in the file path, save locally instead
    # Pass absolute paths to VBS
    #Name of intermediate word doc - update parent directory
    template = r"..\inputs\template_v2-fonts.docx"
    doc_name = r"..\appendix_temp3.docx"
    #Name of final word doc
    new_doc = rf"..\appendix_final_{report_type}_fixedExceedanceMinMax.docx"

####END OF USER INPUTS #######

    # Read location from crosswalk based on field later
    if report_type == 'elevation':
        #If the report_type is elevation, then order the fields in a specific order. (Ex: S_Trinity storages, S_Trinity elevations, etc).
        fields = order_elevation_storage_fields (fields) #Returns a list of tuples with the type of field (elevation or storage). Ex: [("S_TRNTY", 'Storage'), ("S_TRNTY", 'Elevation'), ("S_SHSTA", 'Storage'),  ("S_SHSTA", 'Elevation')]
    elif report_type in ['EC', 'Position', 'Cl']:
        fields = [(field, report_type) for field in fields]
    locations = get_locations(location_cw_path, fields) #Get location names for each field
    location_params = get_locations_params (location_cw_path, fields) #Get the field parameter for each field (Ex: "Storage", "Elevation", "Diversion", "Delivery")
    locations_wytypes = get_location_wytypes(location_cw_path, fields) #Get the wytype to use with each field.

    #compare every run to the baseline run
    comparisons = [["NAA", alt] for alt in alts]
    #Remove first comparison that is NAA and NAA
    comparisons.pop(0)

    """
    For each field, there are:
        - 3 comparison tables per alternative. (Ex: If you have 6 alternatives, you will have 18 tables total for S_TRNTY)
        - 12 monthly exceedance plots. 
        - full simulation period statistics plots (1 long-term avg plot and 5 plots of averages for different wy types.)
    """

    if report_type == 'water supply':
        # two tables per comparison
        num_tables = 2 * len(comparisons)
        # always 10 tables
        num_figures = 10
    else:
        #Each comparison will have 3 tables and will be included for every field/location
        num_tables = (len(comparisons) * 3) * len(fields)
        #Include a figure for each month plus 6 full simulation period statistics plots
        num_figures = (12 + 6)*len(fields)

    # Alt Text strings, in order for tables
    alt_text_tables = ["Alt text table example" for t in range(0,num_tables)]

    # Alt text strings, order for figures
    alt_text_figures = ["This figure shows data also presented in data tables in this file." for f in range(0,num_figures)]

    #Create an instance of a word document
    #Open the word document template. This template has the heading style 2 formatted with numbering to allow the figures
    #to inherit the heading numbering.
    doc = docx.Document(template)
    doc.add_heading(f"Attachment{appendix_prefix}", level = 1) #Add Heading 1 (Attachment XXX)

    # Add caption style for Figure captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.color.rgb = RGBColor(0,0,0)
    obj_font.name = 'Times New Roman'

    # Add caption style for Table captions
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.color.rgb = RGBColor(0, 0, 0)
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    if report_type == 'water supply':

        # calculate fields
        dfs, df_exceedances = calculate_supply_fields(dss_path, s_supply_formulas, wy_flags_path)

        for comparison_index, comparison in enumerate(comparisons):
            if comparison_index == 0:
                doc.add_page_break()
                change_orientation(doc, "landscape")

            # Add heading for first table
            tab_title_prefix = "Table " + appendix_prefix + "-"
            add_caption_water_supply(doc, "Table", tab_title_prefix, "CalSim 3 Water Summary Report, by Region and Type, Long-Term Average and Dry and Critical Year Averages", custom_style="Table Caption")

            # create table
            df_curr_table = dfs.loc[(comparison, ['Long Term', 'Dry and Critical']), :]
            df_curr_table.loc['Description', dfs.loc['Description', :].columns] = dfs.loc['Description', :].values
            df_curr_table.loc['Units', dfs.loc['Units', :].columns] = dfs.loc['Units', :].values

            # first do the region table
            region_table = df_curr_table[['Sacramento River Hydrologic Region', 'San Joaquin River Hydrologic Region (not including Friant-Kern and Madera Canal water users)',
                                          'San Francisco Bay Hydrologic Region', 'Central Coast Hydrologic Region', 'Tulare Lake Hydrologic Region (not including Friant-Kern Canal water users)',
                                          'South Lahontan Hydrologic Region', 'South Coast Hydrologic Region', 'Total For All Regions']]
            t = doc.add_table(2*region_table.shape[1]+len(region_table.columns.get_level_values(0).unique()) + 1, 7)
            format_table_supply(t, region_table, doc, comparison, [21, 45])

            footnote1 = doc.add_paragraph()
            run = footnote1.add_run('CVP = Central Valley Project; SWP = State Water Project; M&I = municipal and industrial; Ag = Agricultural; FRSA = Feather River Service Allocation;  TAF = thousand acre-feet.')
            run.font.size = Pt(9)
            footnote1.paragraph_format.space_before = Pt(1)
            footnote1.paragraph_format.space_after = Pt(1)

            footnote2 = doc.add_paragraph()
            run = footnote2.add_run('Long-term average is the average quantity for the period of October 1921–September 2021. Dry and critical year average is the average quantity for the combination of the State Water Resources Control Board D-1641 40-30-30 dry and critical dry years for the period of October 1921–September 2021.')
            run.font.size = Pt(9)
            footnote2.paragraph_format.space_before = Pt(1)
            footnote2.paragraph_format.space_after = Pt(1)

            # next we do the north and south table
            # these are split only to get the headers to look good, functionally they are one table
            doc.add_page_break()
            add_caption_water_supply(doc, "Table", tab_title_prefix, "CalSim 3 Water Supply Summary Report, by Type, Long-Term Average and Dry and Critical Year Averages", custom_style="Table Caption")

            north_table = df_curr_table[['North of Delta', 'Total CVP North of Delta', 'Total SWP North of Delta', 'Total North of Delta']]
            t = doc.add_table(2 * north_table.shape[1] + len(north_table.columns.get_level_values(0).unique()) + 1, 7)
            format_table_supply(t, north_table, doc, comparison, [])

            doc.add_page_break()
            south_table = df_curr_table[['South of Delta', 'Total CVP South of Delta', 'Total SWP South of Delta', 'Total South of Delta']]
            t = doc.add_table(2 * south_table.shape[1] + len(south_table.columns.get_level_values(0).unique()) + 1, 7)
            format_table_supply(t, south_table, doc, comparison, [])

            footnote1 = doc.add_paragraph()
            run = footnote1.add_run('CVP = Central Valley Project; SWP = State Water Project; M&I = municipal and industrial; Ag = Agricultural; FRSA = Feather River Service Allocation;  TAF = thousand acre-feet.')
            run.font.size = Pt(9)
            footnote1.paragraph_format.space_before = Pt(1)
            footnote1.paragraph_format.space_after = Pt(1)

            footnote2 = doc.add_paragraph()
            run = footnote2.add_run('Long-term average is the average quantity for the period of October 1921–September 2021. Dry and critical year average is the average quantity for the combination of the State Water Resources Control Board D-1641 40-30-30 dry and critical dry years for the period of October 1921–September 2021.')
            run.font.size = Pt(9)
            footnote2.paragraph_format.space_before = Pt(1)
            footnote2.paragraph_format.space_after = Pt(1)

            doc.add_page_break()

        # create plots
        # Check for/create directory to save plots
        plot_directory = "supply_plots"

        if os.path.exists(plot_directory):
            # If the directory already exists, clear it out (Wytype names are different for trinity vs sjr and sac, so it
            # can cause issues if there's old results.
            shutil.rmtree(plot_directory, ignore_errors=True)

        # WYType Labels to use in stat plot titles. (Corresponds to the "Statistics" column value for the last 6 rows in the exceedance tables)
        fields = df_exceedances.columns
        df_exceedance_list = [df_exceedances.loc[scenario] for scenario in df_exceedances.index.get_level_values(0).unique()]
        fig_value = 'Average Volume (TAF)'
        line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
        line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

        # Iterate through each stat and plot month abbreivated name by EC in current type of year
        for field in fields:
            create_month_plot(df_exceedance_list, fig_value, field, plot_directory, alts, line_styles, line_colors, report_type)


            # Center figures in middle of page by adding some new lines above
            p = doc.add_paragraph()
            run = p.add_run()


            #Add figure as a picture
            o_fig = doc.add_picture(plot_directory + "/" + field + ".png")

            # Generate fig title
            fig_title_prefix = "Figure " + appendix_prefix + "-"
            fig_title = field

            # Add title below figure
            add_caption_water_supply(doc, "Figure", fig_title_prefix, fig_title, custom_style="Figure Caption")

            # if we are on the last plot we don't need a page break
            if field == fields[-1]:
                continue
            else:
                doc.add_page_break()

    else:
        for field_index, location in enumerate(fields):
            if field_index ==0:
                doc.add_page_break()

            #Add location heading in default word heading 2 style. This allows the figure numbering to inherit the heading 2 numbering.
            doc.add_heading(locations[field_index], level=2)

            ##### Read DSSReader output ########
            if report_type == 'elevation':
                #For the elevation report
                if location[1] == "Storage":
                    dfs = parse_dssReader_output(dss_path, alts, location[0], report_type)
                elif location[1] == 'Elevation':
                    # converted to elevation (ft) based on the storage elevation relationship from res_info.table in the CalSim 3 wresl code.
                    dfs = parse_dssReader_output(dss_path, alts, location[0], report_type, convert_to_elevation= True, orig_unit = 'TAF', storage_elevation_fn = storage_elevation_table)
            elif report_type == 'Cl':
                dfs  = parse_dssReader_output(dss_path, alts, location[0], report_type, convert_to_cl= True, orig_unit = 'uS/cm')
            elif report_type in ["EC", "Position"]:
                dfs = parse_dssReader_output(dss_path, alts, location[0], report_type)
            else:
                dfs = parse_dssReader_output(dss_path, alts, location, report_type)



            # Get table value name depending on type of report
            if report_type == "flow":
                unit = 'cfs'
                table_value = "Monthly Flow (cfs)"
            elif report_type == "elevation":
                if location[1] == 'Elevation':
                    unit = 'feet'
                    table_value = "End of Month Elevation (feet)"
                elif location[1] == 'Storage':
                    unit = 'TAF'
                    table_value = 'End of Month Storage (TAF)'
            elif report_type == 'diversion':
                unit = 'cfs'
                table_value = f"Monthly {location_params[field_index]} (cfs)"
            elif report_type == 'temperature':
                unit = 'DEG-F'
                table_value = f"Monthly Temperature (DEG-F)"
            elif report_type == 'EC':
                unit = "UMHOS/CM"
                table_value = f"Monthly EC (UMHOS/CM)"
            elif report_type == 'Cl':
                unit = "mg/L"
                table_value = r"Monthly Cl (mg/L)"
            elif report_type == 'Position':
                unit = "KM"
                table_value = r"Monthly Position (KM)"
            else:
                raise ValueError(f"No report type for {report_type} is available. Needs to be implemented.")

            # Get figure value name depending on type of report. This is used in the stat figure captions.
            fig_value = f"Average {location_params[field_index]} ({unit})"

            #Create Exceedance Tables from DSS Reader output
            e_dfs, exc_prob, fig_dfs,il_num_years= create_exceedance_tables(dfs, wy_flags_path, locations_wytypes[field_index], report_type)

            ##### Use docx package to create a document with formatted table objects and save to Word .docx file ###########

            ## Add a table for each run in each comparison for the current field to the doc
            for comparison_index, scenario in enumerate(comparisons):

                #Then third table for each comparison should be first alt minus second alt listed
                comparison_tables = []
                for alt in scenario:
                    #Get exceedance tables for each of the runs in the current comparison
                    comparison_tables.append(e_dfs[alts.index(alt)])
                #Add one more table for second alt minus the baseline
                comparison_tables.append(comparison_tables[1].iloc[:, 1:] - comparison_tables[0].iloc[:, 1:])
                #Add the labels column back into the differenced table
                comparison_tables[-1].insert(0, "Statistic", comparison_tables[0]["Statistic"])

                #Set up Comparison labels to be used in table titles
                comparison_table_labels = ["NAA", scenario[1], scenario[1] + " Minus " + "NAA"]

                for comp_table_index, full_table in enumerate(comparison_tables):
                    #Subset the statistics table to exclude the lowest and highest probability of exceedance (usually 1% and 99% exceedance)
                    table = full_table.loc[~full_table.Statistic.isin([f'{round(exc_prob.iloc[0])}% Exceedance', f'{round(exc_prob.iloc[-1])}% Exceedance'])].copy(deep = True)
                    table.reset_index(inplace = True,drop = True)

                    table_letter = chr(ord('a') + comp_table_index)
                    #table_number = str(comparison_index + 1) #Track the table number you are currently using.

                    # Generate table title
                    table_title_prefix = "Table" + appendix_prefix + "-"

                    table_title = locations[field_index] + ", " + comparison_table_labels[comp_table_index] + ", " + table_value


                    #table_title = "Table " + appendix_prefix + "-" + str(field_index + 1) + "-" + str(comparison_index + 1) + chr(ord('a') + comp_table_index)  +". " + locations[field_index] + ", " + comparison_table_labels[comp_table_index] + ", " + table_value

                    # Add caption above table
                    if not use_lumped_table_captions:
                        if table_letter == 'a': #If this is the first table of the 3 comparison tables, then use the next sequential table number + the letter a
                            add_caption_byfield(doc, "Table", table_title_prefix, table_letter +". " + table_title, custom_style="Table Caption")
                        else: #If this is not the first table of the 3 comparison tables, use the
                            add_caption_byfield(doc, 'Table', table_title_prefix, table_letter +". " + table_title, custom_style='Table Caption', use_prev_number= True)
                    # p = doc.add_paragraph()
                    # run = p.add_run(table_title)
                    # run.font.bold = True
                    # run.font.size = Pt(12)
                    # p.paragraph_format.space_before = Pt(8)
                    # p.paragraph_format.space_after = Pt(1)

                    # add a table to the end and create a reference variable
                    # extra row is so we can add the header row
                    t = doc.add_table(table.shape[0] + 1, table.shape[1])
                    #Format table for report
                    format_table(t, table, doc, report_type)

                #Get the number of years of simulation record from the full exceedance probability/values dataframe for each of the naa and the alt you are comparing
                il_sample_sizes =[]
                for alt in scenario:
                    #Averaging the number of samples we have for each month
                    # gives you approximation of the full period of record length in years.
                    il_sample_sizes.append(np.mean(il_num_years[alts.index(alt)]).tolist())

                #Determine the period of record footnote to include.
                #If the NAA and alternative you are comparing to have different sample sizes, use this footnote.
                if len(np.unique(il_sample_sizes))!=1:
                    s_por_footnote  = f"{scenario[0]} Statistics based on approximately {round(il_sample_sizes[0], 1)}-year simulation period. {scenario[1]} statistics based on approximately {round(il_sample_sizes[1], 1)}-year simulation period."
                #If the NAA and alternative you are comparing to have the same sample size and it is a whole number of years, then use this footnote.
                elif il_sample_sizes[0] == int(il_sample_sizes[0]):
                    s_por_footnote = f" Based on the {int(il_sample_sizes[0])}-year simulation period."
                # If the NAA and alternative you are comparing to have the same sample size and it includes a fraction of a year, then use this footnote.
                else:
                    s_por_footnote = f" Based on the {round(il_sample_sizes[0],1)}-year simulation period."

                # Add footnotes to the final table
                if comp_table_index == (len(comparison_tables) - 1):
                    # Add footnotes at end of table
                    footnote0 = doc.add_paragraph()
                    run = footnote0.add_run("a")
                    run.font.superscript = True
                    run1 = footnote0.add_run(s_por_footnote)
                    run1.font.size = Pt(9)
                    footnote0.paragraph_format.space_after = Pt(1)

                    footnote1 = doc.add_paragraph()
                    run = footnote1.add_run('* All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
                    run.font.size = Pt(9)
                    footnote1.paragraph_format.space_before = Pt(1)
                    footnote1.paragraph_format.space_after = Pt(1)

                    footnote2 = doc.add_paragraph()
                    if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:
                        run = footnote2.add_run(
                        f'* Water Year Types defined by the {locations_wytypes[field_index]} Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
                    else:
                        run = footnote2.add_run(f"* Water Year Types defined by the Trinity Water Year Hydrologic Classification.")
                    run.font.size = Pt(9)
                    footnote2.paragraph_format.space_before = Pt(1)
                    footnote2.paragraph_format.space_after = Pt(1)

                    #Commented out b/c we are using water years now.
                    footnote3 = doc.add_paragraph()
                    run = footnote3.add_run('* Water Year Types results are displayed with water year – year type sorting.')
                    run.font.size = Pt(9)
                    footnote3.paragraph_format.space_before = Pt(1)
                if comparison_index!=0:
                    doc.add_page_break() #Add page break after the a,b,c comparison tables.

            #####Create Monthly EC and full simulation period statistic plots, save locally as images#####

            #Individual Month Plots tables are in fig_dfs

           #Format percent exceedances for labels
            exc_percents = [str(round(x)).split(".")[0] + "%" for x in exc_prob.values]
            ##Remove simulation period statistic rows
            # for fig_index in range(len(fig_dfs)):
            #     fig_dfs[fig_index] = fig_dfs[fig_index][:-6]
            #     #Add formatted exceedance probability percents back to dfs
            #     fig_dfs[fig_index]["exc_prob"] = exc_percents

            #Can plot up to 8 scenarios, these lines prepare linestyle and color
            line_colors = ["k", "b", "m", "orange", "y", "r", "purple", "g", 'c']
            line_styles = ["-", "-.", "--", "-.", "-.", "--", "-.", "-.", ":"]

            # Flip doc to landscape orientation for images
            change_orientation(doc, "landscape")

            #Iterate through the dfs and create a figure for each month
            #Save month plots to directory
            month_directory = "month_plots"

            if os.path.exists(month_directory):
                # If the directory already exists, clear it out to prevent using any old figures by accident from previous field/alternative.
                shutil.rmtree(month_directory, ignore_errors=True)

            for month in fig_dfs[0].columns[1:]:
                create_month_plot(dfs, fig_value, month, month_directory, alts, line_styles, line_colors)

            ##Simulation Period Statistic Plots###
            stat_fig_dfs = copy.deepcopy(e_dfs)

            for stat_fig_index in range(len(stat_fig_dfs)):
                #keep only simulation period statistic rows
                stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index][-6:]
                #Transpose to plot all months at once
                stat_fig_dfs[stat_fig_index] = stat_fig_dfs[stat_fig_index].transpose()
                #Drop first row
                stat_fig_dfs[stat_fig_index].rename(columns=stat_fig_dfs[stat_fig_index].iloc[0], inplace=True)
                stat_fig_dfs[stat_fig_index].drop(stat_fig_dfs[stat_fig_index].index[0], inplace=True)
                #Add abbreviated month name column
                stat_fig_dfs[stat_fig_index]["month"] = ["Oct", "Nov", "Dec", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug",
                                            "Sep"]

            #Check for/create directory to save stat plots
            stat_directory = "stat_plots"

            if os.path.exists(stat_directory):
                # If the directory already exists, clear it out (Wytype names are different for trinity vs sjr and sac, so it
                # can cause issues if there's old results.
                shutil.rmtree(stat_directory, ignore_errors=True)

            #WYType Labels to use in stat plot titles. (Corresponds to the "Statistics" column value for the last 6 rows in the exceedance tables)
            stats = e_dfs[0].Statistic[-6:].values.tolist()
            # stats = ["Full Simulation Period", "Wet Water Years (28%)", "Above Normal Water Years (14%)",
            #          "Below Normal Water Years (18%)",
            #          "Dry Water Years (24%)", 'Critical Water Years (16%)']

            #Iterate through each stat and plot month abbreivated name by EC in current type of year
            for stat in stats:
                create_stat_plot(stat_fig_dfs, fig_value, stat, stat_directory, alts, line_styles, line_colors)

            ##Add saved figures to docx object as images####

            #Get saved month plots, in order from Oct - Sept.
            month_plots = [rf"{str(m).zfill(2)}_{datetime.datetime.strptime(str(m), '%m').strftime('%b')}_monthly_exceedance.png" for m in [10,11,12,1,2,3,4,5,6,7,8,9]]

            #Iterate through each monthly figure in the month plots directory
            for month_index, file in enumerate(month_plots):
                # Center figures in middle of page by adding some new lines above
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
                run.add_break()



                #Add figure as a picture
                o_fig = doc.add_picture(month_directory + "/" + file)

                # Add captions below figure
                f = doc.add_paragraph()
                run = f.add_run(
                    '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
                run.font.size = Pt(9)
                f.paragraph_format.space_before = Pt(1)
                f.paragraph_format.space_after = Pt(1)

                # Add title below figure (Commented out b/c using add_caption_byfield now)
                # title = doc.add_paragraph()
                # title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # run = title.add_run(fig_title)
                # run.font.size = Pt(12)
                # run.font.bold = True

                # Generate fig title
                fig_title_value = location_params[field_index]
                fig_title_prefix = "Figure " + appendix_prefix + "-"
                fig_title = locations[field_index] + ", " + datetime.datetime.strptime(file.split("_", 2)[1],
                                                                                       '%b').strftime(
                    '%B') + " " + fig_title_value
                # Add title below figure
                add_caption_byfield(doc, "Figure", fig_title_prefix, fig_title, custom_style = "Figure Caption")



                #Add page break after each figure
                doc.add_page_break()

            # Add stats plots as well
            #Set the statistics plot titles
            if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:  #For Sac or SJR WYType
                stat_titles = ["Long Term", "Wet Year", "Above Normal Year", "Below Normal Year", "Dry Year", 'Critical Year']
            else: #For Trinity WYType
                stat_titles = ["Long Term", "Extremely Wet Year", "Wet Year", "Normal Year", "Dry Year", "Critically Dry Year"]

            for stat_plot_index, stat_title in enumerate(stat_titles):
                # Center figures in middle of page by adding some new lines above
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
                run.add_break()

                #Add stat figure as image to document
                file = stat_title[:5] + "_Exceedance.png" if stat_title!= 'Long Term' else "Full _exceedance.png"
                doc.add_picture(stat_directory + "/" + file)

                # Add footnotes below figure about water year type definition
                caption0 = doc.add_paragraph()
                if locations_wytypes[field_index] in ['40-30-30', '60-20-20']:
                    run = caption0.add_run(
                        f'*As defined by the {locations_wytypes[field_index]} Index Water Year Hydrologic Classification (SWRCB D-1641, 1999).')
                else:
                    run = caption0.add_run(
                        f"*As defined by the Trinity Water Year Hydrologic Classification.")
                run.font.size = Pt(9)
                caption0.paragraph_format.space_before = Pt(1)
                caption0.paragraph_format.space_after = Pt(1)

                # Commented out b/c we are using water years now.
                caption1 = doc.add_paragraph()
                run = caption1.add_run('*These results are displayed with water year - year type sorting.')
                run.font.size = Pt(9)
                caption1.paragraph_format.space_before = Pt(1)
                caption1.paragraph_format.space_after = Pt(1)

                #Add footnotes below figure about climate change scenario
                caption2 = doc.add_paragraph()
                run = caption2.add_run(
                    '*All scenarios are simulated at 2022 Median climate condition and 15 cm sea level rise.')
                run.font.size = Pt(9)
                caption2.paragraph_format.space_before = Pt(1)

                # Generate fig title
                fig_title_prefix = "Figure " + appendix_prefix + "."
                fig_title = locations[field_index] + ", " +  stat_title + " " + fig_value
                #Add fig title as the figure caption below figure.
                add_caption_byfield(doc, "Figure", fig_title_prefix, fig_title, custom_style="Figure Caption")

                #No need for the page break if it's the final plot of the document
                if stat_plot_index == (len(stat_titles) - 1) and field_index == (len(fields) - 1):
                    continue
                else:
                    doc.add_page_break()

                #Flip orientation back to portrait for the next group of tables
                if stat_plot_index == (len(stat_titles) - 1):
                    # Flip doc to landscape orientation for images
                    change_orientation(doc, "portrait")

    # Save docx object to word doc
    doc.save(doc_name)

    ##### Use Python to Run VBS Script that adds alt text to table in saved docx file #######

    # Format alt text for all tables as one string to be passed to vbs
    alt_text_string_tables = ("+").join(alt_text_tables)
    alt_text_string_tables = alt_text_string_tables.replace(" ", "_")

    # Format alt text for all figures as one string to be passed to vbs
    alt_text_string_figures = ("+").join(alt_text_figures)
    alt_text_string_figures = alt_text_string_figures.replace(" ", "_")

    #Run vbs script
    #Arguments are existing document, new document to be saved to, alt text for all tables, number of tables, alt text for all figures, number of figures
    #This will fail if Microsoft Word has document open in the background
    #try opening Task Manager and Ending MS Word Background Task, then rerun

    try: #Call the vbs script for table and figure alt text
        result = subprocess.call("cscript.exe add_alt_text.vbs " + doc_name + " " + new_doc + " " + alt_text_string_tables +  " " + str(num_tables) + " " + alt_text_string_figures + " " + str(num_figures))

    except: #If you have too many figures in the document, the above subprocess call will fail. Use workaround where the alt text for tables and figures are called separately.
        #Call script to add table alt text first
        result = subprocess.call("cscript.exe add_alt_text.vbs " + doc_name + " " + doc_name.replace("temp.docx", 'temp2.docx') + " " + alt_text_string_tables + " " + str( num_tables) + " " + "xx"+ " " + str(0))
        #Then call script to add figure alt text
        result = subprocess.call(
            "cscript.exe add_alt_text.vbs " + doc_name.replace("temp.docx", 'temp2.docx') + " " + new_doc + " " + "xxx" + " " + str(0) + " " + alt_text_string_figures + " " + str(num_figures))

    #Remove temporary doc if process ran successfully
    if result == 1:
        print("VBS script did not run successfully. Try using task manager to end MS Word Background Task and then rerun")
    else:
        #Instructions on how to finish formatting numbered captions.
        print("After running this script, \n1. Open Word file and Ctrl+A to select all. Then F9 to update caption numbering.\
        \n2. For the Heading 2 Numbering, you may have to adjust it to match the appendix_prefix variable (Ex: 'F.2.2') by right clicking and selecting 'Adjust List Indents'. \nThen modify the numbering to match appendix_prefix under 'Enter formatting for number:'")
